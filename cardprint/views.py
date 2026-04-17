"""
Card Print Views
================
Page views + API endpoints for the Print Cards workflow:
    Approved → Generate List → Finalized → Pool
  + Generate Card editor page + PDF generation API.

ARCHITECTURE RULES (same as reprint):
- Views are ULTRA-THIN: parse request → call service → return JsonResponse.
- All mutations delegate to PrintWorkflowService / GenerateCardService.
"""
import json
import logging
import re
import uuid
import io
import base64
from typing import List

from django.shortcuts import get_object_or_404, redirect, render
from django.http import HttpResponse, JsonResponse, FileResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_http_methods
from django.db import transaction
from django.db.models import Count, Q, Max, F
from django.utils import timezone
from django.utils.timezone import localtime
from django.utils.dateparse import parse_datetime
from django.core.files.base import ContentFile
from django.urls import reverse

from idcards.models import IDCard, IDCardTable
from core.services.permission_service import PermissionService, api_require_permission
from core.services.activity_service import ActivityService
from core.views.base import get_user_role, require_any_admin
from core.services.super_mode_service import SuperModeService
from core.services import IDCardService
from accounts.rate_limit import rate_limit

from .models import (
    PrintRequest,
    CardTemplate,
    CardTemplateDoc,
    default_template_json,
    validate_field_mappings,
    validate_template_json,
)
from .services import PrintWorkflowService, GenerateCardService

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------

def _print_access_denied():
    return JsonResponse(
        {'status': 'error', 'message': 'Access denied. You are not assigned to this client.'},
        status=403,
    )


def _parse_offset_limit(request, *, default_limit=100, max_limit=200):
    """Parse and clamp offset/limit query params for list endpoints."""
    try:
        offset = int(request.GET.get('offset', 0))
    except (ValueError, TypeError):
        offset = 0
    try:
        limit = int(request.GET.get('limit', default_limit))
    except (ValueError, TypeError):
        limit = default_limit

    offset = max(offset, 0)
    limit = min(max(limit, 1), max_limit)
    return offset, limit


def _normalize_hex_color(value, default='#111111'):
    raw = str(value or '').strip()
    if not raw:
        return default

    v = raw[1:] if raw.startswith('#') else raw
    if re.fullmatch(r'[0-9a-fA-F]{3}', v):
        v = ''.join(ch * 2 for ch in v)
    if not re.fullmatch(r'[0-9a-fA-F]{6}', v):
        return default
    return f'#{v.upper()}'


def _parse_local_datetime_filter(value):
    """Parse datetime-local input safely into an aware datetime."""
    if not value:
        return None
    dt = parse_datetime(value)
    if not dt:
        return None
    if timezone.is_naive(dt):
        dt = timezone.make_aware(dt, timezone.get_current_timezone())
    return dt


def _check_print_table_scope(user, table_id):
    """Check user has access to the client owning this table."""
    table = get_object_or_404(IDCardTable.objects.select_related('group'), id=table_id)
    if not PermissionService.is_super_admin(user):
        staff_profile = getattr(user, 'staff_profile', None)
        if staff_profile and staff_profile.staff_type == 'admin_staff':
            if not staff_profile.assigned_clients.filter(id=table.group.client_id).exists():
                return None, _print_access_denied()
        elif user.role in ('client', 'client_staff'):
            from client.services import ClientAccessService
            if not ClientAccessService.can_access_table(user, table):
                return None, _print_access_denied()
    return table, None


def _build_ordered_fields(table, fd=None, fd_upper=None):
    """Build ordered fields list for a card's field_data against table.fields."""
    if fd is None:
        fd = {}
    if fd_upper is None:
        fd_upper = {k.upper(): v for k, v in fd.items()}
    ordered = []
    for field in table.fields:
        fname = field['name']
        ftype = field.get('type', 'text')
        fval = fd.get(fname, '') or fd_upper.get(fname.upper(), '')
        ordered.append({'name': fname, 'type': ftype, 'value': fval})
    return ordered


def _get_selected_generate_field_names(table, template_obj=None):
    """Return selected field names for generate-list display; fallback to all fields."""
    if template_obj is None:
        template_obj = _get_template_for_table(table, create_default=False)

    all_field_names = [f.get('name') for f in (table.fields or []) if f.get('name')]
    if not template_obj:
        return all_field_names

    cfg = template_obj.field_config or {}
    front = cfg.get('front_fields') or []
    back = cfg.get('back_fields') or [] if cfg.get('is_two_sided') else []

    valid = set(all_field_names)
    selected = []
    for name in front + back:
        if name in valid and name not in selected:
            selected.append(name)

    return selected if selected else all_field_names


def _filter_ordered_fields_by_names(ordered_fields, allowed_names):
    """Filter ordered_fields to the selected names while preserving row order."""
    if not ordered_fields or not allowed_names:
        return ordered_fields
    allowed = set(allowed_names)
    return [f for f in ordered_fields if f.get('name') in allowed]


def _safe_template_pdf_url(file_field):
    """Return a usable template file URL only when the underlying file exists."""
    if not file_field:
        return ''

    name = str(getattr(file_field, 'name', '') or '').strip()
    if not name:
        return ''

    storage = getattr(file_field, 'storage', None)
    try:
        if storage and not storage.exists(name):
            return ''
        return file_field.url
    except Exception:
        logger.warning('Template file missing/unreadable: %s', name, exc_info=True)
        return ''


def _table_templates_qs(table):
    return CardTemplate.objects.filter(table=table).order_by('-is_default', '-is_active', '-version', '-updated_at', '-id')


def _next_template_version(table):
    max_version = CardTemplate.objects.filter(table=table).aggregate(v=Max('version')).get('v') or 0
    return int(max_version) + 1


def _set_default_template(table, template_id):
    CardTemplate.objects.filter(table=table, is_default=True).exclude(id=template_id).update(is_default=False)
    CardTemplate.objects.filter(table=table, id=template_id).update(is_default=True)


def _set_active_template(table, template_id):
    CardTemplate.objects.filter(table=table, is_active=True).exclude(id=template_id).update(is_active=False)
    CardTemplate.objects.filter(table=table, id=template_id).update(is_active=True)


def _deepcopy_json(value, fallback=None):
    try:
        return json.loads(json.dumps(value))
    except Exception:
        return fallback


def _template_field_type_map(table):
    return {
        str(item.get('name') or '').strip(): str(item.get('type') or 'text').strip().lower()
        for item in (table.fields or [])
        if isinstance(item, dict) and str(item.get('name') or '').strip()
    }


def _looks_like_supported_image_src(value):
    raw = str(value or '').strip().lower()
    if not raw:
        return False
    return (
        raw.startswith('data:image/')
        or raw.startswith('/media/')
        or raw.startswith('media/')
        or raw.startswith('http://')
        or raw.startswith('https://')
        or raw.endswith('.png')
        or raw.endswith('.jpg')
        or raw.endswith('.jpeg')
        or raw.endswith('.webp')
        or raw.endswith('.gif')
    )


def _validate_template_json_for_table(table, template_json) -> List[str]:
    errors = []

    schema_err = validate_template_json(template_json)
    if schema_err:
        return [schema_err]

    if not isinstance(template_json, dict):
        return ['template_json must be an object']

    canvas = template_json.get('canvas') if isinstance(template_json.get('canvas'), dict) else {}
    canvas_w = float(canvas.get('width') or 350.0)
    canvas_h = float(canvas.get('height') or 200.0)

    if canvas_w < 100.0 or canvas_w > 1600.0:
        errors.append('template_json.canvas.width must be between 100 and 1600')
    if canvas_h < 60.0 or canvas_h > 1200.0:
        errors.append('template_json.canvas.height must be between 60 and 1200')

    field_type_map = _template_field_type_map(table)
    known_fields = set(field_type_map.keys())

    elements = template_json.get('elements') if isinstance(template_json.get('elements'), list) else []
    for idx, item in enumerate(elements):
        if not isinstance(item, dict):
            errors.append(f'template_json.elements[{idx}] must be an object')
            continue

        elem_type = str(item.get('type') or '').strip().lower()
        field_name = str(item.get('field') or '').strip()

        try:
            x = float(item.get('x'))
            y = float(item.get('y'))
            w = float(item.get('width'))
            h = float(item.get('height'))
        except (TypeError, ValueError):
            errors.append(f'template_json.elements[{idx}] position/size must be numeric')
            continue

        if x < 0 or y < 0 or w <= 0 or h <= 0:
            errors.append(f'template_json.elements[{idx}] has invalid geometry values')
        if (x + w) > canvas_w or (y + h) > canvas_h:
            errors.append(f'template_json.elements[{idx}] must stay within canvas bounds')

        if elem_type in ('text', 'image'):
            if field_name not in known_fields:
                errors.append(f'template_json.elements[{idx}].field "{field_name}" does not exist in table schema')
                continue

            if elem_type == 'image':
                field_type = field_type_map.get(field_name, 'text')
                if not GenerateCardService._is_image_field(field_type, field_name):
                    errors.append(
                        f'template_json.elements[{idx}].field "{field_name}" is not an image-compatible field'
                    )

        if elem_type == 'background':
            src = str(item.get('src') or '').strip()
            if not _looks_like_supported_image_src(src):
                errors.append(f'template_json.elements[{idx}].src is not a supported image source')

    # Keep payload readable for API clients.
    return errors[:20]


def _clone_template_as_new_version(
    base_template,
    *,
    name,
    is_two_sided,
    template_json,
    field_mappings,
    field_config,
    font_size,
    font_family,
    is_default=None,
):
    table = base_template.table
    should_be_default = base_template.is_default if is_default is None else bool(is_default)
    with transaction.atomic():
        CardTemplate.objects.select_for_update().filter(table=table)
        next_version = _next_template_version(table)
        new_template = CardTemplate.objects.create(
            table=table,
            name=str(name or base_template.name or 'Template')[:120],
            version=next_version,
            is_active=True,
            is_default=False,
            parent_template=base_template,
            template_json=template_json,
            front_pdf=base_template.front_pdf,
            back_pdf=base_template.back_pdf,
            is_two_sided=bool(is_two_sided),
            field_config=field_config,
            field_mappings=field_mappings,
            font_size=font_size,
            font_family=font_family,
        )
        _set_active_template(table, new_template.id)
        if should_be_default:
            _set_default_template(table, new_template.id)
    return new_template


def _ensure_default_template(table):
    tmpl = CardTemplate.objects.filter(table=table, is_default=True, is_active=True).order_by('-version', '-id').first()
    if tmpl:
        return tmpl

    fallback = CardTemplate.objects.filter(table=table).order_by('-is_active', '-version', '-updated_at', '-id').first()
    if fallback:
        with transaction.atomic():
            _set_active_template(table, fallback.id)
            _set_default_template(table, fallback.id)
        return CardTemplate.objects.filter(id=fallback.id).first() or fallback

    return CardTemplate.objects.create(
        table=table,
        name='Default Template',
        version=_next_template_version(table),
        is_active=True,
        is_default=True,
        template_json=default_template_json(),
    )


def _get_template_for_table(table, template_id=None, create_default=True):
    if template_id is not None:
        try:
            tid = int(template_id)
        except (TypeError, ValueError):
            return None
        return CardTemplate.objects.filter(table=table, id=tid).first()
    if create_default:
        return _ensure_default_template(table)
    return (
        CardTemplate.objects.filter(table=table, is_active=True).order_by('-is_default', '-version', '-updated_at', '-id').first()
        or _table_templates_qs(table).first()
    )


def _sanitize_template_json(raw):
    src = raw if isinstance(raw, dict) else {}
    canvas_raw = src.get('canvas') if isinstance(src.get('canvas'), dict) else {}

    try:
        canvas_w = float(canvas_raw.get('width') or 350)
    except (TypeError, ValueError):
        canvas_w = 350.0
    try:
        canvas_h = float(canvas_raw.get('height') or 200)
    except (TypeError, ValueError):
        canvas_h = 200.0

    canvas_w = max(100.0, min(1600.0, canvas_w))
    canvas_h = max(60.0, min(1200.0, canvas_h))

    unit = str(canvas_raw.get('unit') or 'px').strip().lower()
    if unit not in ('px',):
        unit = 'px'

    default_real_w = 85.6
    default_real_h = 54.0
    try:
        real_w_mm = float(canvas_raw.get('realWidthMM') or default_real_w)
    except (TypeError, ValueError):
        real_w_mm = default_real_w
    try:
        real_h_mm = float(canvas_raw.get('realHeightMM') or default_real_h)
    except (TypeError, ValueError):
        real_h_mm = default_real_h
    real_w_mm = max(10.0, min(400.0, real_w_mm))
    real_h_mm = max(10.0, min(400.0, real_h_mm))

    try:
        safe_margin = float(canvas_raw.get('safeMargin') or 10.0)
    except (TypeError, ValueError):
        safe_margin = 10.0
    try:
        bleed = float(canvas_raw.get('bleed') or 5.0)
    except (TypeError, ValueError):
        bleed = 5.0
    safe_margin = max(0.0, min(min(canvas_w, canvas_h) / 2.0, safe_margin))
    bleed = max(0.0, min(min(canvas_w, canvas_h) / 2.0, bleed))

    layout_raw = canvas_raw.get('printLayout') if isinstance(canvas_raw.get('printLayout'), dict) else {}
    layout_mode = str(layout_raw.get('mode') or '1').strip().lower()
    if layout_mode not in ('1', '2', '4', 'custom'):
        layout_mode = '1'
    try:
        layout_cols = int(layout_raw.get('columns') or (2 if layout_mode in ('2', '4') else 1))
    except (TypeError, ValueError):
        layout_cols = 2 if layout_mode in ('2', '4') else 1
    try:
        layout_rows = int(layout_raw.get('rows') or (2 if layout_mode == '4' else 1))
    except (TypeError, ValueError):
        layout_rows = 2 if layout_mode == '4' else 1
    layout_cols = max(1, min(12, layout_cols))
    layout_rows = max(1, min(12, layout_rows))

    try:
        margin_mm = float(layout_raw.get('marginMM') or 8.0)
    except (TypeError, ValueError):
        margin_mm = 8.0
    try:
        gap_x_mm = float(layout_raw.get('gapXMM') or 4.0)
    except (TypeError, ValueError):
        gap_x_mm = 4.0
    try:
        gap_y_mm = float(layout_raw.get('gapYMM') or 4.0)
    except (TypeError, ValueError):
        gap_y_mm = 4.0
    margin_mm = max(0.0, min(40.0, margin_mm))
    gap_x_mm = max(0.0, min(40.0, gap_x_mm))
    gap_y_mm = max(0.0, min(40.0, gap_y_mm))

    page_size = str(layout_raw.get('pageSize') or 'a4').strip().lower()
    if page_size not in ('a4',):
        page_size = 'a4'

    out_elements = []
    elements = src.get('elements') if isinstance(src.get('elements'), list) else []
    for item in elements:
        if not isinstance(item, dict):
            continue
        elem_type = str(item.get('type') or 'text').strip().lower()
        if elem_type not in ('text', 'image', 'background', 'rectangle'):
            continue

        field_name = str(item.get('field') or '').strip()
        if elem_type in ('text', 'image') and not field_name:
            continue

        src_value = ''
        if elem_type == 'background':
            src_value = str(item.get('src') or '').strip()
            if not src_value:
                continue

        label = str(item.get('label') or '').strip()[:120]
        side = str(item.get('side') or 'front').strip().lower()
        if side not in ('front', 'back', 'both'):
            side = 'front'

        def _num(key, default):
            try:
                return float(item.get(key, default))
            except (TypeError, ValueError):
                return float(default)

        x_default = 0.0 if elem_type == 'background' else 20.0
        y_default = 0.0 if elem_type == 'background' else 40.0
        width_default = canvas_w if elem_type == 'background' else 120.0
        x = max(0.0, min(canvas_w, _num('x', x_default)))
        y = max(0.0, min(canvas_h, _num('y', y_default)))
        width = max(10.0, min(canvas_w, _num('width', width_default)))
        height_default = canvas_h if elem_type == 'background' else (50.0 if elem_type == 'image' else (40.0 if elem_type == 'rectangle' else 24.0))
        height = max(10.0, min(canvas_h, _num('height', height_default)))

        try:
            font_size = float(item.get('fontSize') or 12.0)
        except (TypeError, ValueError):
            font_size = 12.0
        font_size = max(6.0, min(72.0, font_size))

        align = str(item.get('textAlign') or item.get('text_align') or item.get('align') or 'left').strip().lower()
        if align not in ('left', 'center', 'right'):
            align = 'left'

        font_family = str(item.get('fontFamily') or item.get('font_family') or 'Arial').strip()[:80]
        font_group = str(item.get('fontGroup') or item.get('font_group') or '').strip()[:60]
        font_face = str(item.get('fontFace') or item.get('font_face') or '').strip()[:60]

        font_weight_raw = str(item.get('fontWeight') or item.get('font_weight') or '400').strip().lower()
        if font_weight_raw == 'normal':
            font_weight_raw = '400'
        elif font_weight_raw == 'bold':
            font_weight_raw = '700'
        try:
            font_weight_num = int(round(float(font_weight_raw)))
        except (TypeError, ValueError):
            font_weight_num = 400
        font_weight_num = max(100, min(900, int(round(font_weight_num / 100.0) * 100)))
        font_weight = str(font_weight_num)

        font_style = str(item.get('fontStyle') or item.get('font_style') or 'normal').strip().lower()
        if font_style not in ('normal', 'italic', 'oblique'):
            font_style = 'normal'

        try:
            line_height = float(item.get('lineHeight') or item.get('line_height') or 1.2)
        except (TypeError, ValueError):
            line_height = 1.2
        line_height = max(0.6, min(3.0, line_height))

        try:
            letter_spacing = float(item.get('letterSpacing') or item.get('letter_spacing') or 0.0)
        except (TypeError, ValueError):
            letter_spacing = 0.0
        letter_spacing = max(-10.0, min(20.0, letter_spacing))

        text_mode = str(item.get('textMode') or item.get('text_mode') or 'artistic').strip().lower()
        if text_mode not in ('artistic', 'paragraph'):
            text_mode = 'artistic'

        color = _normalize_hex_color(item.get('color') or '#111111')

        clean_item = {
            'id': str(item.get('id') or uuid.uuid4().hex[:12])[:24],
            'type': elem_type,
            'label': label,
            'field': field_name,
            'x': round(x, 2),
            'y': round(y, 2),
            'width': round(width, 2),
            'height': round(height, 2),
            'fontSize': round(font_size, 2),
            'align': align,
            'textAlign': align,
            'fontFamily': font_family,
            'fontGroup': font_group,
            'fontFace': font_face,
            'fontWeight': font_weight,
            'fontStyle': font_style,
            'lineHeight': round(line_height, 3),
            'letterSpacing': round(letter_spacing, 3),
            'textMode': text_mode,
            'color': color,
            'side': side,
            'showLabel': bool(item.get('showLabel', True)),
            'locked': bool(item.get('locked', elem_type == 'background')),
        }

        if elem_type == 'background':
            clean_item['field'] = ''
            clean_item['label'] = label or 'Background'
            clean_item['src'] = src_value[:2000000]
        elif elem_type == 'rectangle':
            clean_item['field'] = ''
            clean_item['label'] = label or 'Rectangle'

        out_elements.append(clean_item)

        if len(out_elements) >= 500:
            break

    clean = {
        'canvas': {
            'width': round(canvas_w, 2),
            'height': round(canvas_h, 2),
            'unit': unit,
            'realWidthMM': round(real_w_mm, 4),
            'realHeightMM': round(real_h_mm, 4),
            'safeMargin': round(safe_margin, 2),
            'bleed': round(bleed, 2),
            'printLayout': {
                'mode': layout_mode,
                'columns': layout_cols,
                'rows': layout_rows,
                'marginMM': round(margin_mm, 2),
                'gapXMM': round(gap_x_mm, 2),
                'gapYMM': round(gap_y_mm, 2),
                'pageSize': page_size,
            },
        },
        'elements': out_elements,
    }

    err = validate_template_json(clean)
    if err:
        return default_template_json()
    return clean


def _template_list_item(tmpl):
    template_json = tmpl.template_json if isinstance(tmpl.template_json, dict) else default_template_json()
    elements = template_json.get('elements') if isinstance(template_json.get('elements'), list) else []
    return {
        'id': tmpl.id,
        'name': tmpl.name,
        'version': int(getattr(tmpl, 'version', 1) or 1),
        'is_active': bool(getattr(tmpl, 'is_active', True)),
        'is_default': bool(getattr(tmpl, 'is_default', False)),
        'parent_template_id': getattr(tmpl, 'parent_template_id', None),
        'usage_count': int(getattr(tmpl, 'usage_count', 0) or 0),
        'last_used_at': tmpl.last_used_at.isoformat() if getattr(tmpl, 'last_used_at', None) else '',
        'is_two_sided': bool(tmpl.is_two_sided),
        'updated_at': tmpl.updated_at.isoformat() if tmpl.updated_at else '',
        'element_count': len(elements),
    }


def _clamp_float(value, minimum, maximum, default):
    try:
        num = float(value)
    except (TypeError, ValueError):
        return default
    if num < minimum:
        return minimum
    if num > maximum:
        return maximum
    return num


def _sanitize_name_list(values):
    if not isinstance(values, list):
        return []
    out = []
    seen = set()
    for raw in values:
        name = str(raw or '').strip()
        if not name or name in seen:
            continue
        if len(name) > 120:
            name = name[:120]
        out.append(name)
        seen.add(name)
        if len(out) >= 400:
            break
    return out


def _editor_field_config_payload(field_config):
    cfg = field_config if isinstance(field_config, dict) else {}
    orientation = str(cfg.get('card_orientation') or 'landscape').strip().lower()
    if orientation not in ('landscape', 'portrait'):
        orientation = 'landscape'
    return {
        'is_two_sided': bool(cfg.get('is_two_sided', False)),
        'card_orientation': orientation,
        'front_fields': _sanitize_name_list(cfg.get('front_fields') or []),
        'back_fields': _sanitize_name_list(cfg.get('back_fields') or []),
        'doc_page_settings': _sanitize_doc_page_settings(cfg.get('doc_page_settings'), orientation),
    }


def _sanitize_doc_page_settings(raw, orientation='landscape'):
    settings = raw if isinstance(raw, dict) else {}
    safe_orientation = 'portrait' if orientation == 'portrait' else 'landscape'
    card_w_mm, card_h_mm = GenerateCardService.dimensions_for_orientation_mm(safe_orientation)

    margins = settings.get('margins_mm') if isinstance(settings.get('margins_mm'), dict) else {}
    left = _clamp_float(margins.get('left'), 0.0, 25.0, 3.0)
    right = _clamp_float(margins.get('right'), 0.0, 25.0, 3.0)
    top = _clamp_float(margins.get('top'), 0.0, 25.0, 3.0)
    bottom = _clamp_float(margins.get('bottom'), 0.0, 25.0, 3.0)

    max_horizontal = max(4.0, card_w_mm - 8.0)
    if (left + right) > max_horizontal:
        ratio = max_horizontal / max(0.1, left + right)
        left = round(left * ratio, 2)
        right = round(right * ratio, 2)

    max_vertical = max(4.0, card_h_mm - 8.0)
    if (top + bottom) > max_vertical:
        ratio = max_vertical / max(0.1, top + bottom)
        top = round(top * ratio, 2)
        bottom = round(bottom * ratio, 2)

    def _sanitize_mm_list(values, max_mm, max_items=24):
        if isinstance(values, str):
            src = [chunk.strip() for chunk in values.split(',')]
        elif isinstance(values, list):
            src = values
        else:
            src = []

        out = []
        seen = set()
        for raw_val in src:
            val = round(_clamp_float(raw_val, 0.0, max_mm, -1.0), 2)
            if val < 0.0:
                continue
            key = f'{val:.2f}'
            if key in seen:
                continue
            seen.add(key)
            out.append(val)
            if len(out) >= max_items:
                break
        return out

    wrap_mode = str(settings.get('wrap_mode') or 'margin').strip().lower()
    if wrap_mode not in ('margin', 'box'):
        wrap_mode = 'margin'

    return {
        'margins_mm': {
            'left': round(left, 2),
            'right': round(right, 2),
            'top': round(top, 2),
            'bottom': round(bottom, 2),
        },
        'line_gap_mm': round(_clamp_float(settings.get('line_gap_mm'), 0.0, 12.0, 2.5), 2),
        'wrap_mode': wrap_mode,
        'snap_to_guides': bool(settings.get('snap_to_guides', True)),
        'guides_x_mm': _sanitize_mm_list(settings.get('guides_x_mm'), card_w_mm),
        'guides_y_mm': _sanitize_mm_list(settings.get('guides_y_mm'), card_h_mm),
    }


def _sanitize_doc_layout_name(value):
    name = re.sub(r'\s+', ' ', str(value or '').strip())
    if not name:
        return ''
    return name[:80]


MAX_SAVED_DOC_LAYOUTS = 50


def _sanitize_doc_layout_library(raw):
    """Legacy snapshot validator for one-time migration from field_config."""
    if not isinstance(raw, list):
        return []

    out = []
    seen_ids = set()
    for item in raw:
        if not isinstance(item, dict):
            continue
        layout_id = str(item.get('id') or '').strip()
        name = _sanitize_doc_layout_name(item.get('name'))
        saved_at = str(item.get('saved_at') or '').strip()[:40]
        snapshot = item.get('snapshot') if isinstance(item.get('snapshot'), dict) else None
        if not re.fullmatch(r'[A-Za-z0-9_-]{6,40}', layout_id):
            continue
        if not name or not snapshot:
            continue
        if layout_id in seen_ids:
            continue
        seen_ids.add(layout_id)
        out.append({
            'id': layout_id,
            'name': name,
            'saved_at': saved_at,
            'snapshot': snapshot,
        })
        if len(out) >= 20:
            break
    return out


def _saved_doc_layout_meta(tmpl, table_id):
    docs = list(
        CardTemplateDoc.objects
        .filter(template=tmpl)
        .order_by('-updated_at', '-created_at', '-id')[:MAX_SAVED_DOC_LAYOUTS]
    )
    meta = []
    for item in docs:
        download_url = ''
        try:
            download_url = reverse(
                'cardprint:api_template_doc_layout_download',
                kwargs={'table_id': table_id, 'layout_id': item.layout_id},
            )
        except Exception:
            download_url = ''
        meta.append({
            'id': item.layout_id,
            'name': item.name,
            'saved_at': item.updated_at.isoformat() if item.updated_at else '',
            'has_doc_file': bool(item.docx_file),
            'download_url': download_url,
        })
    return meta


def _snapshot_to_docx_bytes(snapshot, name):
    """Build a real DOCX file from the current editor snapshot."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt
    except Exception:
        return None, 'python-docx is not available for DOC save'

    if not isinstance(snapshot, dict):
        return None, 'Invalid snapshot payload'

    orientation = str(snapshot.get('card_orientation') or 'landscape').strip().lower()
    if orientation not in ('landscape', 'portrait'):
        orientation = 'landscape'

    page_w_mm = 87.0 if orientation == 'landscape' else 57.0
    page_h_mm = 57.0 if orientation == 'landscape' else 87.0
    page_settings = _sanitize_doc_page_settings(snapshot.get('doc_page_settings'), orientation)
    margins = page_settings.get('margins_mm') if isinstance(page_settings.get('margins_mm'), dict) else {}

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE if orientation == 'landscape' else WD_ORIENT.PORTRAIT
    section.page_width = Mm(page_w_mm)
    section.page_height = Mm(page_h_mm)
    section.left_margin = Mm(_clamp_float(margins.get('left'), 0.0, 25.0, 3.0))
    section.right_margin = Mm(_clamp_float(margins.get('right'), 0.0, 25.0, 3.0))
    section.top_margin = Mm(_clamp_float(margins.get('top'), 0.0, 25.0, 3.0))
    section.bottom_margin = Mm(_clamp_float(margins.get('bottom'), 0.0, 25.0, 3.0))

    doc.add_heading(str(name or 'Saved DOC'), level=1)
    doc.add_paragraph(f'Card Size: {page_w_mm:.0f}mm x {page_h_mm:.0f}mm')

    def _write_side(side_key, side_label):
        model = snapshot.get(f'editable_design_{side_key}') if isinstance(snapshot.get(f'editable_design_{side_key}'), dict) else None
        if not model:
            return

        doc.add_heading(side_label, level=2)
        lines = [x for x in (model.get('lines') or []) if isinstance(x, dict)]
        lines.sort(key=lambda x: (float(x.get('y_mm') or 0.0), float(x.get('x_mm') or 0.0)))
        for line in lines:
            txt = str(line.get('text') or '').strip()
            if not txt:
                continue
            p = doc.add_paragraph(txt)
            align = str(line.get('text_align') or 'left').strip().lower()
            if align == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            style_font_size = _clamp_float(line.get('font_size_pt'), 6.0, 72.0, 11.0)
            style_line_height = _clamp_float(line.get('line_height'), 0.8, 3.0, 1.15)
            if p.runs:
                run = p.runs[0]
                run.font.size = Pt(style_font_size)
                fam = str(line.get('font_family') or '').strip()
                if fam:
                    run.font.name = fam[:80]
                weight = str(line.get('font_weight') or '400').lower()
                run.bold = weight in ('700', '600', 'bold', 'semibold')

            try:
                p.paragraph_format.line_spacing = style_line_height
            except Exception:
                pass

        images = [x for x in (model.get('images') or []) if isinstance(x, dict)]
        for idx, image in enumerate(images[:8], start=1):
            data_url = str(image.get('data_url') or '').strip()
            if not data_url.startswith('data:image/') or ',' not in data_url:
                continue
            try:
                encoded = data_url.split(',', 1)[1]
                raw = base64.b64decode(encoded)
                width_mm = _clamp_float(image.get('w_mm'), 5.0, page_w_mm - 10.0, 25.0)
                doc.add_paragraph(f'Image {idx}')
                doc.add_picture(io.BytesIO(raw), width=Mm(width_mm))
            except Exception:
                continue

    _write_side('front', 'Front Side')
    if bool(snapshot.get('is_two_sided', False)):
        _write_side('back', 'Back Side')

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue(), None


def _migrate_legacy_doc_layouts(tmpl, user=None):
    """Normalize field_config to the current doc-layout format."""
    cfg = dict(tmpl.field_config) if isinstance(tmpl.field_config, dict) else {}
    if 'doc_layout_library' in cfg:
        cfg.pop('doc_layout_library', None)
        tmpl.field_config = cfg
        tmpl.save(update_fields=['field_config', 'updated_at'])
    return cfg


def _trim_saved_doc_layouts(tmpl):
    stale = list(
        CardTemplateDoc.objects
        .filter(template=tmpl)
        .order_by('-updated_at', '-created_at', '-id')[MAX_SAVED_DOC_LAYOUTS:]
    )
    for item in stale:
        if item.docx_file:
            item.docx_file.delete(save=False)
        item.delete()


def _doc_layout_snapshot_from_template(tmpl):
    cfg = tmpl.field_config if isinstance(tmpl.field_config, dict) else {}
    mappings_raw = tmpl.field_mappings if isinstance(tmpl.field_mappings, dict) else {'front': {}, 'back': {}}
    try:
        mappings = json.loads(json.dumps(mappings_raw))
    except Exception:
        mappings = {'front': {}, 'back': {}}
    orientation = str(cfg.get('card_orientation') or 'landscape').strip().lower()
    if orientation not in ('landscape', 'portrait'):
        orientation = 'landscape'

    front_fields = _sanitize_name_list(cfg.get('front_fields') or [])
    back_fields = _sanitize_name_list(cfg.get('back_fields') or []) if tmpl.is_two_sided else []

    raw_style = cfg.get('docx_text_style') if isinstance(cfg.get('docx_text_style'), dict) else {}
    style_align = str(raw_style.get('align') or 'left').strip().lower()
    if style_align not in ('left', 'center', 'right'):
        style_align = 'left'

    docx_style = {
        'font_family': str(raw_style.get('font_family') or tmpl.font_family or 'Arial').strip()[:80] or 'Arial',
        'font_size_pt': _clamp_float(raw_style.get('font_size_pt'), 6.0, 72.0, float(tmpl.font_size or 11)),
        'line_height': _clamp_float(raw_style.get('line_height'), 0.8, 3.0, 1.15),
        'char_spacing_pt': _clamp_float(raw_style.get('char_spacing_pt'), -5.0, 20.0, 0.0),
        'font_weight': str(raw_style.get('font_weight') or 'normal').strip().lower(),
        'font_color_hex': _normalize_hex_color(raw_style.get('font_color_hex') or '#111111'),
        'align': style_align,
    }
    if docx_style['font_weight'] not in ('normal', 'semibold', 'bold'):
        docx_style['font_weight'] = 'normal'

    return {
        'is_two_sided': bool(tmpl.is_two_sided),
        'card_orientation': orientation,
        'doc_page_settings': _sanitize_doc_page_settings(cfg.get('doc_page_settings'), orientation),
        'front_fields': front_fields,
        'back_fields': back_fields,
        'field_mappings': mappings,
        'font_size': int(max(6, min(72, int(tmpl.font_size or 11)))),
        'font_family': str(tmpl.font_family or 'Arial').strip()[:50] or 'Arial',
        'docx_text_style': docx_style,
        'show_guides': bool(cfg.get('show_guides', True)),
    }


def _apply_doc_layout_snapshot_to_template(tmpl, snapshot):
    if not isinstance(snapshot, dict):
        return 'Invalid saved DOC snapshot'

    is_two_sided = bool(snapshot.get('is_two_sided', False))
    mappings = snapshot.get('field_mappings') if isinstance(snapshot.get('field_mappings'), dict) else {'front': {}, 'back': {}}
    mapping_err = validate_field_mappings(mappings)
    if mapping_err:
        return mapping_err

    tmpl.is_two_sided = is_two_sided
    tmpl.field_mappings = mappings

    try:
        font_size = int(snapshot.get('font_size', tmpl.font_size or 11) or 11)
    except (TypeError, ValueError):
        font_size = int(tmpl.font_size or 11)
    tmpl.font_size = max(6, min(72, font_size))

    font_family = str(snapshot.get('font_family', tmpl.font_family or 'Arial') or '').strip()
    tmpl.font_family = (font_family[:50] if font_family else (tmpl.font_family or 'Arial'))

    cfg = tmpl.field_config if isinstance(tmpl.field_config, dict) else {}
    orientation = str(snapshot.get('card_orientation') or 'landscape').strip().lower()
    if orientation not in ('landscape', 'portrait'):
        orientation = 'landscape'
    cfg['card_orientation'] = orientation
    cfg['is_two_sided'] = is_two_sided
    cfg['doc_page_settings'] = _sanitize_doc_page_settings(snapshot.get('doc_page_settings'), orientation)
    cfg['front_fields'] = _sanitize_name_list(snapshot.get('front_fields') or [])
    cfg['back_fields'] = _sanitize_name_list(snapshot.get('back_fields') or []) if is_two_sided else []
    cfg['show_guides'] = bool(snapshot.get('show_guides', cfg.get('show_guides', True)))

    raw_style = snapshot.get('docx_text_style') if isinstance(snapshot.get('docx_text_style'), dict) else {}
    style_weight = str(raw_style.get('font_weight') or 'normal').strip().lower()
    if style_weight not in ('normal', 'semibold', 'bold'):
        style_weight = 'normal'
    style_align = str(raw_style.get('align') or 'left').strip().lower()
    if style_align not in ('left', 'center', 'right'):
        style_align = 'left'

    cfg['docx_text_style'] = {
        'font_family': str(raw_style.get('font_family') or tmpl.font_family or 'Arial').strip()[:80] or 'Arial',
        'font_size_pt': _clamp_float(raw_style.get('font_size_pt'), 6.0, 72.0, float(tmpl.font_size or 11)),
        'line_height': _clamp_float(raw_style.get('line_height'), 0.8, 3.0, 1.15),
        'char_spacing_pt': _clamp_float(raw_style.get('char_spacing_pt'), -5.0, 20.0, 0.0),
        'font_weight': style_weight,
        'font_color_hex': _normalize_hex_color(raw_style.get('font_color_hex') or '#111111'),
        'align': style_align,
    }

    # TODO: Replace with new JSON-based template editor
    cfg.pop('editable_design_front', None)
    cfg.pop('editable_design_back', None)
    cfg.pop('mapping_confidence', None)

    tmpl.field_config = cfg
    return None


def _template_payload(tmpl):
    field_config = _migrate_legacy_doc_layouts(tmpl)
    card_orientation = field_config.get('card_orientation') or 'landscape'
    if card_orientation not in ('landscape', 'portrait'):
        card_orientation = 'landscape'
    front_pdf_url = _safe_template_pdf_url(tmpl.front_pdf)
    back_pdf_url = _safe_template_pdf_url(tmpl.back_pdf)
    raw_docx_style = field_config.get('docx_text_style') if isinstance(field_config.get('docx_text_style'), dict) else {}
    try:
        style_font_size = float(raw_docx_style.get('font_size_pt') or tmpl.font_size or 11)
    except (TypeError, ValueError):
        style_font_size = float(tmpl.font_size or 11)
    try:
        style_line_height = float(raw_docx_style.get('line_height') or 1.15)
    except (TypeError, ValueError):
        style_line_height = 1.15
    try:
        style_char_spacing = float(raw_docx_style.get('char_spacing_pt') or 0.0)
    except (TypeError, ValueError):
        style_char_spacing = 0.0
    style_weight = str(raw_docx_style.get('font_weight') or 'normal').strip().lower()
    if style_weight not in ('normal', 'semibold', 'bold'):
        style_weight = 'normal'
    style_align = str(raw_docx_style.get('align') or 'left').strip().lower()
    if style_align not in ('left', 'center', 'right'):
        style_align = 'left'
    style_color = _normalize_hex_color(raw_docx_style.get('font_color_hex') or '#111111')
    docx_style = {
        'font_family': str(raw_docx_style.get('font_family') or tmpl.font_family or 'Arial').strip()[:80] or 'Arial',
        'font_size_pt': max(6.0, min(72.0, style_font_size)),
        'line_height': max(0.8, min(3.0, style_line_height)),
        'char_spacing_pt': max(-5.0, min(20.0, style_char_spacing)),
        'font_weight': style_weight,
        'font_color_hex': style_color,
        'text_align': style_align,
    }
    front_fields = field_config.get('front_fields') or []
    back_fields = field_config.get('back_fields') or []
    doc_layout_meta = _saved_doc_layout_meta(tmpl, tmpl.table_id)
    active_doc_layout_id = str(field_config.get('active_doc_layout_id') or '').strip()
    if not any(item['id'] == active_doc_layout_id for item in doc_layout_meta):
        active_doc_layout_id = ''

    return {
        'id': tmpl.id,
        'name': tmpl.name,
        'version': int(getattr(tmpl, 'version', 1) or 1),
        'is_active': bool(getattr(tmpl, 'is_active', True)),
        'is_default': bool(getattr(tmpl, 'is_default', False)),
        'parent_template_id': getattr(tmpl, 'parent_template_id', None),
        'usage_count': int(getattr(tmpl, 'usage_count', 0) or 0),
        'last_used_at': tmpl.last_used_at.isoformat() if getattr(tmpl, 'last_used_at', None) else '',
        'is_two_sided': tmpl.is_two_sided,
        'template_json': _sanitize_template_json(tmpl.template_json),
        'field_mappings': tmpl.field_mappings or {'front': {}, 'back': {}},
        'font_size': tmpl.font_size,
        'font_family': tmpl.font_family,
        'docx_style': docx_style,
        'card_orientation': card_orientation,
        'has_front_pdf': bool(front_pdf_url),
        'has_back_pdf': bool(back_pdf_url),
        'front_pdf_url': front_pdf_url,
        'back_pdf_url': back_pdf_url,
        'front_fields': front_fields,
        'back_fields': back_fields,
        'show_guides': bool(field_config.get('show_guides', True)),
        'doc_page_settings': _sanitize_doc_page_settings(field_config.get('doc_page_settings'), card_orientation),
        'doc_layout_library': doc_layout_meta,
        'active_doc_layout_id': active_doc_layout_id,
    }


# ---------------------------------------------------------------------------
# PAGE VIEW
# ---------------------------------------------------------------------------

@login_required
@require_any_admin
def print_cards(request, table_id):
    """Print Cards workflow page with tabs: Generate List | Finalized."""
    table = get_object_or_404(
        IDCardTable.objects.select_related('group__client'), id=table_id,
    )
    user = request.user
    if not PermissionService.can_access_client(user, table.group.client_id):
        return redirect('manage_clients')

    current_step = request.GET.get('step', 'generate_list')
    if current_step not in ('generate_list', 'finalized'):
        current_step = 'generate_list'

    # Existing field configuration drives which columns are shown in Generate List.
    template_obj = _get_template_for_table(table, create_default=False)
    selected_generate_field_names = _get_selected_generate_field_names(table, template_obj)

    # Step counts for tabs
    counts = PrintRequest.objects.filter(table=table).aggregate(
        fn=Count('id', filter=Q(status='finalized')),
        gl=Count('id', filter=Q(status='generate_list')),
    )
    step_counts = {
        'finalized': counts['fn'],
        'generate_list': counts['gl'],
    }
    approved_count = IDCard.objects.filter(table=table, status='approved').count()

    # Items for the current step
    generate_items = []
    finalized_items = []
    generate_total = 0
    finalized_total = 0

    if current_step == 'generate_list':
        base_qs = PrintRequest.objects.filter(table=table, status='generate_list')
        generate_total = base_qs.count()
        pr_qs = base_qs.select_related('card', 'requested_by').order_by('-updated_at')[:200]
        for idx, pr in enumerate(pr_qs):
            card = pr.card
            fd = card.field_data or {}
            fd_upper = {k.upper(): v for k, v in fd.items()}
            ordered_fields = _build_ordered_fields(table, fd, fd_upper)
            ordered_fields = _filter_ordered_fields_by_names(ordered_fields, selected_generate_field_names)
            req_by = pr.requested_by
            generate_items.append({
                'pr_id': pr.id,
                'card_id': card.id,
                'sr_no': idx + 1,
                'status': card.status,
                'status_display': card.get_status_display(),
                'requested_by_name': req_by.get_full_name() or req_by.username if req_by else 'System',
                'moved_at': localtime(pr.updated_at).strftime('%d %b %Y %H:%M'),
                'ordered_fields': ordered_fields,
            })

    elif current_step == 'finalized':
        base_qs = PrintRequest.objects.filter(table=table, status='finalized')
        finalized_total = base_qs.count()
        pr_qs = base_qs.select_related('card', 'requested_by').order_by('-updated_at')[:200]
        for idx, pr in enumerate(pr_qs):
            card = pr.card
            fd = card.field_data or {}
            fd_upper = {k.upper(): v for k, v in fd.items()}
            ordered_fields = _build_ordered_fields(table, fd, fd_upper)
            req_by = pr.requested_by
            finalized_items.append({
                'pr_id': pr.id,
                'card_id': card.id,
                'sr_no': idx + 1,
                'status': card.status,
                'status_display': card.get_status_display(),
                'requested_by_name': req_by.get_full_name() or req_by.username if req_by else 'System',
                'finalized_at': localtime(pr.updated_at).strftime('%d %b %Y %H:%M'),
                'ordered_fields': ordered_fields,
            })

    # Generate-card editor data (for inline modal)
    template_data = _template_payload(template_obj) if template_obj else {
        'id': None,
        'name': '',
        'is_two_sided': False,
        'template_json': default_template_json(),
        'field_mappings': {'front': {}, 'back': {}},
        'font_size': 11,
        'font_family': 'Arial',
        'docx_style': {
            'font_family': 'Arial',
            'font_size_pt': 11,
            'line_height': 1.15,
            'char_spacing_pt': 0,
            'font_weight': 'normal',
            'font_color_hex': '#111111',
        },
        'card_orientation': 'landscape',
        'has_front_pdf': False,
        'has_back_pdf': False,
        'front_pdf_url': '',
        'back_pdf_url': '',
        'front_fields': [],
        'back_fields': [],
        'doc_page_settings': _sanitize_doc_page_settings({}, 'landscape'),
    }
    front_pdf_url = template_data.get('front_pdf_url') or ''
    back_pdf_url = template_data.get('back_pdf_url') or ''
    field_config = _editor_field_config_payload(template_obj.field_config if template_obj else {})

    import json as _json
    context = {
        'active_page': 'manage_clients',
        'user_role': get_user_role(user),
        'table': table,
        'group': table.group,
        'client': table.group.client,
        'current_step': current_step,
        'step_counts': step_counts,
        'approved_count': approved_count,
        'generate_items': generate_items,
        'finalized_items': finalized_items,
        'generate_total': generate_total,
        'generate_has_more': generate_total > len(generate_items),
        'generate_display_fields': [
            f for f in (table.fields or [])
            if f.get('name') in set(selected_generate_field_names)
        ],
        'finalized_total': finalized_total,
        'finalized_has_more': finalized_total > len(finalized_items),
        'table_fields_json': _json.dumps(table.fields if table.fields else []),
        'field_config_json': _json.dumps(field_config),
        'has_template_front_pdf': bool(front_pdf_url),
        'has_template_back_pdf': bool(back_pdf_url),
        'template_data_json': _json.dumps(template_data),
        'front_pdf_url': front_pdf_url,
        'back_pdf_url': back_pdf_url,
    }
    return render(request, 'cardprint/print-cards.html', context)


# ---------------------------------------------------------------------------
# API VIEWS
# ---------------------------------------------------------------------------

@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_print_send(request, table_id):
    """Send approved cards to generate list.

    Body: { "card_ids": [1, 2, 3] }
    Only cards in 'approved' status are accepted.
    """
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    card_ids = data.get('card_ids', [])
    if not card_ids:
        return JsonResponse({'status': 'error', 'message': 'No cards selected'}, status=400)

    # Only allow cards in 'approved' status
    valid_ids = list(
        IDCard.objects.filter(
            id__in=card_ids, table=table, status='approved',
        ).values_list('id', flat=True)
    )
    if not valid_ids:
        return JsonResponse(
            {'status': 'error', 'message': 'No approved cards found in selection'},
            status=400,
        )

    result = PrintWorkflowService.create_requests(table, valid_ids, request.user)

    if not result.success:
        return JsonResponse({'status': 'error', 'message': result.message}, status=400)

    # Move approved cards → download status so they leave the approve list
    if result.data['created'] > 0:
        IDCard.objects.filter(
            id__in=valid_ids,
            table=table,
            status='approved',
        ).update(status='download')
        for card_id in valid_ids:
            ActivityService.log(
                'card_status',
                'Card moved from Approved to Download (sent to generate list)',
                user=request.user,
                target_model='IDCard',
                target_id=card_id,
                target_name=f'Card #{card_id}',
            )

    return JsonResponse({
        'status': 'ok',
        'message': f"{result.data['created']} card(s) added to generate list"
                   + (f" ({result.data['skipped']} already in list)" if result.data['skipped'] else ''),
        'created': result.data['created'],
        'skipped': result.data['skipped'],
    })


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_print_step_counts(request, table_id):
    """Return step counts for the generate/finalized workflow tabs."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    counts = PrintRequest.objects.filter(table=table).aggregate(
        gl=Count('id', filter=Q(status='generate_list')),
        fn=Count('id', filter=Q(status='finalized')),
        po=Count('id', filter=Q(status='pool')),
    )
    approved_count = IDCard.objects.filter(table=table, status='approved').count()
    payload = {
        'status': 'ok',
        'generate_list': counts['gl'],
        'finalized': counts['fn'],
        'pool': counts['po'],
        'approved': approved_count,
    }
    return JsonResponse(payload)


# ---------------------------------------------------------------------------
# 2-STEP API VIEWS: Generate, Finalized List, Mark Pool, Pool List
# ---------------------------------------------------------------------------


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_print_generate_list(request, table_id):
    """List generate_list items with pagination and search."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    query = request.GET.get('q', '').strip()
    from_raw = (request.GET.get('from') or '').strip()
    to_raw = (request.GET.get('to') or '').strip()
    from_dt = _parse_local_datetime_filter(from_raw)
    to_dt = _parse_local_datetime_filter(to_raw)
    offset, limit = _parse_offset_limit(request, default_limit=100, max_limit=200)

    pr_qs = PrintRequest.objects.filter(
        table=table, status='generate_list',
    ).select_related('card', 'requested_by').order_by('-updated_at')

    selected_generate_field_names = _get_selected_generate_field_names(table)

    if query:
        pr_qs = IDCardService._apply_search_filter(
            pr_qs,
            query,
            table=table,
            json_field='card__field_data',
            id_lookup='card__id',
        )

    if from_dt:
        pr_qs = pr_qs.filter(updated_at__gte=from_dt)
    if to_dt:
        pr_qs = pr_qs.filter(updated_at__lte=to_dt)

    total = pr_qs.count()
    batch = list(pr_qs[offset:offset + limit + 1])
    has_more = len(batch) > limit
    if has_more:
        batch = batch[:limit]

    items = []
    for idx, pr in enumerate(batch):
        card = pr.card
        fd = card.field_data or {}
        fd_upper = {k.upper(): v for k, v in fd.items()}
        ordered_fields = _build_ordered_fields(table, fd, fd_upper)
        ordered_fields = _filter_ordered_fields_by_names(ordered_fields, selected_generate_field_names)
        req_by = pr.requested_by
        items.append({
            'pr_id': pr.id,
            'card_id': card.id,
            'sr_no': offset + idx + 1,
            'status': card.status,
            'status_display': card.get_status_display(),
            'requested_by_name': req_by.get_full_name() or req_by.username if req_by else 'System',
            'moved_at': localtime(pr.updated_at).strftime('%d %b %Y %H:%M'),
            'ordered_fields': ordered_fields,
        })

    payload = {
        'status': 'ok',
        'items': items,
        'total': total,
        'has_more': has_more,
        'offset': offset,
        'limit': limit,
    }
    return JsonResponse(payload)


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_field_config_save(request, table_id):
    """Save field_config (selected fields per side) for a table's CardTemplate.

    Body: {
        "is_two_sided": true/false,
        "front_fields": ["Name", "Photo", ...],
        "back_fields": ["Father Name", ...]
    }
    """
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    is_two_sided = bool(data.get('is_two_sided', False))
    front_fields = data.get('front_fields', [])
    back_fields = data.get('back_fields', []) if is_two_sided else []

    if not front_fields:
        return JsonResponse({'status': 'error', 'message': 'Select at least one front field'}, status=400)

    # Validate field names against table fields
    valid_names = {f['name'] for f in (table.fields or [])}
    for fname in front_fields + back_fields:
        if fname not in valid_names:
            return JsonResponse({'status': 'error', 'message': f'Invalid field: {fname}'}, status=400)

    tmpl = _ensure_default_template(table)
    existing_cfg = tmpl.field_config if isinstance(tmpl.field_config, dict) else {}
    card_orientation = existing_cfg.get('card_orientation') or 'landscape'
    if card_orientation not in ('landscape', 'portrait'):
        card_orientation = 'landscape'
    tmpl.field_config = {
        'is_two_sided': is_two_sided,
        'front_fields': front_fields,
        'back_fields': back_fields,
        'card_orientation': card_orientation,
    }
    tmpl.is_two_sided = is_two_sided
    tmpl.save()

    return JsonResponse({'status': 'ok', 'message': 'Field configuration saved'})


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_finalized_list')
def api_print_finalized_list(request, table_id):
    """List finalized print items with pagination and search."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    query = request.GET.get('q', '').strip()
    from_raw = (request.GET.get('from') or '').strip()
    to_raw = (request.GET.get('to') or '').strip()
    from_dt = _parse_local_datetime_filter(from_raw)
    to_dt = _parse_local_datetime_filter(to_raw)
    offset, limit = _parse_offset_limit(request, default_limit=100, max_limit=200)

    pr_qs = PrintRequest.objects.filter(
        table=table, status='finalized',
    ).select_related('card', 'requested_by').order_by('-updated_at')

    if query:
        pr_qs = IDCardService._apply_search_filter(
            pr_qs,
            query,
            table=table,
            json_field='card__field_data',
            id_lookup='card__id',
        )

    if from_dt:
        pr_qs = pr_qs.filter(updated_at__gte=from_dt)
    if to_dt:
        pr_qs = pr_qs.filter(updated_at__lte=to_dt)

    total = pr_qs.count()
    batch = list(pr_qs[offset:offset + limit + 1])
    has_more = len(batch) > limit
    if has_more:
        batch = batch[:limit]

    items = []
    for idx, pr in enumerate(batch):
        card = pr.card
        fd = card.field_data or {}
        fd_upper = {k.upper(): v for k, v in fd.items()}
        ordered_fields = _build_ordered_fields(table, fd, fd_upper)
        req_by = pr.requested_by
        items.append({
            'pr_id': pr.id,
            'card_id': card.id,
            'sr_no': offset + idx + 1,
            'status': card.status,
            'status_display': card.get_status_display(),
            'requested_by_name': req_by.get_full_name() or req_by.username if req_by else 'System',
            'finalized_at': localtime(pr.updated_at).strftime('%d %b %Y %H:%M'),
            'ordered_fields': ordered_fields,
        })

    payload = {
        'status': 'ok',
        'items': items,
        'total': total,
        'has_more': has_more,
        'offset': offset,
        'limit': limit,
    }
    return JsonResponse(payload)


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_finalized_list')
def api_print_mark_pool(request, table_id):
    """Move finalized items to pool.

    Body: { "request_ids": [1, 2, 3] }
    """
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    request_ids = data.get('request_ids', [])
    if not request_ids:
        return JsonResponse({'status': 'error', 'message': 'No items selected'}, status=400)

    valid_ids = list(
        PrintRequest.objects.filter(
            id__in=request_ids, table=table, status='finalized',
        ).values_list('id', flat=True)
    )
    if not valid_ids:
        return JsonResponse(
            {'status': 'error', 'message': 'No valid finalized items found'},
            status=400,
        )

    result = PrintWorkflowService.bulk_mark_pool(valid_ids, request.user)
    if not result.success:
        return JsonResponse({'status': 'error', 'message': result.message}, status=400)
    return JsonResponse({
        'status': 'ok',
        'message': f"{result.data['updated']} item(s) moved to pool",
        'updated': result.data['updated'],
        'skipped': result.data['skipped'],
    })


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_print_retrieve_generate(request, table_id):
    """Move generate-list items back to approved list.

    Body: { "request_ids": [1, 2, 3] }
    """
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    request_ids = data.get('request_ids', [])
    if not request_ids:
        return JsonResponse({'status': 'error', 'message': 'No items selected'}, status=400)

    prs = list(
        PrintRequest.objects.filter(
            id__in=request_ids,
            table=table,
            status='generate_list',
        ).values('id', 'card_id')
    )
    if not prs:
        return JsonResponse({'status': 'error', 'message': 'No valid generate-list items found'}, status=400)

    valid_request_ids = [row['id'] for row in prs]
    card_ids = [row['card_id'] for row in prs]

    with transaction.atomic():
        # Move cards back to approved list.
        IDCard.objects.filter(
            id__in=card_ids,
            table=table,
        ).update(status='approved')

        # Remove from active print workflow.
        updated = PrintRequest.objects.filter(
            id__in=valid_request_ids,
            table=table,
            status='generate_list',
        ).update(status='pool', updated_at=timezone.now())

    for card_id in card_ids:
        ActivityService.log(
            'card_status',
            'Card moved from Generate List to Approved',
            user=request.user,
            target_model='IDCard',
            target_id=card_id,
            target_name=f'Card #{card_id}',
        )

    return JsonResponse({
        'status': 'ok',
        'message': f'{updated} item(s) moved back to approved list',
        'updated': updated,
        'skipped': len(request_ids) - updated,
    })


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_finalized_list')
def api_print_retrieve_finalized(request, table_id):
    """Move finalized items back to pending list.

    Body: { "request_ids": [1, 2, 3] }
    """
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    request_ids = data.get('request_ids', [])
    if not request_ids:
        return JsonResponse({'status': 'error', 'message': 'No items selected'}, status=400)

    prs = list(
        PrintRequest.objects.filter(
            id__in=request_ids,
            table=table,
            status='finalized',
        ).values('id', 'card_id')
    )
    if not prs:
        return JsonResponse({'status': 'error', 'message': 'No valid finalized items found'}, status=400)

    valid_request_ids = [row['id'] for row in prs]
    card_ids = [row['card_id'] for row in prs]

    with transaction.atomic():
        # Move cards back to pending list.
        IDCard.objects.filter(
            id__in=card_ids,
            table=table,
        ).update(status='pending')

        # Remove from active print workflow.
        updated = PrintRequest.objects.filter(
            id__in=valid_request_ids,
            table=table,
            status='finalized',
        ).update(status='pool', updated_at=timezone.now())

    for card_id in card_ids:
        ActivityService.log(
            'card_status',
            'Card moved from Finalized to Pending',
            user=request.user,
            target_model='IDCard',
            target_id=card_id,
            target_name=f'Card #{card_id}',
        )

    return JsonResponse({
        'status': 'ok',
        'message': f'{updated} item(s) moved back to pending list',
        'updated': updated,
        'skipped': len(request_ids) - updated,
    })


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_print_pool_list(request, table_id):
    """List pool items with pagination and search."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    query = request.GET.get('q', '').strip()
    offset, limit = _parse_offset_limit(request, default_limit=100, max_limit=200)

    pr_qs = PrintRequest.objects.filter(
        table=table, status='pool',
    ).select_related('card', 'requested_by').order_by('-updated_at')

    if query:
        pr_qs = IDCardService._apply_search_filter(
            pr_qs,
            query,
            table=table,
            json_field='card__field_data',
            id_lookup='card__id',
        )

    total = pr_qs.count()
    batch = list(pr_qs[offset:offset + limit + 1])
    has_more = len(batch) > limit
    if has_more:
        batch = batch[:limit]

    items = []
    for idx, pr in enumerate(batch):
        card = pr.card
        fd = card.field_data or {}
        fd_upper = {k.upper(): v for k, v in fd.items()}
        ordered_fields = _build_ordered_fields(table, fd, fd_upper)
        req_by = pr.requested_by
        items.append({
            'pr_id': pr.id,
            'card_id': card.id,
            'sr_no': offset + idx + 1,
            'status': card.status,
            'status_display': card.get_status_display(),
            'requested_by_name': req_by.get_full_name() or req_by.username if req_by else 'System',
            'pool_at': localtime(pr.updated_at).strftime('%d %b %Y %H:%M'),
            'ordered_fields': ordered_fields,
        })

    return JsonResponse({
        'status': 'ok',
        'items': items,
        'total': total,
        'has_more': has_more,
        'offset': offset,
        'limit': limit,
    })

# ===========================================================================
# GENERATE CARD � PAGE VIEWS
# ===========================================================================
@login_required
@require_any_admin
def generate_card(request, table_id):
    """Generate Card editor for a specific table."""
    table = get_object_or_404(
        IDCardTable.objects.select_related('group__client'), id=table_id,
    )
    user = request.user
    if not PermissionService.can_access_client(user, table.group.client_id):
        return redirect('manage_clients')

    template_obj = _ensure_default_template(table)
    generate_count = PrintRequest.objects.filter(table=table, status='generate_list').count()
    field_config = _editor_field_config_payload(template_obj.field_config or {})

    import json as _json
    template_data = _template_payload(template_obj)

    table_fields = table.fields if hasattr(table, 'fields') and table.fields else []
    context = {
        'active_page': 'manage_clients',
        'user_role': get_user_role(user),
        'table': table,
        'group': table.group,
        'client': table.group.client,
        'template': template_obj,
        'template_data_json': _json.dumps(template_data),
        'table_fields_json': _json.dumps(table_fields),
        'field_config_json': _json.dumps(field_config),
        'front_pdf_url': _safe_template_pdf_url(template_obj.front_pdf),
        'back_pdf_url': _safe_template_pdf_url(template_obj.back_pdf),
        'generate_count': generate_count,
    }
    return render(request, 'cardprint/generate-card.html', context)


# ===========================================================================
# GENERATE CARD � API ENDPOINTS
# ===========================================================================

@require_http_methods(["GET", "POST", "PUT"])
@login_required
@api_require_permission('perm_print_list')
def api_templates(request, ref_id):
    """Template list/create/update API.

    GET  /api/templates/<table_id>/       -> list templates for table
    POST /api/templates/<table_id>/       -> create template for table
    PUT  /api/templates/<template_id>/    -> update template by id
    """
    if request.method in ('GET', 'POST'):
        table, err = _check_print_table_scope(request.user, ref_id)
        if err:
            return err

        if request.method == 'GET':
            templates = [_template_list_item(t) for t in _table_templates_qs(table)]
            return JsonResponse({'status': 'ok', 'templates': templates})

        try:
            data = json.loads(request.body or '{}')
        except (json.JSONDecodeError, ValueError):
            return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

        name = str((data or {}).get('name') or '').strip()[:120]
        if not name:
            name = f'Template {timezone.now().strftime("%H%M%S")}'

        template_json = _sanitize_template_json((data or {}).get('template_json') or default_template_json())
        validation_errors = _validate_template_json_for_table(table, template_json)
        if validation_errors:
            return JsonResponse({'status': 'error', 'message': '; '.join(validation_errors)}, status=400)

        is_two_sided = bool((data or {}).get('is_two_sided', False))
        orientation = str((data or {}).get('card_orientation') or 'landscape').strip().lower()
        if orientation not in ('landscape', 'portrait'):
            orientation = 'landscape'
        font_family = str((data or {}).get('font_family') or 'Arial').strip()[:50] or 'Arial'
        try:
            font_size = int((data or {}).get('font_size', 11) or 11)
        except (TypeError, ValueError):
            font_size = 11
        font_size = max(6, min(72, font_size))

        requested_default = data.get('is_default')

        with transaction.atomic():
            CardTemplate.objects.select_for_update().filter(table=table)
            is_default = bool(requested_default) if requested_default is not None else not CardTemplate.objects.filter(table=table, is_default=True).exists()
            tmpl = CardTemplate.objects.create(
                table=table,
                name=name,
                version=_next_template_version(table),
                is_active=True,
                is_default=False,
                is_two_sided=is_two_sided,
                template_json=template_json,
                font_family=font_family,
                font_size=font_size,
                field_config={
                    'card_orientation': orientation,
                    'is_two_sided': is_two_sided,
                },
            )
            _set_active_template(table, tmpl.id)
            if is_default:
                _set_default_template(table, tmpl.id)

        return JsonResponse({
            'status': 'ok',
            'message': 'Template created',
            'template': _template_payload(tmpl),
            'templates': [_template_list_item(t) for t in _table_templates_qs(table)],
            'active_template_id': tmpl.id,
        })

    tmpl = CardTemplate.objects.filter(id=ref_id).select_related('table__group').first()
    if not tmpl:
        return JsonResponse({'status': 'error', 'message': 'Template not found'}, status=404)

    table, err = _check_print_table_scope(request.user, tmpl.table_id)
    if err:
        return err

    try:
        data = json.loads(request.body or '{}')
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    updated_name = str(data.get('name') or tmpl.name or '').strip()[:120]
    if not updated_name:
        return JsonResponse({'status': 'error', 'message': 'Template name is required'}, status=400)

    is_two_sided = bool(data.get('is_two_sided', tmpl.is_two_sided))
    updated_font_family = str(data.get('font_family') or tmpl.font_family or 'Arial').strip()[:50] or 'Arial'
    try:
        updated_font_size = int(data.get('font_size', tmpl.font_size or 11) or 11)
    except (TypeError, ValueError):
        updated_font_size = int(tmpl.font_size or 11)
    updated_font_size = max(6, min(72, updated_font_size))

    raw_template_json = data.get('template_json', tmpl.template_json)
    updated_template_json = _sanitize_template_json(raw_template_json)
    validation_errors = _validate_template_json_for_table(table, updated_template_json)
    if validation_errors:
        return JsonResponse({'status': 'error', 'message': '; '.join(validation_errors)}, status=400)

    updated_cfg = _deepcopy_json(tmpl.field_config, fallback={}) if isinstance(tmpl.field_config, dict) else {}
    updated_orientation = str(data.get('card_orientation') or updated_cfg.get('card_orientation') or 'landscape').strip().lower()
    if updated_orientation not in ('landscape', 'portrait'):
        updated_orientation = 'landscape'
    updated_cfg['card_orientation'] = updated_orientation
    updated_cfg['is_two_sided'] = is_two_sided

    updated_mappings = _deepcopy_json(tmpl.field_mappings, fallback={'front': {}, 'back': {}}) if isinstance(tmpl.field_mappings, dict) else {'front': {}, 'back': {}}
    if 'field_mappings' in data:
        mapping_err = validate_field_mappings(data.get('field_mappings'))
        if mapping_err:
            return JsonResponse({'status': 'error', 'message': mapping_err}, status=400)
        updated_mappings = data.get('field_mappings')

    requested_default = data.get('is_default')
    new_template = _clone_template_as_new_version(
        tmpl,
        name=updated_name,
        is_two_sided=is_two_sided,
        template_json=updated_template_json,
        field_mappings=updated_mappings,
        field_config=updated_cfg,
        font_size=updated_font_size,
        font_family=updated_font_family,
        is_default=requested_default if requested_default is not None else tmpl.is_default,
    )

    return JsonResponse({
        'status': 'ok',
        'message': 'Template updated (new version created)',
        'template': _template_payload(new_template),
        'templates': [_template_list_item(t) for t in _table_templates_qs(table)],
        'active_template_id': new_template.id,
    })


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_template_detail(request, template_id):
    """Return full details for a single template."""
    tmpl = CardTemplate.objects.filter(id=template_id).first()
    if not tmpl:
        return JsonResponse({'status': 'error', 'message': 'Template not found'}, status=404)

    _table, err = _check_print_table_scope(request.user, tmpl.table_id)
    if err:
        return err

    return JsonResponse({'status': 'ok', 'template': _template_payload(tmpl)})


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_template_duplicate(request, template_id):
    """Duplicate a template into a new versioned template row."""
    base = CardTemplate.objects.filter(id=template_id).select_related('table__group').first()
    if not base:
        return JsonResponse({'status': 'error', 'message': 'Template not found'}, status=404)

    table, err = _check_print_table_scope(request.user, base.table_id)
    if err:
        return err

    try:
        data = json.loads(request.body or '{}')
    except (json.JSONDecodeError, ValueError):
        data = {}

    duplicate_name = str(data.get('name') or f'{base.name} Copy').strip()[:120] or f'{base.name} Copy'
    activate = bool(data.get('activate', False))
    make_default = bool(data.get('is_default', False))

    with transaction.atomic():
        CardTemplate.objects.select_for_update().filter(table=table)
        new_template = CardTemplate.objects.create(
            table=table,
            name=duplicate_name,
            version=_next_template_version(table),
            is_active=activate,
                is_default=False,
            parent_template=base,
            template_json=_deepcopy_json(base.template_json, fallback=default_template_json()) or default_template_json(),
            front_pdf=base.front_pdf,
            back_pdf=base.back_pdf,
            is_two_sided=bool(base.is_two_sided),
            field_config=_deepcopy_json(base.field_config, fallback={}) or {},
            field_mappings=_deepcopy_json(base.field_mappings, fallback={'front': {}, 'back': {}}) or {'front': {}, 'back': {}},
            font_size=int(base.font_size or 11),
            font_family=str(base.font_family or 'Arial')[:50] or 'Arial',
        )
        if activate:
            _set_active_template(table, new_template.id)
        if make_default:
            _set_default_template(table, new_template.id)

    return JsonResponse({
        'status': 'ok',
        'message': 'Template duplicated',
        'template': _template_payload(new_template),
        'templates': [_template_list_item(t) for t in _table_templates_qs(table)],
        'active_template_id': new_template.id if activate else (_get_template_for_table(table, create_default=True).id),
    })


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_template_set_default(request, template_id):
    """Mark a template as default for its table."""
    tmpl = CardTemplate.objects.filter(id=template_id).select_related('table__group').first()
    if not tmpl:
        return JsonResponse({'status': 'error', 'message': 'Template not found'}, status=404)

    table, err = _check_print_table_scope(request.user, tmpl.table_id)
    if err:
        return err

    with transaction.atomic():
        CardTemplate.objects.select_for_update().filter(table=table)
        _set_default_template(table, tmpl.id)
        _set_active_template(table, tmpl.id)

    refreshed = CardTemplate.objects.filter(id=tmpl.id).first()
    return JsonResponse({
        'status': 'ok',
        'message': 'Default template updated',
        'template': _template_payload(refreshed or tmpl),
        'templates': [_template_list_item(t) for t in _table_templates_qs(table)],
        'active_template_id': tmpl.id,
    })

@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_template_get(request, table_id):
    """Return template settings for a table.

    Optional query param: template_id
    """
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    template_id_raw = request.GET.get('template_id')
    tmpl = _get_template_for_table(table, template_id=template_id_raw, create_default=True)
    if not tmpl:
        return JsonResponse({'status': 'error', 'message': 'Template not found for this table'}, status=404)

    templates = [_template_list_item(t) for t in _table_templates_qs(table)]
    return JsonResponse({
        'status': 'ok',
        'template': _template_payload(tmpl),
        'templates': templates,
        'active_template_id': tmpl.id,
    })


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_template_save(request, table_id):
    """Save table template settings for PDF generation."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    template_id = data.get('template_id')
    tmpl = _get_template_for_table(table, template_id=template_id, create_default=True)
    if not tmpl:
        return JsonResponse({'status': 'error', 'message': 'Template not found for this table'}, status=404)

    updated_name = str(data.get('name') or tmpl.name or '').strip()[:120] or tmpl.name
    updated_is_two_sided = bool(data.get('is_two_sided', False))

    raw_template_json = data.get('template_json')
    updated_template_json = _sanitize_template_json(raw_template_json if raw_template_json is not None else tmpl.template_json)
    validation_errors = _validate_template_json_for_table(table, updated_template_json)
    if validation_errors:
        return JsonResponse({'status': 'error', 'message': '; '.join(validation_errors)}, status=400)

    updated_mappings = _deepcopy_json(tmpl.field_mappings, fallback={'front': {}, 'back': {}}) if isinstance(tmpl.field_mappings, dict) else {'front': {}, 'back': {}}
    raw_mappings = data.get('field_mappings')
    if raw_mappings is not None:
        mapping_err = validate_field_mappings(raw_mappings)
        if mapping_err:
            return JsonResponse({'status': 'error', 'message': mapping_err}, status=400)
        updated_mappings = raw_mappings

    try:
        updated_font_size = int(data.get('font_size', tmpl.font_size or 11) or 11)
    except (TypeError, ValueError):
        updated_font_size = int(tmpl.font_size or 11)
    updated_font_size = max(6, min(72, updated_font_size))

    updated_font_family = str(data.get('font_family', tmpl.font_family or 'Arial') or '').strip()
    updated_font_family = (updated_font_family[:50] if updated_font_family else (tmpl.font_family or 'Arial'))

    card_orientation = data.get('card_orientation', 'landscape')
    if card_orientation not in ('landscape', 'portrait'):
        card_orientation = 'landscape'

    cfg = _deepcopy_json(tmpl.field_config, fallback={}) if isinstance(tmpl.field_config, dict) else {}
    front_fields = _sanitize_name_list(data.get('front_fields', cfg.get('front_fields') or []))
    back_fields = _sanitize_name_list(data.get('back_fields', cfg.get('back_fields') or [])) if updated_is_two_sided else []
    cfg['front_fields'] = front_fields
    cfg['back_fields'] = back_fields
    cfg['card_orientation'] = card_orientation
    cfg['is_two_sided'] = updated_is_two_sided
    cfg['show_guides'] = bool(data.get('show_guides', cfg.get('show_guides', True)))
    cfg['doc_page_settings'] = _sanitize_doc_page_settings(
        data.get('doc_page_settings', cfg.get('doc_page_settings')),
        card_orientation,
    )

    try:
        style_font_size = float(data.get('docx_font_size_pt', updated_font_size) or updated_font_size)
    except (TypeError, ValueError):
        style_font_size = float(updated_font_size)
    style_font_size = max(6.0, min(72.0, style_font_size))

    try:
        line_height = float(data.get('docx_line_height', 1.15) or 1.15)
    except (TypeError, ValueError):
        line_height = 1.15
    line_height = max(0.8, min(3.0, line_height))

    try:
        char_spacing = float(data.get('docx_char_spacing_pt', 0.0) or 0.0)
    except (TypeError, ValueError):
        char_spacing = 0.0
    char_spacing = max(-5.0, min(20.0, char_spacing))

    style_family = str(data.get('docx_font_family', updated_font_family or 'Arial') or '').strip()[:80]
    if not style_family:
        style_family = 'Arial'
    style_weight = str(data.get('docx_font_weight', 'normal') or 'normal').strip().lower()
    if style_weight not in ('normal', 'semibold', 'bold'):
        style_weight = 'normal'
    style_align = str(data.get('docx_text_align', 'left') or 'left').strip().lower()
    if style_align not in ('left', 'center', 'right'):
        style_align = 'left'
    style_color = _normalize_hex_color(data.get('docx_font_color_hex', '#111111'))

    cfg['docx_text_style'] = {
        'font_family': style_family,
        'font_size_pt': style_font_size,
        'line_height': line_height,
        'char_spacing_pt': char_spacing,
        'font_weight': style_weight,
        'font_color_hex': style_color,
        'align': style_align,
    }

    # TODO: Replace with new JSON-based template editor
    cfg.pop('editable_design_front', None)
    cfg.pop('editable_design_back', None)
    cfg.pop('mapping_confidence', None)

    requested_default = data.get('is_default')
    new_template = _clone_template_as_new_version(
        tmpl,
        name=updated_name,
        is_two_sided=updated_is_two_sided,
        template_json=updated_template_json,
        field_mappings=updated_mappings,
        field_config=cfg,
        font_size=updated_font_size,
        font_family=updated_font_family,
        is_default=requested_default if requested_default is not None else tmpl.is_default,
    )

    return JsonResponse({
        'status': 'ok',
        'message': 'Template settings saved (new version created)',
        'template': _template_payload(new_template),
        'templates': [_template_list_item(t) for t in _table_templates_qs(table)],
        'active_template_id': new_template.id,
    })


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_template_doc_layout_save(request, table_id):
    """Save current editor state as a named DOC with persisted DOCX file."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body or '{}')
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    name = _sanitize_doc_layout_name((data or {}).get('name'))
    if not name:
        return JsonResponse({'status': 'error', 'message': 'Please enter a DOC name'}, status=400)

    tmpl = _ensure_default_template(table)
    cfg = _migrate_legacy_doc_layouts(tmpl, request.user)
    snapshot = _doc_layout_snapshot_from_template(tmpl)

    try:
        snapshot_size = len(json.dumps(snapshot))
    except Exception:
        snapshot_size = 0
    if snapshot_size > 2_500_000:
        return JsonResponse(
            {'status': 'error', 'message': 'DOC snapshot is too large. Reduce embedded images and try again.'},
            status=400,
        )

    layout_id = uuid.uuid4().hex[:12]
    doc_bytes, doc_err = _snapshot_to_docx_bytes(snapshot, name)
    if doc_err:
        return JsonResponse({'status': 'error', 'message': doc_err}, status=500)

    doc_obj = CardTemplateDoc(
        template=tmpl,
        layout_id=layout_id,
        name=name,
        snapshot=snapshot,
        created_by=request.user,
    )
    filename = f'table_{table.id}_{layout_id}.docx'
    doc_obj.docx_file.save(filename, ContentFile(doc_bytes), save=False)
    doc_obj.save()

    cfg.pop('doc_layout_library', None)
    cfg['active_doc_layout_id'] = layout_id
    tmpl.field_config = cfg
    tmpl.save(update_fields=['field_config', 'updated_at'])
    _trim_saved_doc_layouts(tmpl)

    return JsonResponse({'status': 'ok', 'message': 'DOC saved', 'template': _template_payload(tmpl)})


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_template_doc_layout_list(request, table_id):
    """Return persisted saved DOC list for the current table template."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    tmpl = _ensure_default_template(table)
    cfg = _migrate_legacy_doc_layouts(tmpl, request.user)
    docs = _saved_doc_layout_meta(tmpl, table.id)
    active_doc_layout_id = str(cfg.get('active_doc_layout_id') or '').strip()
    if not any(item['id'] == active_doc_layout_id for item in docs):
        active_doc_layout_id = ''

    return JsonResponse({'status': 'ok', 'docs': docs, 'active_doc_layout_id': active_doc_layout_id})


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_template_doc_layout_apply(request, table_id, layout_id):
    """Apply a previously saved DOC snapshot to the current template."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    tmpl = _ensure_default_template(table)
    cfg = _migrate_legacy_doc_layouts(tmpl, request.user)
    selected = CardTemplateDoc.objects.filter(template=tmpl, layout_id=str(layout_id or '').strip()).first()
    if not selected:
        return JsonResponse({'status': 'error', 'message': 'Saved DOC not found'}, status=404)

    snapshot = selected.snapshot if isinstance(selected.snapshot, dict) else None
    apply_err = _apply_doc_layout_snapshot_to_template(tmpl, snapshot)
    if apply_err:
        return JsonResponse({'status': 'error', 'message': apply_err}, status=400)

    updated_cfg = dict(tmpl.field_config) if isinstance(tmpl.field_config, dict) else {}
    updated_cfg.pop('doc_layout_library', None)
    updated_cfg['active_doc_layout_id'] = selected.layout_id
    tmpl.field_config = updated_cfg
    tmpl.save()
    selected.save(update_fields=['updated_at'])

    return JsonResponse({'status': 'ok', 'message': 'DOC loaded', 'template': _template_payload(tmpl)})


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_template_doc_layout_download(request, table_id, layout_id):
    """Download a saved DOCX file for a previously saved editor DOC layout."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    tmpl = _ensure_default_template(table)
    _migrate_legacy_doc_layouts(tmpl, request.user)

    selected = CardTemplateDoc.objects.filter(template=tmpl, layout_id=str(layout_id or '').strip()).first()
    if not selected:
        return JsonResponse({'status': 'error', 'message': 'Saved DOC not found'}, status=404)

    name = re.sub(r'[^A-Za-z0-9_-]+', '_', str(selected.name or 'saved_doc')).strip('_') or 'saved_doc'
    filename = f'{name}.docx'

    file_missing = not selected.docx_file or not selected.docx_file.name
    if not file_missing:
        try:
            file_missing = not selected.docx_file.storage.exists(selected.docx_file.name)
        except Exception:
            file_missing = True

    if file_missing:
        snapshot = selected.snapshot if isinstance(selected.snapshot, dict) else None
        if not snapshot:
            return JsonResponse({'status': 'error', 'message': 'Saved DOC file is unavailable'}, status=404)
        doc_bytes, doc_err = _snapshot_to_docx_bytes(snapshot, selected.name)
        if doc_err:
            return JsonResponse({'status': 'error', 'message': doc_err}, status=500)
        regen_name = f'table_{table.id}_{selected.layout_id}.docx'
        selected.docx_file.save(regen_name, ContentFile(doc_bytes), save=False)
        selected.save(update_fields=['docx_file', 'updated_at'])

    response = FileResponse(selected.docx_file.open('rb'), as_attachment=True, filename=filename)
    response.block_size = SuperModeService.download_block_size_bytes(request.user)
    return response


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_template_upload_pdf(request, table_id, side):
    """Upload front or back design PDF file."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    if side not in ('front', 'back'):
        return JsonResponse({'status': 'error', 'message': 'Invalid side'}, status=400)

    template_file = request.FILES.get('pdf') or request.FILES.get('template')
    if not template_file:
        return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)

    if not template_file.name.lower().endswith('.pdf'):
        return JsonResponse({'status': 'error', 'message': 'File must be a design PDF (.pdf)'}, status=400)

    if template_file.size > 20 * 1024 * 1024:
        return JsonResponse({'status': 'error', 'message': 'File too large (max 20 MB)'}, status=400)

    tmpl = _ensure_default_template(table)
    cfg = dict(tmpl.field_config) if isinstance(tmpl.field_config, dict) else {}
    if side == 'front':
        if tmpl.front_pdf:
            tmpl.front_pdf.delete(save=False)
        tmpl.front_pdf = template_file
        cfg.pop('editable_design_front', None)
    else:
        if tmpl.back_pdf:
            tmpl.back_pdf.delete(save=False)
        tmpl.back_pdf = template_file
        cfg.pop('editable_design_back', None)
    tmpl.field_config = cfg
    tmpl.save()

    template_url = (tmpl.front_pdf.url if side == 'front' else tmpl.back_pdf.url)
    return JsonResponse({'status': 'ok', 'message': 'Design PDF uploaded', 'pdf_url': template_url})


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
def api_template_clear_pdf(request, table_id, side):
    """Clear saved front/back design PDF template."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    if side not in ('front', 'back'):
        return JsonResponse({'status': 'error', 'message': 'Invalid side'}, status=400)

    tmpl = _ensure_default_template(table)
    cfg = dict(tmpl.field_config) if isinstance(tmpl.field_config, dict) else {}
    mappings = dict(tmpl.field_mappings) if isinstance(tmpl.field_mappings, dict) else {}

    if side == 'front':
        if tmpl.front_pdf:
            tmpl.front_pdf.delete(save=False)
        tmpl.front_pdf = None
        cfg.pop('editable_design_front', None)
    else:
        if tmpl.back_pdf:
            tmpl.back_pdf.delete(save=False)
        tmpl.back_pdf = None
        cfg.pop('editable_design_back', None)

    if isinstance(mappings.get(side), dict):
        mappings[side] = {}
    else:
        mappings.pop(side, None)

    cfg.pop('mapping_confidence', None)
    tmpl.field_config = cfg
    tmpl.field_mappings = mappings
    tmpl.save(update_fields=['front_pdf', 'back_pdf', 'field_config', 'field_mappings', 'updated_at'])
    return JsonResponse({'status': 'ok', 'message': f'{side.capitalize()} design PDF cleared', 'template': _template_payload(tmpl)})


@require_http_methods(["GET"])
@login_required
@api_require_permission('perm_print_list')
def api_generate_card_list(request, table_id):
    """List cards in generate_list status for this table (paginated + searchable)."""
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    query = request.GET.get('q', '').strip()
    offset, limit = _parse_offset_limit(request, default_limit=500, max_limit=500)

    pr_qs = PrintRequest.objects.filter(
        table=table, status='generate_list',
    ).select_related('card', 'requested_by').order_by('created_at')

    selected_generate_field_names = _get_selected_generate_field_names(table)

    if query:
        pr_qs = IDCardService._apply_search_filter(
            pr_qs,
            query,
            table=table,
            json_field='card__field_data',
            id_lookup='card__id',
        )

    total = pr_qs.count()
    batch = list(pr_qs[offset:offset + limit])

    items = []
    for idx, pr in enumerate(batch):
        card = pr.card
        fd = card.field_data or {}
        fd_upper = {k.upper(): v for k, v in fd.items()}
        ordered_fields = _build_ordered_fields(table, fd, fd_upper)
        ordered_fields = _filter_ordered_fields_by_names(ordered_fields, selected_generate_field_names)
        items.append({
            'pr_id': pr.id,
            'card_id': card.id,
            'sr_no': offset + idx + 1,
            'ordered_fields': ordered_fields,
        })

    return JsonResponse({
        'status': 'ok',
        'items': items,
        'total': total,
        'offset': offset,
        'limit': limit,
    })


def _read_binary_buffer(buffer_obj):
    if not buffer_obj:
        return b''
    try:
        if hasattr(buffer_obj, 'seek'):
            buffer_obj.seek(0)
        data = buffer_obj.read()
        if isinstance(data, bytes):
            return data
        if isinstance(data, bytearray):
            return bytes(data)
    except Exception:
        pass
    try:
        data = buffer_obj.getvalue()
        if isinstance(data, bytes):
            return data
        if isinstance(data, bytearray):
            return bytes(data)
    except Exception:
        pass
    return b''


@require_http_methods(["POST"])
@login_required
@api_require_permission('perm_print_list')
@rate_limit(max_requests=5, window_seconds=60, key_prefix='print_generate')
def api_generate_pdf(request, table_id):
    """Generate output PDF for selected generate_list cards.

    Uses saved field mappings and design PDFs.
    On success: returns PDF and moves cards to finalized.
    Body: { "request_ids": [1, 2, 3] }
    """
    table, err = _check_print_table_scope(request.user, table_id)
    if err:
        return err

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    preview_raw = data.get('preview_only', False)
    if isinstance(preview_raw, str):
        preview_only = preview_raw.strip().lower() in {'1', 'true', 'yes', 'y'}
    else:
        preview_only = bool(preview_raw)

    export_format = str(data.get('export_format') or 'pdf').strip().lower()
    if export_format not in {'pdf', 'png', 'zip'}:
        export_format = 'pdf'

    queue_ready = bool(data.get('queue_ready', False))

    batch_size = data.get('batch_size')
    try:
        batch_size = int(batch_size) if batch_size is not None else None
    except (TypeError, ValueError):
        batch_size = None
    if batch_size is not None:
        batch_size = max(1, min(500, batch_size))

    request_ids = data.get('request_ids', [])
    if not request_ids:
        return JsonResponse({'status': 'error', 'message': 'No cards selected'}, status=400)

    template_id = data.get('template_id')
    tmpl = _get_template_for_table(table, template_id=template_id, create_default=False)
    if not tmpl:
        return JsonResponse({'status': 'error', 'message': 'No template configured for this table'}, status=400)

    template_json = tmpl.template_json if isinstance(tmpl.template_json, dict) else {}
    elements = template_json.get('elements') if isinstance(template_json.get('elements'), list) else []

    has_front_elements = False
    has_back_elements = False
    for item in elements:
        if not isinstance(item, dict):
            continue
        side = str(item.get('side') or 'front').strip().lower()
        if side not in ('front', 'back', 'both'):
            side = 'front'
        if side == 'both':
            has_front_elements = True
            has_back_elements = True
            continue
        if side == 'back':
            has_back_elements = True
        else:
            has_front_elements = True

    if not has_front_elements:
        return JsonResponse({'status': 'error', 'message': 'Add at least one front element before generating'}, status=400)
    if tmpl.is_two_sided and not has_back_elements:
        return JsonResponse({'status': 'error', 'message': 'Add at least one back element before generating'}, status=400)

    prs = list(
        PrintRequest.objects.filter(
            id__in=request_ids, table=table, status='generate_list',
        ).select_related('card').order_by('created_at')
    )
    if not prs:
        return JsonResponse({'status': 'error', 'message': 'No valid generate-list cards found'}, status=400)

    if preview_only:
        prs = prs[:1]

    layout_options = data.get('layout') if isinstance(data.get('layout'), dict) else None
    if queue_ready:
        payload = GenerateCardService.build_job_payload(
            table,
            tmpl,
            prs,
            requested_by=request.user,
            layout_options=layout_options,
            export_format=export_format,
            batch_size=batch_size,
        )
        return JsonResponse({'status': 'ok', 'queue_ready': True, 'job': payload})

    pdf_bytes = None
    if export_format == 'zip':
        file_bytes, error = GenerateCardService.build_pdf_zip_for_cards(
            table,
            tmpl,
            prs,
            layout_options=layout_options,
            batch_size=batch_size,
        )
    else:
        file_buffer, error = GenerateCardService.generate(
            table,
            tmpl,
            prs,
            layout_options=layout_options,
            batch_size=batch_size,
        )
        pdf_bytes = _read_binary_buffer(file_buffer)
        file_bytes = pdf_bytes

    if error or not file_bytes:
        logger.error('api_generate_pdf: %s', error)
        return JsonResponse({'status': 'error', 'message': f'Generation failed: {error or "Unknown error"}'}, status=500)

    if export_format == 'png':
        png_pages, png_err = GenerateCardService.render_pdf_pages_to_png(pdf_bytes or file_bytes)
        if png_err or not png_pages:
            logger.error('api_generate_pdf (png): %s', png_err)
            return JsonResponse({'status': 'error', 'message': f'PNG export failed: {png_err or "Unknown error"}'}, status=500)
        if preview_only and len(png_pages) == 1:
            file_bytes = png_pages[0]
        else:
            file_bytes = GenerateCardService.build_png_zip_bytes(
                png_pages,
                file_prefix='preview_card' if preview_only else 'cards',
            )

    if not preview_only:
        # Move cards to finalized
        valid_ids = [pr.id for pr in prs]
        PrintWorkflowService.bulk_generate(valid_ids, request.user)

    usage_increment = max(1, len(prs))
    CardTemplate.objects.filter(id=tmpl.id).update(
        usage_count=F('usage_count') + usage_increment,
        last_used_at=timezone.now(),
    )

    now_stamp = timezone.now().strftime('%Y%m%d_%H%M%S')
    file_prefix = 'preview_card' if preview_only else 'cards'
    if export_format == 'zip':
        safe_filename = f'{file_prefix}_{table.id}_{now_stamp}.zip'
        response = HttpResponse(file_bytes, content_type='application/zip')
    elif export_format == 'png':
        if preview_only and (pdf_bytes is not None and file_bytes and file_bytes[:8] == b'\x89PNG\r\n\x1a\n'):
            safe_filename = f'{file_prefix}_{table.id}_{now_stamp}.png'
            response = HttpResponse(file_bytes, content_type='image/png')
        else:
            safe_filename = f'{file_prefix}_{table.id}_{now_stamp}.zip'
            response = HttpResponse(file_bytes, content_type='application/zip')
    else:
        safe_filename = f'{file_prefix}_{table.id}_{now_stamp}.pdf'
        response = HttpResponse(file_bytes, content_type='application/pdf')

    disposition = 'inline' if preview_only else 'attachment'
    response['Content-Disposition'] = f'{disposition}; filename="{safe_filename}"'
    return response
