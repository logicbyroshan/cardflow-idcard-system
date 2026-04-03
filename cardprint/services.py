"""
Card Print Workflow Service
===========================
All mutations for PrintRequest go through this service.
Views must NOT call .save(), .create(), .delete() directly.
Modelled after ReprintWorkflowService.

Workflow: generate_list (queued) → finalized (via Generate Card PDF) → pool
"""
import io
import logging
import os
import re
import base64
from difflib import SequenceMatcher

from django.db import transaction
from django.utils import timezone

from core.services.base import ServiceResult
from .models import PrintRequest

logger = logging.getLogger(__name__)


class PrintWorkflowService:
    """Service layer for the card-print workflow."""

    ALLOWED_TRANSITIONS = {
        'generate_list': ['finalized'],
        'finalized': ['pool'],
    }

    VALID_STATUSES = ['generate_list', 'finalized', 'pool']

    @classmethod
    def create_requests(cls, table, card_ids, user):
        """Create PrintRequest rows for the given card IDs.

        Skips cards that already have an active (non-pool) print request
        for the same table to avoid duplicates.

        Returns ServiceResult with data: {created: int, skipped: int}
        """
        if not card_ids:
            return ServiceResult(success=False, message='No card IDs provided')

        with transaction.atomic():
            existing_ids = set(
                PrintRequest.objects.select_for_update().filter(
                    table=table,
                    card_id__in=card_ids,
                    status__in=['generate_list', 'finalized'],
                ).values_list('card_id', flat=True)
            )

            to_create = []
            skipped = 0
            for cid in card_ids:
                if cid in existing_ids:
                    skipped += 1
                    continue
                to_create.append(PrintRequest(
                    card_id=cid,
                    table=table,
                    status='generate_list',
                    requested_by=user,
                ))

            if to_create:
                PrintRequest.objects.bulk_create(to_create, ignore_conflicts=True)

        created = len(to_create)
        logger.info(
            'PrintWorkflow: created=%d skipped=%d table=%d user=%s',
            created, skipped, table.id, user.username,
        )
        return ServiceResult(
            success=True,
            message=f'{created} card(s) added to generate list',
            data={'created': created, 'skipped': skipped},
        )

    @classmethod
    def bulk_generate(cls, request_ids, user):
        """Transition generate_list → finalized for a batch of PrintRequest IDs.
        Called after PDF has been generated and downloaded in Generate Card page.

        Returns ServiceResult with data: {updated: int, skipped: int}
        """
        target_status = 'finalized'
        valid_from = [s for s, targets in cls.ALLOWED_TRANSITIONS.items() if target_status in targets]

        with transaction.atomic():
            qs = PrintRequest.objects.select_for_update().filter(
                id__in=request_ids,
                status__in=valid_from,
            )
            updated = qs.update(status=target_status, updated_at=timezone.now())

        skipped = len(request_ids) - updated
        logger.info(
            'PrintWorkflow: finalize updated=%d skipped=%d user=%s',
            updated, skipped, user.username,
        )
        if not updated:
            return ServiceResult(
                success=False,
                message='No generate list items eligible for finalization.',
                data={'updated': 0, 'skipped': skipped},
            )
        return ServiceResult(
            success=True,
            message=f'{updated} item(s) finalized successfully',
            data={'updated': updated, 'skipped': skipped},
        )

    @classmethod
    def bulk_mark_pool(cls, request_ids, user):
        """Transition finalized → pool for a batch of PrintRequest IDs.

        Returns ServiceResult with data: {updated: int, skipped: int}
        """
        target_status = 'pool'
        valid_from = [s for s, targets in cls.ALLOWED_TRANSITIONS.items() if target_status in targets]

        with transaction.atomic():
            qs = PrintRequest.objects.select_for_update().filter(
                id__in=request_ids,
                status__in=valid_from,
            )
            updated = qs.update(status=target_status, updated_at=timezone.now())

        skipped = len(request_ids) - updated
        logger.info(
            'PrintWorkflow: mark_pool updated=%d skipped=%d user=%s',
            updated, skipped, user.username,
        )
        if not updated:
            return ServiceResult(
                success=False,
                message='No finalized items eligible for pool.',
                data={'updated': 0, 'skipped': skipped},
            )
        return ServiceResult(
            success=True,
            message=f'{updated} item(s) moved to pool',
            data={'updated': updated, 'skipped': skipped},
        )

# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION SERVICE
# ─────────────────────────────────────────────────────────────────────────────

IMAGE_FIELD_TYPES = {'photo', 'image', 'mother_photo', 'father_photo', 'signature', 'barcode', 'qr_code'}


class GenerateCardService:
    """
    Generates ID-card PDFs using ReportLab data overlay merged onto design PDFs.
    Card size: 87mm × 57mm (CR80 ID card).
    One page per card side.
    """
    CARD_LANDSCAPE_W_MM = 87.0
    CARD_LANDSCAPE_H_MM = 57.0
    CARD_PORTRAIT_W_MM = 57.0
    CARD_PORTRAIT_H_MM = 87.0

    @classmethod
    def dimensions_for_orientation_mm(cls, orientation):
        if orientation == 'portrait':
            return cls.CARD_PORTRAIT_W_MM, cls.CARD_PORTRAIT_H_MM
        return cls.CARD_LANDSCAPE_W_MM, cls.CARD_LANDSCAPE_H_MM

    @classmethod
    def _resolve_dimensions_mm(cls, template):
        cfg = template.field_config if isinstance(template.field_config, dict) else {}
        orientation = cfg.get('card_orientation') or 'landscape'
        return cls.dimensions_for_orientation_mm(orientation)

    @classmethod
    def generate(cls, table, template, print_requests):
        """
        Generate a multi-page PDF: one page per card (two pages per card for 2-sided).
        Uses ReportLab for data overlay and merges each page onto uploaded design PDF.

        Args:
            table: IDCardTable instance
            template: CardTemplate instance
            print_requests: list of PrintRequest instances (with card pre-fetched)

        Returns:
            (io.BytesIO buffer, None) on success
            (None, error_message) on failure
        """
        try:
            from reportlab.lib.units import mm
            from reportlab.pdfgen import canvas as rl_canvas

            buffer = io.BytesIO()
            card_w_mm, card_h_mm = cls._resolve_dimensions_mm(template)
            card_w_pt = card_w_mm * mm
            card_h_pt = card_h_mm * mm

            c = rl_canvas.Canvas(buffer, pagesize=(card_w_pt, card_h_pt))

            style = cls._pdf_style_from_template(template)
            front_mappings = (template.field_mappings or {}).get('front', {})
            back_mappings = (template.field_mappings or {}).get('back', {})
            field_cfg = template.field_config or {}
            front_allowed = set(field_cfg.get('front_fields') or [])
            back_allowed = set(field_cfg.get('back_fields') or [])
            front_editable = cls._extract_editable_model(field_cfg.get('editable_design_front'))
            back_editable = cls._extract_editable_model(field_cfg.get('editable_design_back')) if template.is_two_sided else None

            # Build field type map from table definition
            field_type_map = {f['name']: f.get('type', 'text') for f in (table.fields or [])}

            for pr in print_requests:
                card = pr.card
                fd = card.field_data or {}
                fd_upper = {k.upper(): v for k, v in fd.items()}

                # Front
                if front_editable:
                    cls._draw_side_from_editable_model(
                        c,
                        front_editable,
                        fd,
                        fd_upper,
                        field_type_map,
                        front_mappings,
                        card_h_pt,
                        style,
                        front_allowed,
                    )
                else:
                    cls._draw_side(
                        c, fd, fd_upper, field_type_map, front_mappings,
                        card_h_pt, style, front_allowed,
                    )
                c.showPage()

                # Back (always render a back page for 2-sided templates)
                if template.is_two_sided:
                    if back_editable:
                        cls._draw_side_from_editable_model(
                            c,
                            back_editable,
                            fd,
                            fd_upper,
                            field_type_map,
                            back_mappings,
                            card_h_pt,
                            style,
                            back_allowed,
                        )
                    else:
                        cls._draw_side(
                            c, fd, fd_upper, field_type_map, back_mappings,
                            card_h_pt, style, back_allowed,
                        )
                    c.showPage()

            c.save()
            buffer.seek(0)

            merged_buffer, merge_err = cls._merge_overlay_with_design(
                template,
                buffer,
                use_front_design=not bool(front_editable),
                use_back_design=not bool(back_editable),
            )
            if merge_err:
                return None, merge_err
            return merged_buffer, None

        except Exception as exc:
            logger.error('GenerateCardService.generate error: %s', exc, exc_info=True)
            return None, str(exc)

    @classmethod
    def _pdf_style_from_template(cls, template):
        cfg = template.field_config if isinstance(template.field_config, dict) else {}
        raw = cfg.get('docx_text_style') if isinstance(cfg.get('docx_text_style'), dict) else {}

        family_raw = str(raw.get('font_family') or template.font_family or 'Arial').strip().lower()
        weight = str(raw.get('font_weight') or 'normal').strip().lower()
        try:
            size = float(raw.get('font_size_pt') or template.font_size or 11)
        except (TypeError, ValueError):
            size = float(template.font_size or 11)
        size = max(6.0, min(72.0, size))

        try:
            line_height = float(raw.get('line_height') or 1.15)
        except (TypeError, ValueError):
            line_height = 1.15
        line_height = max(0.8, min(3.0, line_height))

        align = str(raw.get('align') or 'left').strip().lower()
        if align not in ('left', 'center', 'right'):
            align = 'left'

        if 'times' in family_raw:
            font_name = 'Times-Bold' if weight in ('bold', 'semibold') else 'Times-Roman'
        elif 'courier' in family_raw:
            font_name = 'Courier-Bold' if weight in ('bold', 'semibold') else 'Courier'
        else:
            font_name = 'Helvetica-Bold' if weight in ('bold', 'semibold') else 'Helvetica'

        color_hex = cls._normalize_hex_color(raw.get('font_color_hex'), default='111111')
        r = int(color_hex[0:2], 16) / 255.0
        g = int(color_hex[2:4], 16) / 255.0
        b = int(color_hex[4:6], 16) / 255.0

        return {
            'font_name': font_name,
            'font_size': size,
            'line_height': line_height,
            'align': align,
            'color_rgb': (r, g, b),
        }

    @classmethod
    def _template_pdf_bytes(cls, file_field):
        if not file_field:
            return None
        try:
            with file_field.open('rb') as fobj:
                return fobj.read()
        except Exception as exc:
            logger.warning('Template PDF read failed: %s', exc)
            return None

    @classmethod
    def _extract_editable_model(cls, raw):
        if not isinstance(raw, dict):
            return None
        lines = raw.get('lines') if isinstance(raw.get('lines'), list) else []
        images = raw.get('images') if isinstance(raw.get('images'), list) else []
        if not lines and not images:
            return None
        return {
            'lines': [dict(item) for item in lines if isinstance(item, dict)],
            'images': [dict(item) for item in images if isinstance(item, dict)],
        }

    @classmethod
    def _merge_overlay_with_design(cls, template, overlay_buffer, use_front_design=True, use_back_design=True):
        try:
            from pypdf import PdfReader, PdfWriter
        except Exception:
            # pypdf unavailable: return overlay-only PDF.
            overlay_buffer.seek(0)
            return overlay_buffer, None

        try:
            overlay_buffer.seek(0)
            overlay_reader = PdfReader(overlay_buffer)
            writer = PdfWriter()

            front_design = cls._template_pdf_bytes(getattr(template, 'front_pdf', None)) if use_front_design else None
            back_design = cls._template_pdf_bytes(getattr(template, 'back_pdf', None)) if use_back_design else None

            page_idx = 0
            for _ in range(len(overlay_reader.pages)):
                side = 'front'
                if template.is_two_sided and (page_idx % 2 == 1):
                    side = 'back'

                overlay_page = overlay_reader.pages[page_idx]
                page_idx += 1

                design_bytes = back_design if side == 'back' else front_design
                if design_bytes:
                    design_reader = PdfReader(io.BytesIO(design_bytes))
                    base_page = design_reader.pages[0]
                    base_page.merge_page(overlay_page)
                    writer.add_page(base_page)
                else:
                    writer.add_page(overlay_page)

            out = io.BytesIO()
            writer.write(out)
            out.seek(0)
            return out, None
        except Exception as exc:
            logger.error('PDF merge failed: %s', exc, exc_info=True)
            return None, 'Failed to merge data with design PDF'

    @classmethod
    def _draw_side(
        cls,
        c,
        fd,
        fd_upper,
        field_type_map,
        mappings,
        card_h_pt,
        style,
        allowed_fields=None,
    ):
        """Draw all mapped fields onto the current ReportLab canvas page."""
        from reportlab.lib.units import mm
        from django.conf import settings

        for field_name, mapping in mappings.items():
            if allowed_fields and field_name not in allowed_fields:
                continue
            value = fd.get(field_name) or fd_upper.get(field_name.upper()) or ''
            ftype = field_type_map.get(field_name, 'text')

            x_mm = float(mapping.get('x_mm', 0))
            y_mm = float(mapping.get('y_mm', 0))
            w_mm = float(mapping.get('w_mm', 20))
            h_mm = float(mapping.get('h_mm', 10))

            # ReportLab: (0,0) = bottom-left. Our origin: top-left.
            rl_x = x_mm * mm
            rl_y = card_h_pt - (y_mm + h_mm) * mm  # bottom-left of box in RL coords
            rl_w = w_mm * mm
            rl_h = h_mm * mm

            if cls._is_image_field(ftype, field_name):
                cls._draw_image(c, value, rl_x, rl_y, rl_w, rl_h, settings)
            else:
                render_text = cls._build_render_text_for_mapping(field_name, mapping, value)
                cls._draw_text(c, render_text, rl_x, rl_y, rl_w, rl_h, style)

    @classmethod
    def _build_render_text_for_mapping(cls, field_name, mapping, value):
        label, text_value, show_key = cls._mapping_text_parts(field_name, mapping, value)

        if show_key and label:
            if not label.endswith(':'):
                label = f'{label}:'
            return f'{label}\n{text_value}'
        return text_value

    @classmethod
    def _mapping_text_parts(cls, field_name, mapping, value):
        label = str(mapping.get('label_text') or cls._format_field_label(field_name)).strip()
        placeholder_raw = mapping.get('placeholder', None)
        if placeholder_raw is None:
            placeholder = 'XXXXX'
        else:
            placeholder = str(placeholder_raw).strip()
        show_key_raw = mapping.get('show_key', True)
        if isinstance(show_key_raw, str):
            show_key = show_key_raw.strip().lower() not in ('false', '0', 'no')
        else:
            show_key = bool(show_key_raw)

        text_value = str(value).strip() if value is not None else ''
        if not text_value:
            text_value = placeholder
        return label, text_value, show_key

    @classmethod
    def _extract_key_prefix(cls, text):
        src = str(text or '').strip()
        if not src:
            return ''
        # Entire placeholder/date-mask style text is not a key prefix.
        if re.fullmatch(r'[xX0-9\-\./_\s]{3,}', src):
            return ''
        m = re.match(r'^(.*?(?:[:=-]|\s{2,})\s*)', src)
        if not m:
            # Fallback for common key-placeholder style like "Mobile XXXXX".
            parts = src.split()
            if len(parts) >= 2:
                tail = parts[-1]
                if re.fullmatch(r'[xX*._-]{3,}', tail) or re.fullmatch(r'[xX0-9+\-]{4,}', tail):
                    return ' '.join(parts[:-1]).strip() + ' '
            return ''
        prefix = str(m.group(1) or '').strip()
        if not prefix:
            return ''
        return prefix + (' ' if not prefix.endswith(' ') else '')

    @classmethod
    def _format_field_label(cls, field_name):
        raw = str(field_name or '').strip()
        if not raw:
            return 'Field'
        txt = re.sub(r'[_\-]+', ' ', raw)
        txt = re.sub(r'\s+', ' ', txt).strip()
        return txt.title()

    @classmethod
    def _normalize_hex_color(cls, value, default='111111'):
        raw = str(value or '').strip()
        if not raw:
            return default

        v = raw[1:] if raw.startswith('#') else raw
        if re.fullmatch(r'[0-9a-fA-F]{3}', v):
            v = ''.join(ch * 2 for ch in v)
        if not re.fullmatch(r'[0-9a-fA-F]{6}', v):
            return default
        return v.upper()

    @classmethod
    def _is_image_field(cls, ftype, field_name=''):
        """Detect image-like fields from type/name using tolerant matching."""
        t = (ftype or '').strip().lower()
        n = (field_name or '').strip().lower()
        if t in IMAGE_FIELD_TYPES:
            return True
        if t in {'file', 'img', 'picture'}:
            return True
        if any(key in n for key in ('photo', 'image', 'signature', 'sign', 'barcode', 'qr')):
            return True
        return False

    @classmethod
    def _draw_image(cls, c, value, x, y, w, h, settings):
        """Draw an image scaled to fit within the bounding box (preserveAspectRatio)."""
        if not value or value == 'NOT_FOUND' or str(value).startswith('PENDING:'):
            return
        try:
            from reportlab.lib.utils import ImageReader

            raw_value = str(value).replace('\\', '/')
            if raw_value.startswith('/media/'):
                raw_value = raw_value[len('/media/'):]
            elif raw_value.startswith('media/'):
                raw_value = raw_value[len('media/'):]

            img_path = os.path.join(settings.MEDIA_ROOT, raw_value)
            if not os.path.exists(img_path):
                return
            img_reader = ImageReader(img_path)
            iw, ih = img_reader.getSize()
            if iw <= 0 or ih <= 0:
                return
            # Scale to fit while preserving aspect ratio
            scale = min(w / iw, h / ih)
            draw_w = iw * scale
            draw_h = ih * scale
            # Centre inside the box
            offset_x = (w - draw_w) / 2
            offset_y = (h - draw_h) / 2
            c.drawImage(img_reader, x + offset_x, y + offset_y,
                        width=draw_w, height=draw_h, mask='auto')
        except Exception as exc:
            logger.warning('GenerateCardService: image draw failed for %s: %s', value, exc)

    @classmethod
    def _draw_text(cls, c, text, x, y, w, h, style):
        """Draw text inside a bounding box with basic word-wrap."""
        if not text:
            return
        try:
            font_name = style.get('font_name', 'Helvetica') if isinstance(style, dict) else 'Helvetica'
            font_size = float(style.get('font_size', 11)) if isinstance(style, dict) else 11.0
            line_height_mult = float(style.get('line_height', 1.3)) if isinstance(style, dict) else 1.3
            align = style.get('align', 'left') if isinstance(style, dict) else 'left'
            color_rgb = style.get('color_rgb', (0, 0, 0)) if isinstance(style, dict) else (0, 0, 0)

            c.saveState()
            # Clip to box
            path = c.beginPath()
            path.rect(x, y, w, h)
            c.clipPath(path, stroke=0, fill=0)

            c.setFont(font_name, font_size)
            c.setFillColorRGB(color_rgb[0], color_rgb[1], color_rgb[2])

            line_height = font_size * line_height_mult
            max_text_width = max(1.0, w - 2.0)
            key_prefix_text = str(style.get('key_prefix_text') or '') if isinstance(style, dict) else ''
            value_text = style.get('value_text') if isinstance(style, dict) else None

            if key_prefix_text and value_text is not None:
                prefix_w = c.stringWidth(key_prefix_text, font_name, font_size)
                first_line_width = max(1.0, max_text_width - prefix_w)
                value_lines = cls._wrap_text_to_width(c, font_name, font_size, str(value_text), max_text_width, first_line_width)
                if value_lines:
                    lines = [key_prefix_text + value_lines[0]] + value_lines[1:]
                else:
                    lines = [key_prefix_text.strip()]
                align = 'left'
            else:
                lines = cls._wrap_text_to_width(c, font_name, font_size, text, max_text_width)

            # Draw lines from top of box downward
            # In RL, y is from bottom. Top of box = y + h. First baseline just inside top.
            baseline_y = y + h - font_size
            for line in lines:
                if baseline_y < y:
                    break
                if align == 'right':
                    text_width = c.stringWidth(line, font_name, font_size)
                    draw_x = max(x + 1, x + w - text_width - 1)
                elif align == 'center':
                    text_width = c.stringWidth(line, font_name, font_size)
                    draw_x = x + max(1, (w - text_width) / 2)
                else:
                    draw_x = x + 1
                c.drawString(draw_x, baseline_y, line)
                baseline_y -= line_height

            c.restoreState()
        except Exception as exc:
            logger.warning('GenerateCardService: text draw failed: %s', exc)

    @classmethod
    def _wrap_text_to_width(cls, c, font_name, font_size, text, max_width, first_line_width=None):
        lines = []
        first_limit = max(1.0, float(first_line_width if first_line_width is not None else max_width))
        normal_limit = max(1.0, float(max_width))
        line_limit = first_limit

        for raw_line in str(text or '').split('\n'):
            src = str(raw_line or '').strip()
            if not src:
                lines.append('')
                line_limit = normal_limit
                continue

            words = src.split()
            current = ''
            for word in words:
                candidate = (current + ' ' + word).strip() if current else word
                if c.stringWidth(candidate, font_name, font_size) <= line_limit:
                    current = candidate
                    continue

                if current:
                    lines.append(current)
                    current = ''
                    line_limit = normal_limit

                if c.stringWidth(word, font_name, font_size) <= line_limit:
                    current = word
                    continue

                frag = ''
                for ch in word:
                    cand = frag + ch
                    if c.stringWidth(cand, font_name, font_size) <= line_limit:
                        frag = cand
                    else:
                        if frag:
                            lines.append(frag)
                            line_limit = normal_limit
                        frag = ch
                current = frag

            if current:
                lines.append(current)
            line_limit = normal_limit

        return lines

    @classmethod
    def _reportlab_font_name(cls, family_raw, weight_raw):
        fam = str(family_raw or '').strip().lower()
        weight = str(weight_raw or '').strip().lower()
        is_bold = weight in {'bold', '700', '600', 'semibold'}
        if 'times' in fam:
            return 'Times-Bold' if is_bold else 'Times-Roman'
        if 'courier' in fam:
            return 'Courier-Bold' if is_bold else 'Courier'
        return 'Helvetica-Bold' if is_bold else 'Helvetica'

    @classmethod
    def _style_from_editable_line(cls, line, default_style):
        style = default_style if isinstance(default_style, dict) else {}
        font_size = float(line.get('font_size_pt') or style.get('font_size') or 11)
        font_size = max(6.0, min(72.0, font_size))

        line_height = float(line.get('line_height') or style.get('line_height') or 1.15)
        line_height = max(0.8, min(3.0, line_height))

        align = str(line.get('text_align') or style.get('align') or 'left').strip().lower()
        if align not in {'left', 'center', 'right'}:
            align = 'left'

        weight_raw = line.get('font_weight') or '400'
        family_raw = line.get('font_family') or style.get('font_name') or 'Helvetica'
        font_name = cls._reportlab_font_name(family_raw, weight_raw)

        color_hex = cls._normalize_hex_color(line.get('font_color_hex'), default='111111')
        color_rgb = (
            int(color_hex[0:2], 16) / 255.0,
            int(color_hex[2:4], 16) / 255.0,
            int(color_hex[4:6], 16) / 255.0,
        )

        return {
            'font_name': font_name,
            'font_size': font_size,
            'line_height': line_height,
            'align': align,
            'color_rgb': color_rgb,
        }

    @classmethod
    def _rect_overlap_ratio(cls, ax, ay, aw, ah, bx, by, bw, bh):
        ax2, ay2 = ax + aw, ay + ah
        bx2, by2 = bx + bw, by + bh
        inter_w = max(0.0, min(ax2, bx2) - max(ax, bx))
        inter_h = max(0.0, min(ay2, by2) - max(ay, by))
        inter_area = inter_w * inter_h
        area_a = max(0.0001, aw * ah)
        return inter_area / area_a

    @classmethod
    def _line_center_in_rect(cls, line, rx, ry, rw, rh):
        x = float(line.get('x_mm') or 0.0)
        y = float(line.get('y_mm') or 0.0)
        w = max(0.0, float(line.get('w_mm') or 0.0))
        h = max(0.0, float(line.get('h_mm') or 0.0))
        cx = x + (w / 2.0)
        cy = y + (h / 2.0)
        return (rx <= cx <= (rx + rw)) and (ry <= cy <= (ry + rh))

    @classmethod
    def _line_match_score(cls, line, x_mm, y_mm, w_mm, h_mm):
        lx = float(line.get('x_mm') or 0.0)
        ly = float(line.get('y_mm') or 0.0)
        lw = max(0.0, float(line.get('w_mm') or 0.0))
        lh = max(0.0, float(line.get('h_mm') or 0.0))
        if lw <= 0.0 or lh <= 0.0:
            return None

        overlap = cls._rect_overlap_ratio(lx, ly, lw, lh, x_mm, y_mm, w_mm, h_mm)
        center_inside = cls._line_center_in_rect(line, x_mm, y_mm, w_mm, h_mm)
        if overlap < 0.10 and not center_inside:
            return None

        line_cx = lx + (lw / 2.0)
        line_cy = ly + (lh / 2.0)
        dx_anchor = abs(lx - x_mm) / max(1.0, w_mm)
        dy_anchor = abs(ly - y_mm) / max(1.0, h_mm)
        dx_center = abs(line_cx - (x_mm + (w_mm / 2.0))) / max(1.0, w_mm)
        dy_center = abs(line_cy - (y_mm + (h_mm / 2.0))) / max(1.0, h_mm)

        return (
            (overlap * 1.65)
            + (0.26 if center_inside else 0.0)
            - (0.80 * dx_anchor)
            - (1.35 * dy_anchor)
            - (0.10 * dx_center)
            - (0.08 * dy_center)
        )

    @classmethod
    def _find_best_line_index(cls, lines, x_mm, y_mm, w_mm, h_mm, exclude_indexes=None):
        excludes = set(exclude_indexes or [])
        ranked = []
        for idx, line in enumerate(lines):
            if idx in excludes:
                continue
            score = cls._line_match_score(line, x_mm, y_mm, w_mm, h_mm)
            if score is None:
                continue
            ly = float(line.get('y_mm') or 0.0)
            lx = float(line.get('x_mm') or 0.0)
            ranked.append((score, idx, ly, lx))

        if not ranked:
            return None
        ranked.sort(key=lambda item: (-item[0], item[2], item[3]))
        return ranked[0][1]

    @classmethod
    def _is_source_idx_reliable(cls, lines, source_idx, x_mm, y_mm, w_mm, h_mm):
        if source_idx is None:
            return False
        if source_idx < 0 or source_idx >= len(lines):
            return False
        score = cls._line_match_score(lines[source_idx], x_mm, y_mm, w_mm, h_mm)
        return score is not None and score >= -0.10

    @classmethod
    def _apply_text_mapping_to_lines(cls, lines, x_mm, y_mm, w_mm, h_mm, render_text):
        if not lines:
            return False

        idx = cls._find_best_line_index(lines, x_mm, y_mm, w_mm, h_mm)
        if idx is None:
            return False

        lines[idx]['text'] = render_text

        return True

    @classmethod
    def _draw_inline_data_image(cls, c, data_url, x, y, w, h):
        if not data_url or not str(data_url).startswith('data:image/'):
            return
        try:
            from reportlab.lib.utils import ImageReader

            payload = str(data_url)
            if ',' not in payload:
                return
            encoded = payload.split(',', 1)[1]
            raw = base64.b64decode(encoded)
            img_reader = ImageReader(io.BytesIO(raw))
            iw, ih = img_reader.getSize()
            if iw <= 0 or ih <= 0:
                return

            scale = min(w / iw, h / ih)
            draw_w = iw * scale
            draw_h = ih * scale
            offset_x = (w - draw_w) / 2
            offset_y = (h - draw_h) / 2
            c.drawImage(img_reader, x + offset_x, y + offset_y, width=draw_w, height=draw_h, mask='auto')
        except Exception as exc:
            logger.warning('GenerateCardService: inline image draw failed: %s', exc)

    @classmethod
    def _draw_side_from_editable_model(
        cls,
        c,
        editable_model,
        fd,
        fd_upper,
        field_type_map,
        mappings,
        card_h_pt,
        default_style,
        allowed_fields=None,
    ):
        from reportlab.lib.units import mm
        from django.conf import settings

        lines = [dict(item) for item in (editable_model.get('lines') or []) if isinstance(item, dict)]
        images = [dict(item) for item in (editable_model.get('images') or []) if isinstance(item, dict)]
        used_line_indexes = set()
        mapped_line_indexes = set()

        # Draw template embedded images first. Mapped dynamic photo fields should
        # render after this pass so they replace the template image in preview/output.
        for image in images:
            x_mm = float(image.get('x_mm') or 0.0)
            y_mm = float(image.get('y_mm') or 0.0)
            w_mm = max(0.5, float(image.get('w_mm') or 0.0))
            h_mm = max(0.5, float(image.get('h_mm') or 0.0))
            data_url = str(image.get('data_url') or '')
            if not data_url:
                continue

            rl_x = x_mm * mm
            rl_y = card_h_pt - (y_mm + h_mm) * mm
            rl_w = w_mm * mm
            rl_h = h_mm * mm
            cls._draw_inline_data_image(c, data_url, rl_x, rl_y, rl_w, rl_h)

        for field_name, mapping in mappings.items():
            if allowed_fields and field_name not in allowed_fields:
                continue

            ftype = field_type_map.get(field_name, 'text')
            value = fd.get(field_name) or fd_upper.get(field_name.upper()) or ''
            x_mm = float(mapping.get('x_mm', 0))
            y_mm = float(mapping.get('y_mm', 0))
            w_mm = max(0.5, float(mapping.get('w_mm', 20)))
            h_mm = max(0.5, float(mapping.get('h_mm', 10)))
            label_text, value_text, show_key = cls._mapping_text_parts(field_name, mapping, value)

            if cls._is_image_field(ftype, field_name):
                rl_x = x_mm * mm
                rl_y = card_h_pt - (y_mm + h_mm) * mm
                rl_w = w_mm * mm
                rl_h = h_mm * mm
                cls._draw_image(c, value, rl_x, rl_y, rl_w, rl_h, settings)
                continue

            render_text = cls._build_render_text_for_mapping(field_name, mapping, value)
            source_idx_raw = mapping.get('source_line_idx', None)
            source_idx = None
            try:
                source_idx = int(source_idx_raw)
            except (TypeError, ValueError):
                source_idx = None

            matched = False
            chosen_idx = None
            if source_idx is not None and source_idx not in used_line_indexes and cls._is_source_idx_reliable(lines, source_idx, x_mm, y_mm, w_mm, h_mm):
                chosen_idx = source_idx

            if chosen_idx is None:
                chosen_idx = cls._find_best_line_index(lines, x_mm, y_mm, w_mm, h_mm, exclude_indexes=used_line_indexes)

            if chosen_idx is not None:
                lines[chosen_idx]['w_mm'] = max(float(lines[chosen_idx].get('w_mm') or 0.0), float(w_mm))
                lines[chosen_idx]['h_mm'] = max(float(lines[chosen_idx].get('h_mm') or 0.0), float(h_mm))
                original_text = str(lines[chosen_idx].get('text') or '')
                key_prefix = '' if show_key else cls._extract_key_prefix(original_text)
                if key_prefix:
                    lines[chosen_idx]['text'] = key_prefix + value_text
                    lines[chosen_idx]['__wrap_value_only'] = True
                    lines[chosen_idx]['__key_prefix_text'] = key_prefix
                    lines[chosen_idx]['__value_text'] = value_text
                else:
                    lines[chosen_idx]['text'] = render_text
                    lines[chosen_idx].pop('__wrap_value_only', None)
                    lines[chosen_idx].pop('__key_prefix_text', None)
                    lines[chosen_idx].pop('__value_text', None)

                field_name_l = str(field_name or '').strip().lower()
                if any(k in field_name_l for k in ('mobile', 'contact', 'phone')):
                    lines[chosen_idx]['text_align'] = 'left'
                used_line_indexes.add(chosen_idx)
                mapped_line_indexes.add(chosen_idx)
                matched = True

            if not matched:
                matched = cls._apply_text_mapping_to_lines(lines, x_mm, y_mm, w_mm, h_mm, render_text)
            if not matched:
                lines.append({
                    'text': render_text,
                    'x_mm': x_mm,
                    'y_mm': y_mm,
                    'w_mm': w_mm,
                    'h_mm': h_mm,
                    'font_size_pt': float(default_style.get('font_size', 11)) if isinstance(default_style, dict) else 11.0,
                    'line_height': float(default_style.get('line_height', 1.15)) if isinstance(default_style, dict) else 1.15,
                    'text_align': str(default_style.get('align', 'left')) if isinstance(default_style, dict) else 'left',
                })
                used_line_indexes.add(len(lines) - 1)
                mapped_line_indexes.add(len(lines) - 1)

        # Expand text boxes for wrapped values so long fields (e.g. addresses)
        # can render on multiple lines instead of clipping to one line.
        mm_per_pt = 25.4 / 72.0
        for idx, line in enumerate(lines):
            if idx not in mapped_line_indexes:
                continue
            text = str(line.get('text') or '')
            if not text.strip():
                continue

            line_style = cls._style_from_editable_line(line, default_style)
            font_name = line_style.get('font_name', 'Helvetica')
            font_size = float(line_style.get('font_size', 11.0))
            line_height_mult = float(line_style.get('line_height', 1.15))
            box_w_pt = max(1.0, float(line.get('w_mm') or 1.0) * mm)
            max_text_width = max(1.0, box_w_pt - 2.0)

            key_prefix_text = str(line.get('__key_prefix_text') or '') if line.get('__wrap_value_only') else ''
            value_text_only = str(line.get('__value_text') or text) if line.get('__wrap_value_only') else text
            if key_prefix_text:
                prefix_w = c.stringWidth(key_prefix_text, font_name, font_size)
                first_width = max(1.0, max_text_width - prefix_w)
                wrapped_lines = cls._wrap_text_to_width(c, font_name, font_size, value_text_only, max_text_width, first_width)
                wrapped_count = max(1, len(wrapped_lines))
            else:
                wrapped_lines = cls._wrap_text_to_width(c, font_name, font_size, value_text_only, max_text_width)
                wrapped_count = max(1, len(wrapped_lines))

            needed_lines = max(1, wrapped_count)
            needed_h_mm = (needed_lines * font_size * line_height_mult * mm_per_pt) + 0.6
            current_h_mm = max(0.5, float(line.get('h_mm') or 0.5))
            line['h_mm'] = round(max(current_h_mm, needed_h_mm), 2)

        lines.sort(key=lambda item: (float(item.get('y_mm') or 0.0), float(item.get('x_mm') or 0.0)))
        for line in lines:
            text = str(line.get('text') or '')
            if not text.strip():
                continue

            x_mm = float(line.get('x_mm') or 0.0)
            y_mm = float(line.get('y_mm') or 0.0)
            w_mm = max(0.5, float(line.get('w_mm') or 1.0))
            h_mm = max(0.5, float(line.get('h_mm') or 1.0))

            rl_x = x_mm * mm
            rl_y = card_h_pt - (y_mm + h_mm) * mm
            rl_w = w_mm * mm
            rl_h = h_mm * mm

            line_style = cls._style_from_editable_line(line, default_style)
            if line.get('__wrap_value_only'):
                line_style = dict(line_style)
                line_style['key_prefix_text'] = str(line.get('__key_prefix_text') or '')
                line_style['value_text'] = str(line.get('__value_text') or '')
            cls._draw_text(c, text, rl_x, rl_y, rl_w, rl_h, line_style)


class PdfTemplateAnalyzerService:
    """Extract candidate field positions/styles from a design PDF using PyMuPDF."""

    @classmethod
    def analyze_template(cls, template, side, fields, card_w_mm, card_h_mm, exclude_fields=None):
        try:
            import fitz
        except Exception:
            return None, 'PyMuPDF is not available for PDF analysis'

        exclude = set(str(x).strip() for x in (exclude_fields or []) if str(x).strip())
        pdf_field = template.back_pdf if side == 'back' else template.front_pdf
        pdf_path = getattr(pdf_field, 'path', None) if pdf_field else None
        if not pdf_path or not os.path.exists(pdf_path):
            return {
                'mappings': {},
                'matched_style': {'font_size_pt': 11, 'font_family': ''},
                'engine': 'pymupdf',
            }, None

        try:
            doc = fitz.open(pdf_path)
            if doc.page_count < 1:
                return {
                    'mappings': {},
                    'matched_style': {'font_size_pt': 11, 'font_family': ''},
                    'engine': 'pymupdf',
                }, None

            page = doc.load_page(0)
            line_items = cls._extract_line_items(page)
            image_boxes = cls._extract_image_boxes(page)
            if not line_items:
                return {
                    'mappings': {},
                    'matched_style': {'font_size_pt': 11, 'font_family': ''},
                    'engine': 'pymupdf',
                }, None

            target_fields = [f for f in (fields or []) if isinstance(f, dict) and str(f.get('name', '')).strip()]
            text_fields = [
                f for f in target_fields
                if not GenerateCardService._is_image_field(f.get('type', ''), f.get('name', ''))
                and str(f.get('name', '')).strip() not in exclude
            ]
            image_fields = [
                f for f in target_fields
                if GenerateCardService._is_image_field(f.get('type', ''), f.get('name', ''))
                and str(f.get('name', '')).strip() not in exclude
            ]

            mappings = {}
            used_fields = set()
            used_line_indexes = set()

            font_pts = []
            font_families = {}
            matched_font_pts = []
            matched_font_families = {}

            def add_style_sample(item, weight=1):
                w = max(1, min(6, int(weight or 1)))
                size_pt = cls._clamp(float(item.get('font_size', 11.0)), 6.0, 72.0, 11.0)
                for _ in range(w):
                    font_pts.append(size_pt)

                mapped_family = cls._map_pdf_font_family(item.get('font_name', ''))
                if mapped_family:
                    font_families[mapped_family] = font_families.get(mapped_family, 0) + w

            def add_matched_style_sample(item, weight=1):
                w = max(1, min(10, int(weight or 1)))
                size_pt = cls._clamp(float(item.get('font_size', 11.0)), 6.0, 72.0, 11.0)
                for _ in range(w):
                    matched_font_pts.append(size_pt)

                mapped_family = cls._map_pdf_font_family(item.get('font_name', ''))
                if mapped_family:
                    matched_font_families[mapped_family] = matched_font_families.get(mapped_family, 0) + w

            for item in line_items:
                add_style_sample(item, 1)

            # Pass 1: key/value aware matching.
            for item in line_items:
                raw_text = item.get('text', '')
                parsed = cls._split_key_value_text(raw_text)

                best_field = None
                best_score = 0.0
                best_mode = 'text'

                for field in text_fields:
                    fname = str(field.get('name', '')).strip()
                    if not fname or fname in used_fields:
                        continue

                    text_score = cls._score_field_match(raw_text, fname)
                    key_score = cls._score_field_match(parsed.get('key_text', ''), fname) + 0.08 if parsed.get('has_delimiter') else 0.0
                    score = text_score
                    mode = 'text'
                    if key_score > score:
                        score = key_score
                        mode = 'key'

                    if parsed.get('has_delimiter') and mode == 'key' and parsed.get('value_text'):
                        score += 0.04

                    if score > best_score:
                        best_score = score
                        best_field = field
                        best_mode = mode

                if not best_field or best_score < 0.72:
                    continue

                fname = str(best_field.get('name', '')).strip()
                box = cls._estimate_text_box(item, page.rect.width, page.rect.height, card_w_mm, card_h_mm)

                show_key = True
                label_text = parsed.get('key_text', '').strip() or GenerateCardService._format_field_label(fname)
                placeholder = ''
                if parsed.get('has_delimiter') and parsed.get('value_text'):
                    # Inline sample value usually means value-only rendering on output.
                    show_key = False
                    label_text = ''
                    placeholder = ''
                elif (not parsed.get('has_delimiter')) and best_mode == 'text' and best_score < 0.83:
                    show_key = False
                    label_text = ''
                    placeholder = ''

                mappings[fname] = {
                    'x_mm': box['x_mm'],
                    'y_mm': box['y_mm'],
                    'w_mm': box['w_mm'],
                    'h_mm': box['h_mm'],
                    'label_text': label_text,
                    'placeholder': placeholder,
                    'show_key': show_key,
                }
                used_fields.add(fname)
                used_line_indexes.add(item['index'])
                add_style_sample(item, 3)
                add_matched_style_sample(item, 4)

            # Pass 2: value-only fallback for remaining text fields.
            remaining_fields = [f for f in text_fields if str(f.get('name', '')).strip() not in used_fields]
            key_match_count = len(used_fields)
            # Avoid noisy wrong placements when key labels are already detected.
            # Value-only fallback is mainly for templates that have little/no key labels.
            allow_value_only_fallback = key_match_count == 0
            if remaining_fields and allow_value_only_fallback:
                value_candidates = []
                for item in line_items:
                    if item['index'] in used_line_indexes:
                        continue
                    parsed = cls._split_key_value_text(item.get('text', ''))
                    if parsed.get('has_delimiter'):
                        continue

                    txt = str(item.get('text', '')).strip()
                    if len(txt) < 2:
                        continue

                    # Skip likely field-label lines (already key-like).
                    key_like = 0.0
                    for field in text_fields:
                        key_like = max(key_like, cls._score_field_match(txt, str(field.get('name', '')).strip()))
                    if key_like >= 0.78:
                        continue

                    value_candidates.append(item)

                value_candidates.sort(key=lambda it: (it.get('y0', 0.0), it.get('x0', 0.0)))
                for idx, field in enumerate(remaining_fields):
                    if idx >= len(value_candidates):
                        break
                    item = value_candidates[idx]
                    fname = str(field.get('name', '')).strip()
                    box = cls._estimate_text_box(item, page.rect.width, page.rect.height, card_w_mm, card_h_mm, prefer_value_only=True)
                    mappings[fname] = {
                        'x_mm': box['x_mm'],
                        'y_mm': box['y_mm'],
                        'w_mm': box['w_mm'],
                        'h_mm': box['h_mm'],
                        'label_text': '',
                        'placeholder': '',
                        'show_key': False,
                    }
                    used_fields.add(fname)
                    used_line_indexes.add(item['index'])
                    add_style_sample(item, 2)
                    add_matched_style_sample(item, 3)

            # Image fields: anchor next to matching image label.
            for field in image_fields:
                fname = str(field.get('name', '')).strip()
                if not fname or fname in mappings:
                    continue

                best_item = None
                best_score = 0.0
                for item in line_items:
                    score = cls._score_field_match(item.get('text', ''), fname)
                    if score > best_score:
                        best_score = score
                        best_item = item

                if not best_item or best_score < 0.68:
                    continue

                label_box = cls._estimate_text_box(best_item, page.rect.width, page.rect.height, card_w_mm, card_h_mm)
                img_w = cls._clamp(card_w_mm * 0.23, 14.0, 24.0, 20.0)
                img_h = cls._clamp(card_h_mm * 0.42, 20.0, 35.0, 25.0)
                x_mm = cls._clamp(label_box['x_mm'] + label_box['w_mm'] + 1.5, 0.0, card_w_mm - img_w, max(0.0, card_w_mm - img_w - 3.0))
                y_mm = cls._clamp(label_box['y_mm'] - 0.6, 0.0, card_h_mm - img_h, 4.0)

                mappings[fname] = {
                    'x_mm': round(x_mm, 2),
                    'y_mm': round(y_mm, 2),
                    'w_mm': round(img_w, 2),
                    'h_mm': round(img_h, 2),
                    'label_text': GenerateCardService._format_field_label(fname),
                    'placeholder': '[PHOTO]',
                    'show_key': False,
                }
                add_style_sample(best_item, 1)
                add_matched_style_sample(best_item, 2)

            # Fallback: map remaining image fields to detected image boxes (if any).
            remaining_image_fields = [
                f for f in image_fields
                if str(f.get('name', '')).strip() and str(f.get('name', '')).strip() not in mappings
            ]
            if remaining_image_fields and image_boxes:
                image_boxes = sorted(
                    image_boxes,
                    key=lambda b: ((b.get('x1', 0.0) - b.get('x0', 0.0)) * (b.get('y1', 0.0) - b.get('y0', 0.0))),
                    reverse=True,
                )
                for idx, field in enumerate(remaining_image_fields):
                    if idx >= len(image_boxes):
                        break
                    fname = str(field.get('name', '')).strip()
                    if not fname:
                        continue
                    box_src = image_boxes[idx]
                    box = cls._estimate_image_box(box_src, page.rect.width, page.rect.height, card_w_mm, card_h_mm)
                    mappings[fname] = {
                        'x_mm': box['x_mm'],
                        'y_mm': box['y_mm'],
                        'w_mm': box['w_mm'],
                        'h_mm': box['h_mm'],
                        'label_text': GenerateCardService._format_field_label(fname),
                        'placeholder': '[PHOTO]',
                        'show_key': False,
                    }

            best_family = ''
            best_count = 0
            family_source = matched_font_families if matched_font_families else font_families
            for fam, cnt in family_source.items():
                if cnt > best_count:
                    best_count = cnt
                    best_family = fam

            size_source = matched_font_pts if matched_font_pts else font_pts
            size_pt = cls._dominant_font_size(size_source, 11.0) if size_source else 11.0
            matched_style = {
                'font_size_pt': cls._clamp(size_pt, 8.0, 72.0, 11.0),
                'font_family': best_family,
            }

            return {
                'mappings': mappings,
                'matched_style': matched_style,
                'engine': 'pymupdf',
            }, None
        except Exception as exc:
            logger.error('PdfTemplateAnalyzerService.analyze_template failed: %s', exc, exc_info=True)
            return None, 'Failed to analyze PDF template text'
        finally:
            try:
                doc.close()
            except Exception:
                pass

    @classmethod
    def convert_template_pdf_to_docx(cls, template, side):
        """Convert uploaded design PDF side into a simple editable DOCX document."""
        try:
            import fitz
        except Exception:
            return None, 'PyMuPDF is not available for PDF conversion'

        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.section import WD_ORIENT
        except Exception:
            return None, 'python-docx is not available for Word conversion'

        pdf_field = template.back_pdf if side == 'back' else template.front_pdf
        pdf_path = getattr(pdf_field, 'path', None) if pdf_field else None
        if not pdf_path or not os.path.exists(pdf_path):
            return None, f'Upload {side} design PDF first'

        doc = None
        try:
            doc = fitz.open(pdf_path)
            if doc.page_count < 1:
                return None, 'PDF has no pages'

            page = doc.load_page(0)
            word_doc = Document()

            # Match page orientation so editing starts from a similar canvas.
            section = word_doc.sections[0]
            page_w_in = max(4.0, min(20.0, float(page.rect.width) / 72.0))
            page_h_in = max(4.0, min(20.0, float(page.rect.height) / 72.0))
            if page_w_in > page_h_in:
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(page_w_in)
                section.page_height = Inches(page_h_in)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Inches(page_w_in)
                section.page_height = Inches(page_h_in)

            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)

            lines = cls._extract_line_items(page)
            prev_y = None
            for item in lines:
                y = float(item.get('y0', 0.0))
                if prev_y is not None:
                    gap = y - prev_y
                    # Add a visual break for larger vertical gaps between blocks.
                    if gap > max(12.0, float(item.get('font_size', 11.0)) * 1.6):
                        word_doc.add_paragraph('')

                p = word_doc.add_paragraph()
                run = p.add_run(str(item.get('text', '') or ''))
                font_size = cls._clamp(float(item.get('font_size', 11.0)), 8.0, 48.0, 11.0)
                run.font.size = Pt(font_size)
                font_name = str(item.get('font_name', '') or '')
                if font_name:
                    run.font.name = font_name.split(',')[0][:64]
                if 'bold' in font_name.lower():
                    run.bold = True
                prev_y = y

            # Add detected embedded images to keep editable photo placeholders handy.
            image_count = 0
            for img_meta in page.get_images(full=True):
                if image_count >= 6:
                    break
                try:
                    xref = img_meta[0]
                    extracted = doc.extract_image(xref)
                    img_bytes = extracted.get('image')
                    if not img_bytes:
                        continue
                    from io import BytesIO
                    p = word_doc.add_paragraph('')
                    p.add_run('Image ' + str(image_count + 1) + ':')
                    word_doc.add_picture(BytesIO(img_bytes), width=Inches(1.6))
                    image_count += 1
                except Exception:
                    continue

            out = io.BytesIO()
            word_doc.save(out)
            out.seek(0)
            return out.getvalue(), None
        except Exception as exc:
            logger.error('PdfTemplateAnalyzerService.convert_template_pdf_to_docx failed: %s', exc, exc_info=True)
            return None, 'Failed to convert PDF to editable Word file'
        finally:
            try:
                if doc is not None:
                    doc.close()
            except Exception:
                pass

    @classmethod
    def build_editable_design_model(cls, template, side, card_w_mm, card_h_mm):
        """Return editable line/image blocks in card-mm coordinates for in-place editor mode."""
        try:
            import fitz
        except Exception:
            return None, 'PyMuPDF is not available for editable conversion'

        pdf_field = template.back_pdf if side == 'back' else template.front_pdf
        pdf_path = getattr(pdf_field, 'path', None) if pdf_field else None
        if not pdf_path or not os.path.exists(pdf_path):
            return None, f'Upload {side} design PDF first'

        doc = None
        try:
            doc = fitz.open(pdf_path)
            if doc.page_count < 1:
                return None, 'PDF has no pages'

            page = doc.load_page(0)
            page_w_pt = float(page.rect.width)
            page_h_pt = float(page.rect.height)
            sx = card_w_mm / float(max(1.0, page_w_pt))
            sy = card_h_mm / float(max(1.0, page_h_pt))

            lines = []
            for item in cls._extract_line_items(page):
                x0 = float(item.get('x0', 0.0))
                y0 = float(item.get('y0', 0.0))
                x1 = float(item.get('x1', x0 + 20.0))
                y1 = float(item.get('y1', y0 + 8.0))

                raw_family = str(item.get('font_name', '') or '').strip()
                ui_family = cls._map_pdf_font_family(raw_family)
                family = ui_family or raw_family.split(',')[0][:64] or 'Arial'
                weight = '700' if re.search(r'(bold|black|heavy)', raw_family, re.I) else ('600' if re.search(r'semibold|demibold|medium', raw_family, re.I) else '400')
                color_hex = '#' + GenerateCardService._normalize_hex_color(item.get('font_color_hex') or '111111', default='111111')
                text_align = cls._infer_line_alignment(x0, x1, page_w_pt)

                lines.append({
                    'text': str(item.get('text', '') or ''),
                    'x_mm': round(cls._clamp(x0 * sx, 0.0, card_w_mm, 0.0), 2),
                    'y_mm': round(cls._clamp(y0 * sy, 0.0, card_h_mm, 0.0), 2),
                    'w_mm': round(cls._clamp((x1 - x0) * sx, 2.0, card_w_mm, 20.0), 2),
                    'h_mm': round(cls._clamp((y1 - y0) * sy, 2.0, card_h_mm, 8.0), 2),
                    'font_size_pt': round(cls._clamp(float(item.get('font_size', 11.0)), 6.0, 72.0, 11.0), 2),
                    'font_family': family,
                    'font_weight': weight,
                    'font_color_hex': color_hex,
                    'text_align': text_align,
                })

            images = []
            seen_rects = set()

            def _append_image_rect(rect, image_bytes):
                if not image_bytes:
                    return False
                rkey = (
                    round(float(rect.x0), 2),
                    round(float(rect.y0), 2),
                    round(float(rect.x1), 2),
                    round(float(rect.y1), 2),
                )
                if rkey in seen_rects:
                    return False
                seen_rects.add(rkey)

                b64 = base64.b64encode(image_bytes).decode('ascii')
                data_url = 'data:image/png;base64,' + b64
                images.append({
                    'x_mm': round(cls._clamp(float(rect.x0) * sx, 0.0, card_w_mm, 0.0), 2),
                    'y_mm': round(cls._clamp(float(rect.y0) * sy, 0.0, card_h_mm, 0.0), 2),
                    'w_mm': round(cls._clamp((float(rect.x1) - float(rect.x0)) * sx, 2.0, card_w_mm, 20.0), 2),
                    'h_mm': round(cls._clamp((float(rect.y1) - float(rect.y0)) * sy, 2.0, card_h_mm, 20.0), 2),
                    'data_url': data_url,
                })
                return True

            for img_meta in page.get_images(full=True):
                if len(images) >= 8:
                    break
                try:
                    xref = img_meta[0]
                    rects = page.get_image_rects(xref)
                    if not rects:
                        continue
                    rect = rects[0]
                    rkey = (round(float(rect.x0), 2), round(float(rect.y0), 2), round(float(rect.x1), 2), round(float(rect.y1), 2))
                    if rkey in seen_rects:
                        continue
                    seen_rects.add(rkey)

                    extracted = doc.extract_image(xref)
                    img_bytes = extracted.get('image')
                    if not img_bytes:
                        continue
                    _append_image_rect(rect, img_bytes)
                except Exception:
                    continue

            # Fallback 1: some PDFs don't expose xref images but do expose image blocks.
            if len(images) < 8:
                for block in cls._extract_image_boxes(page):
                    if len(images) >= 8:
                        break
                    try:
                        rect = fitz.Rect(
                            float(block.get('x0', 0.0)),
                            float(block.get('y0', 0.0)),
                            float(block.get('x1', 0.0)),
                            float(block.get('y1', 0.0)),
                        )
                        if rect.width < 4 or rect.height < 4:
                            continue
                        clip = rect & page.rect
                        if clip.width < 4 or clip.height < 4:
                            continue
                        pix = page.get_pixmap(clip=clip, matrix=fitz.Matrix(2, 2), alpha=False)
                        img_bytes = pix.tobytes('png')
                        _append_image_rect(clip, img_bytes)
                    except Exception:
                        continue

            # Fallback 2: preserve mapped photo/signature regions if still missing.
            if len(images) < 8 and isinstance(template.field_mappings, dict):
                side_mappings = template.field_mappings.get(side) or {}
                if isinstance(side_mappings, dict):
                    for field_name, mapping in side_mappings.items():
                        if len(images) >= 8:
                            break
                        if not GenerateCardService._is_image_field('', str(field_name or '')):
                            continue
                        if not isinstance(mapping, dict):
                            continue
                        try:
                            x_mm = float(mapping.get('x_mm', 0.0))
                            y_mm = float(mapping.get('y_mm', 0.0))
                            w_mm = float(mapping.get('w_mm', 0.0))
                            h_mm = float(mapping.get('h_mm', 0.0))
                        except Exception:
                            continue
                        if w_mm <= 0.5 or h_mm <= 0.5:
                            continue

                        x0_pt = x_mm / sx
                        y0_pt = y_mm / sy
                        x1_pt = (x_mm + w_mm) / sx
                        y1_pt = (y_mm + h_mm) / sy
                        rect = fitz.Rect(x0_pt, y0_pt, x1_pt, y1_pt) & page.rect
                        if rect.width < 4 or rect.height < 4:
                            continue
                        try:
                            pix = page.get_pixmap(clip=rect, matrix=fitz.Matrix(2, 2), alpha=False)
                            img_bytes = pix.tobytes('png')
                            _append_image_rect(rect, img_bytes)
                        except Exception:
                            continue

            return {
                'engine': 'pymupdf-editable',
                'page_mm': {
                    'width': round(card_w_mm, 2),
                    'height': round(card_h_mm, 2),
                },
                'lines': lines,
                'images': images,
            }, None
        except Exception as exc:
            logger.error('PdfTemplateAnalyzerService.build_editable_design_model failed: %s', exc, exc_info=True)
            return None, 'Failed to build editable design model'
        finally:
            try:
                if doc is not None:
                    doc.close()
            except Exception:
                pass

    @classmethod
    def _extract_line_items(cls, page):
        text_dict = page.get_text('dict') or {}
        out = []
        idx = 0
        for block in text_dict.get('blocks', []):
            if block.get('type') != 0:
                continue
            for line in block.get('lines', []):
                spans = line.get('spans', []) or []
                if not spans:
                    continue

                text_parts = []
                x0 = y0 = None
                x1 = y1 = None
                font_name_counts = {}
                color_counts = {}
                size_weight_sum = 0.0
                size_sum = 0.0

                for sp in spans:
                    txt = str(sp.get('text', '') or '')
                    if txt:
                        text_parts.append(txt)

                    bbox = sp.get('bbox') or [0, 0, 0, 0]
                    sx0, sy0, sx1, sy1 = [float(v) for v in bbox]
                    x0 = sx0 if x0 is None else min(x0, sx0)
                    y0 = sy0 if y0 is None else min(y0, sy0)
                    x1 = sx1 if x1 is None else max(x1, sx1)
                    y1 = sy1 if y1 is None else max(y1, sy1)

                    fname = str(sp.get('font', '') or '').strip()
                    if fname:
                        font_name_counts[fname] = font_name_counts.get(fname, 0) + max(1, len(txt.strip()))

                    color_hex = cls._pdf_color_to_hex(sp.get('color', None), default='111111')
                    if color_hex:
                        color_counts[color_hex] = color_counts.get(color_hex, 0) + max(1, len(txt.strip()))

                    size = float(sp.get('size', 11.0) or 11.0)
                    w = max(1, len(txt.strip()))
                    size_weight_sum += w
                    size_sum += size * w

                raw_text = ''.join(text_parts).strip()
                if not raw_text:
                    continue

                font_name = ''
                font_count = 0
                for fn, cnt in font_name_counts.items():
                    if cnt > font_count:
                        font_name = fn
                        font_count = cnt

                font_color_hex = '111111'
                color_count = 0
                for hex_val, cnt in color_counts.items():
                    if cnt > color_count:
                        font_color_hex = hex_val
                        color_count = cnt

                font_size = (size_sum / size_weight_sum) if size_weight_sum > 0 else 11.0
                out.append({
                    'index': idx,
                    'text': raw_text,
                    'x0': float(x0 or 0.0),
                    'y0': float(y0 or 0.0),
                    'x1': float(x1 or 0.0),
                    'y1': float(y1 or 0.0),
                    'font_name': font_name,
                    'font_size': font_size,
                    'font_color_hex': font_color_hex,
                })
                idx += 1

        out.sort(key=lambda it: (it.get('y0', 0.0), it.get('x0', 0.0)))
        return out

    @classmethod
    def _extract_image_boxes(cls, page):
        text_dict = page.get_text('dict') or {}
        out = []
        for block in text_dict.get('blocks', []):
            if block.get('type') != 1:
                continue
            bbox = block.get('bbox') or [0, 0, 0, 0]
            try:
                x0, y0, x1, y1 = [float(v) for v in bbox]
            except Exception:
                continue
            if x1 <= x0 or y1 <= y0:
                continue
            out.append({'x0': x0, 'y0': y0, 'x1': x1, 'y1': y1})
        return out

    @classmethod
    def _pdf_color_to_hex(cls, value, default='111111'):
        # PyMuPDF span colors are typically packed integers (0xRRGGBB).
        try:
            if isinstance(value, (tuple, list)) and len(value) >= 3:
                parts = []
                for ch in value[:3]:
                    c = float(ch)
                    if c <= 1.0:
                        c = c * 255.0
                    parts.append(max(0, min(255, int(round(c)))))
                return f'{parts[0]:02X}{parts[1]:02X}{parts[2]:02X}'

            packed = int(value)
            if packed < 0:
                packed = packed & 0xFFFFFF
            return f'{packed & 0xFFFFFF:06X}'
        except Exception:
            return default

    @classmethod
    def _infer_line_alignment(cls, x0, x1, page_w_pt):
        width = max(0.0, float(x1) - float(x0))
        page_w = max(1.0, float(page_w_pt))
        left_gap = max(0.0, float(x0))
        right_gap = max(0.0, page_w - float(x1))

        # Symmetric margins usually indicate centered text blocks.
        center_tol = max(10.0, page_w * 0.04)
        if width <= page_w * 0.86 and abs(left_gap - right_gap) <= center_tol:
            return 'center'

        # Tight right margin indicates right alignment.
        if right_gap <= max(8.0, page_w * 0.03) and left_gap > (right_gap + 6.0):
            return 'right'

        return 'left'

    @classmethod
    def _normalize(cls, value):
        txt = str(value or '').lower()
        txt = re.sub(r'\(.*?\)', ' ', txt)
        txt = re.sub(r'[_\-]+', ' ', txt)
        txt = re.sub(r'[^a-z0-9\s]', ' ', txt)
        txt = re.sub(r'\s+', ' ', txt).strip()
        return txt

    @classmethod
    def _normalize_key(cls, value):
        return cls._normalize(value).replace(' ', '')

    @classmethod
    def _split_key_value_text(cls, raw_text):
        text = str(raw_text or '').strip()
        if not text:
            return {'has_delimiter': False, 'key_text': '', 'value_text': ''}

        m = re.match(r'^(.{1,48}?)(?:\s*[:=\-|]\s*|\s{2,})(.{0,200})$', text)
        if not m:
            return {'has_delimiter': False, 'key_text': '', 'value_text': text}

        key_text = str(m.group(1) or '').strip()
        value_text = str(m.group(2) or '').strip()
        if not key_text or len(key_text) < 2:
            return {'has_delimiter': False, 'key_text': '', 'value_text': text}
        return {'has_delimiter': True, 'key_text': key_text, 'value_text': value_text}

    @classmethod
    def _score_field_match(cls, text, field_name):
        text_norm = cls._normalize(text)
        field_norm = cls._normalize(field_name)
        text_key = cls._normalize_key(text)
        field_key = cls._normalize_key(field_name)

        if not text_norm or not field_norm:
            return 0.0
        if text_norm == field_norm or text_key == field_key:
            return 1.0
        if field_norm in text_norm and len(field_norm) >= 3:
            return 0.95
        if text_norm in field_norm and len(text_norm) >= 3:
            return 0.88

        t_tokens = [t for t in text_norm.split(' ') if t]
        f_tokens = [t for t in field_norm.split(' ') if t]
        token_score = 0.0
        if t_tokens and f_tokens:
            common = sum(1 for t in f_tokens if t in t_tokens)
            if common > 0:
                token_score = common / float(max(len(t_tokens), len(f_tokens)))

        ratio = SequenceMatcher(None, text_key, field_key).ratio()
        return max(token_score * 0.82, ratio * 0.86)

    @classmethod
    def _estimate_text_box(cls, item, page_w_pt, page_h_pt, card_w_mm, card_h_mm, prefer_value_only=False):
        sx = card_w_mm / float(max(1.0, page_w_pt))
        sy = card_h_mm / float(max(1.0, page_h_pt))

        x0 = float(item.get('x0', 0.0))
        y0 = float(item.get('y0', 0.0))
        x1 = float(item.get('x1', x0 + 20.0))
        y1 = float(item.get('y1', y0 + 8.0))

        raw_w_mm = max(4.0, (x1 - x0) * sx)
        raw_h_mm = max(2.0, (y1 - y0) * sy)

        if prefer_value_only:
            w_mm = max(card_w_mm * 0.26, raw_w_mm * 1.5)
        else:
            w_mm = max(16.0, raw_w_mm * 1.32)
        h_mm = max(4.8, raw_h_mm * 2.0)

        w_mm = cls._clamp(w_mm, 12.0, card_w_mm * 0.78, min(34.0, card_w_mm * 0.5))
        h_mm = cls._clamp(h_mm, 3.5, card_h_mm * 0.26, 8.0)

        x_mm = cls._clamp(x0 * sx, 0.0, card_w_mm - w_mm, 0.0)
        y_mm = cls._clamp((y0 * sy) - (h_mm * 0.15), 0.0, card_h_mm - h_mm, 0.0)
        return {
            'x_mm': round(x_mm, 2),
            'y_mm': round(y_mm, 2),
            'w_mm': round(w_mm, 2),
            'h_mm': round(h_mm, 2),
        }

    @classmethod
    def _estimate_image_box(cls, item, page_w_pt, page_h_pt, card_w_mm, card_h_mm):
        sx = card_w_mm / float(max(1.0, page_w_pt))
        sy = card_h_mm / float(max(1.0, page_h_pt))

        x0 = float(item.get('x0', 0.0))
        y0 = float(item.get('y0', 0.0))
        x1 = float(item.get('x1', x0 + 40.0))
        y1 = float(item.get('y1', y0 + 50.0))

        w_mm = cls._clamp((x1 - x0) * sx, 12.0, card_w_mm * 0.45, 20.0)
        h_mm = cls._clamp((y1 - y0) * sy, 16.0, card_h_mm * 0.6, 25.0)
        x_mm = cls._clamp(x0 * sx, 0.0, card_w_mm - w_mm, max(0.0, card_w_mm - w_mm - 2.0))
        y_mm = cls._clamp(y0 * sy, 0.0, card_h_mm - h_mm, 4.0)

        return {
            'x_mm': round(x_mm, 2),
            'y_mm': round(y_mm, 2),
            'w_mm': round(w_mm, 2),
            'h_mm': round(h_mm, 2),
        }

    @classmethod
    def _map_pdf_font_family(cls, raw_name):
        src = str(raw_name or '').lower()
        if not src:
            return ''
        if 'noto sans condensed' in src:
            return 'Noto Sans Condensed'
        if 'noto sans' in src:
            return 'Noto Sans'
        if 'nirmala' in src:
            return 'Nirmala UI'
        if 'times' in src:
            return 'Times New Roman'
        if 'calibri' in src:
            return 'Calibri'
        if 'arial rounded' in src:
            return 'Arial Rounded MT Bold'
        if 'arial black' in src:
            return 'Arial Black'
        if 'arial narrow' in src:
            return 'Arial Narrow'
        if 'arial unicode' in src:
            return 'Arial Unicode MS'
        if 'arial' in src or 'helvetica' in src:
            return 'Arial'
        return ''

    @classmethod
    def _median(cls, values, fallback):
        vals = sorted(float(v) for v in (values or []) if v is not None)
        if not vals:
            return fallback
        n = len(vals)
        mid = n // 2
        if n % 2 == 1:
            return vals[mid]
        return (vals[mid - 1] + vals[mid]) / 2.0

    @classmethod
    def _dominant_font_size(cls, values, fallback):
        vals = [float(v) for v in (values or []) if v is not None]
        if not vals:
            return fallback

        buckets = {}
        for v in vals:
            k = int(round(v))
            buckets[k] = buckets.get(k, 0) + 1

        best_size = fallback
        best_count = -1
        for k, cnt in buckets.items():
            if cnt > best_count:
                best_size = float(k)
                best_count = cnt

        return best_size

    @classmethod
    def _clamp(cls, value, min_v, max_v, fallback):
        try:
            num = float(value)
        except Exception:
            return fallback
        if num < min_v:
            return min_v
        if num > max_v:
            return max_v
        return num
