"""
Word Export — Images Mixin

Image embedding, VML conversion, and empty placeholder rendering
for Word document generation.
"""
import logging
import re
from io import BytesIO

from django.core.files.storage import default_storage

from .utils import is_valid_image_path

logger = logging.getLogger(__name__)


class WordImagesMixin:
    """Mixin providing image handling methods for Word exports."""

    # Border is intentionally limited to portrait photos only.
    BORDERED_IMAGE_SUBTYPES = {'photo', 'rel_photo', 'mother_photo', 'father_photo'}
    PHOTO_BORDER_PX = 1  # 1px border fallback
    PHOTO_BORDER_COLOR = (0, 0, 0)
    WORD_BORDER_PT = 0.5

    def _should_add_photo_border(self, image_subtype=None, field_name=None):
        """Return True when the image should be rendered with a 1pt border."""
        subtype = str(image_subtype or '').strip().lower()
        if subtype in self.BORDERED_IMAGE_SUBTYPES:
            return True

        # Fallback: if subtype is missing, infer from field name safely.
        name = str(field_name or '').strip().lower()
        if not name:
            return False

        if re.search(r'\b(?:father|mother)\b.*\b(?:photo|image|pic|picture)\b', name):
            return True
        if re.search(r'\b(?:photo|image|pic|picture)\b', name) and not re.search(r'\b(?:signature|sign|barcode|qr)\b', name):
            return True
        return False

    @staticmethod
    def _build_word_image_stream_from_image(src_img, Image, ImageOps):
        """Build the image stream with border drawing from an already opened PIL Image.
        
        Uses Pillow native ImageDraw.rectangle for fast C-accelerated border drawing
        instead of slow pure-Python loops.
        """
        from PIL import ImageDraw

        if src_img.mode != 'RGB':
            src_img = src_img.convert('RGB')
        bordered = src_img.copy()
        w, h = bordered.size
        if w >= 2 and h >= 2:
            draw = ImageDraw.Draw(bordered)
            border_px = WordImagesMixin.PHOTO_BORDER_PX
            edge = WordImagesMixin.PHOTO_BORDER_COLOR
            max_t = min(border_px, w // 2, h // 2)
            if max_t > 0:
                # Draws outline expanding inwards, ensuring pixel-for-pixel visual equivalence
                draw.rectangle([0, 0, w - 1, h - 1], outline=edge, width=max_t)

        out_stream = BytesIO()
        bordered.save(out_stream, format='JPEG', quality=95)
        out_stream.seek(0)
        return out_stream

    @staticmethod
    def _build_word_image_stream(img_data, Image, ImageOps, add_photo_border=False):
        """Return a BytesIO stream for Word embedding.

        Some Word compatibility renderers ignore VML stroke on image shapes.
        For photo-like columns, add a 1pt-equivalent inner border in the
        bitmap (without changing dimensions) so border remains visible.
        
        Backward compatible wrapper that reuses the optimized image processor.
        """
        if not add_photo_border:
            return BytesIO(img_data)

        # Optimization B: Reuse already opened image stream safely.
        with Image.open(BytesIO(img_data)) as src_img:
            src_img.load()
            return WordImagesMixin._build_word_image_stream_from_image(src_img, Image, ImageOps)

    def _add_image_to_cell(self, cell, img_path, Cm, Pt, RGBColor,
                           WD_ALIGN_PARAGRAPH, parse_xml, nsdecls, Image, ImageOps,
                           fixed_width_cm=None, fixed_height_cm=None,
                           image_subtype=None, field_name=None, cancel_check=None,
                           prefetched_bytes=None, skip_storage_read=False):
        """Add an image to a cell using VML for Word 2007+ compatibility.

        Border is drawn natively by Word's VML renderer (stroked="t") — no
        Pillow re-encode, no extra CPU, and the border scales correctly at any
        DPI and zoom level.

        For missing/pending images: draws an empty bordered rectangle.

        Args:
            prefetched_bytes:   Pre-read image bytes from the concurrent batch
                                pre-fetch (may be None if the image wasn't found).
            skip_storage_read:  When True the function trusts prefetched_bytes
                                as the authoritative source and never opens
                                default_storage — even if prefetched_bytes is None.
        """
        if fixed_width_cm is None:
            fixed_width_cm = self.IMAGE_DEFAULT_WIDTH_CM
        if fixed_height_cm is None:
            fixed_height_cm = self.IMAGE_HEIGHT_CM

        # Determine if this column needs a photo border.
        add_photo_border = self._should_add_photo_border(
            image_subtype=image_subtype,
            field_name=field_name,
        )

        if add_photo_border:
            # 30 twips margins on all sides (about 1.5pt) to ensure the 0.5pt border is fully visible on all sides
            self._set_cell_margins(cell, parse_xml, nsdecls, 30, 30, 30, 30)
            # Subtract 2 * 0.5pt (approx 0.035 cm) from image content dimensions so the total size including 0.5pt border is exactly target size (e.g. 2.5 cm)
            img_h = max(0.1, fixed_height_cm - 0.035)
            img_w = max(0.1, fixed_width_cm - 0.035)
        else:
            self._set_cell_margins(cell, parse_xml, nsdecls, 0, 0, 0, 0)
            img_h = fixed_height_cm
            img_w = fixed_width_cm

        self._set_cell_vertical_align(cell, parse_xml, nsdecls)

        if callable(cancel_check) and cancel_check():
            raise Exception('Export cancelled')

        # ── Resolve image bytes ───────────────────────────────────────────────
        # Priority: pre-fetched bytes (from concurrent batch) → storage read.
        if skip_storage_read:
            # Batch pre-fetch was the authoritative IO; trust its result.
            img_data = prefetched_bytes
        else:
            # No pre-fetch available — fall back to a direct storage read.
            img_data = prefetched_bytes  # will be None here in normal flow
            if img_data is None and is_valid_image_path(img_path):
                try:
                    with default_storage.open(img_path, 'rb') as img_file:
                        img_data = img_file.read()
                except Exception:
                    img_data = None

        if img_data and len(img_data) >= 100:
            try:
                if add_photo_border:
                    # Draw Pillow border and encode as JPEG (extremely fast + high quality)
                    # We do not use slower PNG format to avoid touching 100% CPU usage.
                    img_stream = self._build_word_image_stream(img_data, Image, ImageOps, add_photo_border=True)
                else:
                    # Non-photo: embed raw bytes directly (0% CPU overhead, no re-compression)
                    img_stream = BytesIO(img_data)

                para = cell.paragraphs[0]
                run = para.add_run()
                inline_shape = run.add_picture(
                    img_stream,
                    height=Cm(img_h),
                    width=Cm(img_w),
                )
                img_stream.close()

                if add_photo_border:
                    self._add_drawingml_border(inline_shape, parse_xml)

                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set paragraph spacing with 40 twips (2pt) padding before/after and lineRule="auto" to prevent MS Word clipping top/bottom edges of the image
                pPr = para._p.get_or_add_pPr()
                from docx.oxml.ns import qn as _qn
                for existing in pPr.findall(_qn('w:spacing')):
                    pPr.remove(existing)
                spacing = parse_xml(r'<w:spacing {} w:before="40" w:after="40" w:lineRule="auto"/>'.format(nsdecls('w')))
                pPr.append(spacing)
                return
            except Exception as e:
                logger.warning("Word export: Image embed error for %s: %s", img_path, e)
        else:
            if img_data is not None:
                logger.warning("Word export: Image too small/empty: %s", img_path)
            elif is_valid_image_path(img_path) and not skip_storage_read:
                logger.warning("Word export: Image missing or unreadable: %s", img_path)

        # Missing/pending image → draw an empty bordered rectangle
        # (no placeholder image, no text — just empty space with a border).
        # Keep placeholder at least ROW_HEIGHT_CM so student photo image rows remain visually
        # consistent even when the source image is missing. For relation photos, keep them at
        # their defined smaller height.
        is_rel_photo = (image_subtype == 'rel_photo') or (field_name and any(k in field_name.lower() for k in ('rel', 'mother', 'father', 'parent', 'guardian')))
        if is_rel_photo:
            placeholder_h = fixed_height_cm or 2.0
        else:
            placeholder_h = max(fixed_height_cm or 0, getattr(self, 'ROW_HEIGHT_CM', 2.5))
        self._add_empty_image_box(cell, Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                                   parse_xml, nsdecls, fixed_width_cm, placeholder_h)

    def _add_empty_image_box(self, cell, Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                              parse_xml, nsdecls, fixed_width_cm, fixed_height_cm=None):
        """Draw an empty bordered rectangle for missing/pending images.
        
        Uses a 1-row, 1-col inner table with fixed dimensions and
        a thin black border to represent an empty image placeholder.
        """
        from docx.enum.table import WD_TABLE_ALIGNMENT

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set paragraph spacing with 40 twips (2pt) padding before/after and lineRule="auto" to prevent MS Word clipping top/bottom edges of the VML box
        pPr = para._p.get_or_add_pPr()
        from docx.oxml.ns import qn as _qn
        for existing in pPr.findall(_qn('w:spacing')):
            pPr.remove(existing)
        spacing = parse_xml(r'<w:spacing {} w:before="40" w:after="40" w:lineRule="auto"/>'.format(nsdecls('w')))
        pPr.append(spacing)

        # Clear any default text
        for run in para.runs:
            run.clear()

        # Create a tiny inline table as the bordered box
        # We use VML rectangle via a Run for better compatibility
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement

        # Width/Height in EMU (1 cm = 360000 EMU)
        effective_h = fixed_height_cm if fixed_height_cm else self.IMAGE_HEIGHT_CM
        box_w_emu = int(fixed_width_cm * 360000)
        box_h_emu = int(effective_h * 360000)
        box_w_pt = fixed_width_cm * 28.3465
        box_h_pt = effective_h * 28.3465

        # Create VML rectangle shape directly
        run = para.add_run()
        pict_xml = (
            '<w:pict '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office">'
            f'<v:rect style="width:{box_w_pt:.1f}pt;height:{box_h_pt:.1f}pt" '
            f'filled="f" strokecolor="#000000" strokeweight="{self.WORD_BORDER_PT:.1f}pt">'
            '</v:rect></w:pict>'
        )
        from lxml import etree
        pict_elem = etree.fromstring(pict_xml)
        run._r.append(pict_elem)

    def _add_drawingml_border(self, inline_shape, parse_xml):
        """Add a native DrawingML border (outline) to the inline shape."""
        namespaces = {
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        }
        inline_el = inline_shape._inline
        spPr = inline_el.find('.//pic:spPr', namespaces=namespaces)
        if spPr is not None:
            # Remove existing outlines if any
            for existing in spPr.findall('a:ln', namespaces=namespaces):
                spPr.remove(existing)
            width_emu = int(getattr(self, 'WORD_BORDER_PT', 1.0) * 12700)
            ln_xml = (
                f'<a:ln w="{width_emu}" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                '  <a:solidFill>'
                '    <a:srgbClr val="000000"/>'
                '  </a:solidFill>'
                '</a:ln>'
            )
            ln_el = parse_xml(ln_xml)
            spPr.append(ln_el)
    
    def _convert_to_vml(self, run, inline_shape, add_border=False):
        """Convert an inline DrawingML image to VML for backward compatibility.
        
        VML images are visible in Word's Normal/Draft view and compatible
        with Word 97-2003 format (.doc).  This replaces the <w:drawing>
        element with a <w:pict> VML element referencing the same image
        relationship.
        """
        from lxml import etree  # type: ignore[attr-defined]
        from docx.oxml.ns import qn
        
        # Get dimensions from inline shape (EMU → points)
        width_pt = inline_shape.width / 914400.0 * 72.0
        height_pt = inline_shape.height / 914400.0 * 72.0
        
        # Extract relationship ID from DrawingML blip element
        drawing_elem = run._r.find(qn('w:drawing'))
        blip = drawing_elem.find('.//' + qn('a:blip'))
        rId = blip.get(qn('r:embed'))
        
        # Remove the DrawingML element from the run
        run._r.remove(drawing_elem)
        
        # Create VML <w:pict> element (universally compatible).
        # Use <v:rect> so stroke is consistently rendered around the image.
        border_attrs = (
            f'stroked="t" strokecolor="#000000" strokeweight="{self.WORD_BORDER_PT:.1f}pt"'
            if add_border else
            'stroked="f"'
        )
        vml_xml = (
            '<w:pict '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f'<v:rect {border_attrs} '
            f'style="width:{width_pt:.1f}pt;height:{height_pt:.1f}pt">'
            f'<v:imagedata r:id="{rId}" o:title=""/>'
            '</v:rect></w:pict>'
        )
        pict_elem = etree.fromstring(vml_xml)
        run._r.append(pict_elem)
