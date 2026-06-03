"""
Word Export Module

Handles DOCX file generation for ID card data.
This module is READ-ONLY - it never mutates data.

Features:
- Landscape A4 format with 0.5cm margins
- Header with institution name, table name, date, and branding
- Footer with note, timestamp, and page numbers
- Table with text fields on left, image fields on right
- Auto-sized columns based on content
- 7 entries per page with proper pagination
- Image embedding with borders
- Phase 4: Uses THUMBNAILS for optimized export performance
- Output format: .docx (Word 2007+) only
"""
import logging
from typing import Optional
from dataclasses import dataclass

from django.http import HttpResponse
from django.db.models import QuerySet

from .word_styles import WordStylesMixin
from .word_tables import WordTablesMixin
from .word_images import WordImagesMixin

from .utils import (
    separate_fields_by_type,
    generate_export_filename,
    sort_cards_for_export,
    get_class_field_name,
    
)

logger = logging.getLogger(__name__)


@dataclass
class WordExportResult:
    """Result of a Word export operation."""
    success: bool
    message: str = ''
    response: Optional[HttpResponse] = None
    filename: str = ''
    card_count: int = 0


class WordExporter(WordStylesMixin, WordTablesMixin, WordImagesMixin):
    """
    Word document exporter — combines styles, tables, and image handling.
    
    Features:
    - Landscape orientation for data tables
    - Exports both text and image fields
    - Professional formatting with headers/footers
    - 6 entries per page (safe for wrapped column headers)
    
    Usage:
        exporter = WordExporter()
        result = exporter.export_cards(table, cards)
        if result.success:
            return result.response
    """
    
    ENTRIES_PER_PAGE = 6
    IMAGE_HEIGHT_CM = 2.5
    IMAGE_DEFAULT_WIDTH_CM = 1.9  # 3:4 portrait default
    ROW_HEIGHT_CM = 2.5
    PAGE_WIDTH_CM = 28.7  # A4 landscape usable width with 0.5 cm side margins
    # Page height budget (landscape A4 = 21cm)
    PAGE_HEIGHT_CM = 21.0
    TOP_MARGIN_CM = 0.5
    BOTTOM_MARGIN_CM = 0.5
    FOOTER_EFFECTIVE_CM = 0.5  # footer_distance setting
    # Phase 3: DOCX always uses ORIGINAL images for full quality print.
    # PDF uses thumbnails. ZIP uses originals.
    
    def export_cards(
        self,
        table,
        cards: QuerySet,
        doc_format: str = 'docx',
        status: str = '',
        template_id: int = None,
        allow_large: bool = False,
        progress_callback=None,
        cancel_check=None,
        user=None,
    ) -> WordExportResult:
        """
        Export cards to Word format (.docx only).
        
        Args:
            table: IDCardTable instance
            cards: QuerySet of IDCard instances
            doc_format: Ignored — always outputs .docx (Word 2007+)
            
        Returns:
            WordExportResult with HttpResponse if successful
        """
        try:
            from docx import Document
            from docx.shared import Cm, Pt, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            from docx.oxml.ns import nsdecls, qn
            from docx.oxml import parse_xml, OxmlElement
        except ImportError:
            return WordExportResult(
                success=False,
                message='python-docx library not installed. Run: pip install python-docx'
            )
        
        try:
            from PIL import Image, ImageOps
        except ImportError:
            return WordExportResult(
                success=False,
                message='Pillow library not installed. Run: pip install Pillow'
            )
        
        if not cards.exists():
            return WordExportResult(
                success=False,
                message='No cards to export!'
            )

        try:
            # Separate fields by type: text fields first, then image fields
            field_info = separate_fields_by_type(table.fields or [])
            text_fields = field_info['text']
            image_fields = field_info['image']
            ordered_fields = text_fields + image_fields

            # Column count guard — beyond 25 columns the document is unreadable
            MAX_WORD_COLUMNS = 25
            if len(ordered_fields) > MAX_WORD_COLUMNS:
                return WordExportResult(
                    success=False,
                    message=(
                        f'Word export supports a maximum of {MAX_WORD_COLUMNS} columns '
                        f'({len(ordered_fields)} selected). Remove some fields from the '
                        f'table configuration to proceed, or use Excel export instead.'
                    )
                )

            # Get institution name
            institution_name = "Institution"
            if table.group and table.group.client:
                institution_name = table.group.client.name
            
            # Create document
            doc = Document()
            
            # Setup page (landscape A4 with margins)
            self._setup_page(doc, Cm, WD_ORIENT, parse_xml, nsdecls)
            
            # Add header
            self._add_header(
                doc, institution_name, table.name,
                Cm, Pt, RGBColor, WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH,
                parse_xml, nsdecls
            )
            
            # Add footer
            self._add_footer(
                doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                parse_xml, nsdecls, OxmlElement, qn,
                template_id=template_id,
            )
            
            # Sort cards for export (Class → Name, or Name only)
            cards_list = sort_cards_for_export(cards, table.fields)
            if callable(cancel_check) and cancel_check():
                return WordExportResult(success=False, message='Export cancelled')
            num_cols = 1 + len(ordered_fields)  # Sr No + fields
            column_widths = self._calculate_column_widths(
                ordered_fields, cards_list, num_cols, cancel_check=cancel_check
            )
            if column_widths is None:
                return WordExportResult(success=False, message='Export cancelled')
            
            # Remove default empty paragraph
            if doc.paragraphs:
                p = doc.paragraphs[0]._element
                p.getparent().remove(p)
            
            # Create tables with data (page-break per N rows)
            class_field_name = get_class_field_name(table.fields)
            self._create_data_tables(
                doc, cards_list, ordered_fields, column_widths, num_cols,
                Cm, Pt, RGBColor, WD_TABLE_ALIGNMENT, WD_ALIGN_PARAGRAPH,
                parse_xml, nsdecls, OxmlElement, qn, Image, ImageOps,
                class_field_name=class_field_name,
                progress_callback=progress_callback,
                cancel_check=cancel_check,
            )
            
            # Set Word 97-2003 compatibility mode
            self._set_compatibility_mode(doc)
            
            # Save document to a temporary file to avoid building a giant
            # in-memory bytes object. Streaming back from disk reduces peak
            # memory and allows cooperative cancellation earlier.
            import tempfile
            import os
            from django.http import StreamingHttpResponse

            # Always output .docx (Word 2007+ OOXML format). The legacy .doc
            # format (Word 2003) has been removed — all modern Word versions
            # open .docx natively and it produces smaller, faster files.
            extension = 'docx'
            filename = generate_export_filename(table.name, extension, client_name=institution_name, status=status)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.' + extension)
            tmp_name = tmp.name
            tmp.close()
            try:
                # Write OOXML to disk (python-docx supports saving to filename)
                logger.info("DOCX_SAVE_START table=%s cards=%d file=%s", table.id, len(cards_list), tmp_name)
                import time
                save_started = time.monotonic()
                doc.save(tmp_name)
                logger.info(
                    "DOCX_SAVE_DONE table=%s cards=%d elapsed=%.2fs file=%s",
                    table.id,
                    len(cards_list),
                    time.monotonic() - save_started,
                    tmp_name,
                )

                file_size = 0
                try:
                    file_size = os.path.getsize(tmp_name)
                except Exception:
                    file_size = 0

                def _iter_file():
                    try:
                        with open(tmp_name, 'rb') as fh:
                            while True:
                                chunk = fh.read(1024 * 1024)
                                if not chunk:
                                    break
                                yield chunk
                    finally:
                        try:
                            os.unlink(tmp_name)
                        except Exception:
                            pass

                response = StreamingHttpResponse(_iter_file(), content_type=content_type)
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                if file_size:
                    response['Content-Length'] = str(file_size)

                return WordExportResult(
                    success=True,
                    response=response,
                    filename=filename,
                    card_count=len(cards_list)
                )
            except Exception:
                # Ensure temp file cleanup on error
                try:
                    os.unlink(tmp_name)
                except Exception:
                    pass
                raise
            
        except Exception as e:
            if isinstance(e, Exception) and 'cancel' in str(e).lower():
                return WordExportResult(success=False, message='Export cancelled')
            logger.error("Word export failed: %s", e, exc_info=True)
            return WordExportResult(
                success=False,
                message='Word export failed. Please try again or contact support.'
            )

