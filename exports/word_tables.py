"""
Word Export — Tables Mixin

Data table creation, column width calculation, header/data row rendering,
and text preparation for Word document generation.
"""
import logging

from django.core.files.storage import default_storage

from mediafiles.services import ImageService

from .utils import is_valid_image_path, format_field_value, humanize_label
from .column_spec import get_column_spec, is_nowrap_column, classify_column

logger = logging.getLogger(__name__)


class WordTablesMixin:
    """Mixin providing data table construction and row rendering methods."""

    def _get_image_render_width_cm(self, cards_list, field_name):
        """Compute rendered image width at IMAGE_HEIGHT_CM from first valid image.
        
        Used to size image columns to: rendered_width + 0.1 cm.
        Falls back to 3:4 portrait ratio (standard ID photo) if no image found.
        """
        DEFAULT_RATIO = 0.75  # 3:4 portrait
        try:
            from PIL import Image as PILImage
            for card in cards_list[:10]:
                img_path = ImageService.get_image_path_for_export(
                    card=card, field_name=field_name,
                    prefer_thumbnail=False, fallback_to_field_data=True
                )
                if img_path and is_valid_image_path(img_path):
                    try:
                        with default_storage.open(img_path, 'rb') as f:
                            with PILImage.open(f) as img:
                                w, h = img.size
                                if h > 0:
                                    return self.IMAGE_HEIGHT_CM * (w / h)
                    except Exception:
                        pass
        except Exception:
            pass
        return self.IMAGE_HEIGHT_CM * DEFAULT_RATIO
    
    def _calculate_column_widths(self, ordered_fields, cards, num_cols, cancel_check=None):
        """Calculate optimal column widths based on column_spec intelligence + data.

        Uses ``column_spec`` semantic min/max bounds per field category,
        combined with P90 content length for proportional distribution.

        Algorithm:
        - Image columns: fixed width from subtype dimensions + 0.1 cm.
        - Text columns: share remaining page width proportionally,
          clamped by word_min_cm / word_max_cm from column_spec.
        """
        # Step 1: Compute fixed widths for image columns
        image_widths = {}
        for idx, field in enumerate(ordered_fields):
            if field['is_image']:
                render_w = field.get('image_width_cm', self.IMAGE_DEFAULT_WIDTH_CM)
                image_widths[1 + idx] = render_w + 0.1

        # Step 2: Remaining page width for text + Sr No columns
        used_by_images = sum(image_widths.values())

        # Guard: prevent image columns from crowding out text columns.
        # Each text column needs at least 1.0 cm; if images would leave
        # less than that, scale all image widths down proportionally.
        num_text_cols = num_cols - len(image_widths)
        min_text_budget = num_text_cols * 1.0  # 1 cm floor per text col
        if image_widths and (self.PAGE_WIDTH_CM - used_by_images) < min_text_budget:
            max_image_budget = max(
                self.PAGE_WIDTH_CM - min_text_budget,
                len(image_widths) * 0.8,  # never shrink below 0.8 cm per image
            )
            _scale = max_image_budget / used_by_images
            image_widths = {k: round(v * _scale, 2) for k, v in image_widths.items()}
            used_by_images = sum(image_widths.values())

        remaining = max(self.PAGE_WIDTH_CM - used_by_images, min_text_budget)

        # Step 3: Collect P90 value lengths per text column and compute weights
        text_weights = {}
        text_specs = {}

        # Sr No column
        sr_spec = get_column_spec('SR NO')
        text_weights[0] = max(sr_spec.pref_chars, 4)
        text_specs[0] = sr_spec

        for idx, field in enumerate(ordered_fields):
            if not field['is_image']:
                name = field['name']
                ftype = field.get('type', 'text')
                spec = get_column_spec(name, ftype)
                text_specs[1 + idx] = spec

                lengths = [len(name)]  # header length as baseline
                for card in cards:
                    if callable(cancel_check) and cancel_check():
                        return None
                    fd = card.field_data or {}
                    val = str(fd.get(name, '') or '')
                    if val:
                        lengths.append(len(val))

                # 90th percentile
                lengths.sort()
                p90_idx = max(0, int(len(lengths) * 0.9) - 1)
                representative = lengths[p90_idx] if lengths else spec.pref_chars

                # Clamp to spec char range
                representative = max(representative, spec.min_chars)
                if spec.max_chars > 0:
                    representative = min(representative, spec.max_chars)

                # Safe text-column boost: verbose categories tend to need
                # slightly more room for clean wrapping.
                verbose_weight_boost = {
                    'full_name': 1.20,
                    'parent_name': 1.15,
                    'guardian_name': 1.15,
                    'spouse_name': 1.15,
                    'reporting_manager': 1.15,
                    'address': 1.30,
                    'email': 1.10,
                    'allergies': 1.20,
                    'medical_condition': 1.20,
                    'department': 1.08,
                    'designation': 1.08,
                    'course': 1.08,
                    'branch': 1.08,
                }
                boost = verbose_weight_boost.get(spec.category, 1.0)
                # Global long-text boost: for any wrapping column with large
                # content envelope, keep more horizontal room under dense
                # layouts so values stay readable.
                if spec.wrap and spec.max_chars >= 20:
                    boost = max(boost, 1.18)
                representative = max(spec.min_chars, int(round(representative * boost)))
                if spec.max_chars > 0:
                    representative = min(representative, spec.max_chars)

                # Density score: if data is significantly longer than the
                # category's preferred baseline, grant limited extra weight.
                density = representative / max(float(spec.pref_chars or 1), 1.0)
                density = max(0.85, min(density, 1.60))
                representative = int(round(representative * density))

                # ── Longest-single-word floor ─────────────────────────────
                # Min column width must fit the longest unbreakable word so
                # text is never clipped mid-character in a fixed-layout table.
                max_word_len = max((len(w) for w in name.split()), default=len(name))
                for _card in cards[:min(len(cards), 100)]:
                    _fd = _card.field_data or {}
                    _val = str(_fd.get(name, '') or '').strip()
                    for _word in _val.split():
                        if len(_word) > max_word_len:
                            max_word_len = len(_word)
                if spec.max_chars > 0:
                    max_word_len = min(max_word_len, spec.max_chars)
                representative = max(representative, max_word_len)

                text_weights[1 + idx] = max(representative, 3)

        total_text_w = sum(text_weights.values()) or 1

        # ── Dense-table word_max_cm overrides (>20 data columns) ──────
        # Tighten wide-column caps when the table is dense so narrow
        # categorical columns are not squeezed out.
        _DENSE_WORD_MAX: dict = {
            'full_name': 4.5,
            'parent_name': 4.0,
            'guardian_name': 4.0,
            'spouse_name': 4.0,
            'reporting_manager': 3.5,
            'email': 3.5,
            # Keep address readable in dense exports while still bounded.
            'address': 3.2,
            'allergies': 2.8,
            'medical_condition': 2.8,
            'department': 3.0,
            'designation': 3.0,
            'course': 2.8,
            'branch': 2.8,
        }
        _dense_word = num_cols > 20  # >20 data columns (includes Sr No)

        # Step 4: Build final widths with bounded allocation.
        # Goal: keep total width inside page margins while avoiding extremely
        # narrow text columns that cause ugly wrapping/clipping.
        column_widths = {}
        text_min_cm = {}
        text_target_cm = {}
        text_max_cm = {}

        for col_idx in range(num_cols):
            if col_idx in image_widths:
                column_widths[col_idx] = image_widths[col_idx]
            elif col_idx in text_weights:
                raw_cm = (text_weights[col_idx] / total_text_w) * remaining
                spec = text_specs.get(col_idx, sr_spec)
                # Effective max: tighter cap in dense-table mode
                eff_max_cm = spec.word_max_cm
                if _dense_word and spec.category in _DENSE_WORD_MAX:
                    eff_max_cm = min(eff_max_cm, _DENSE_WORD_MAX[spec.category])
                elif _dense_word and spec.wrap:
                    # Generic dense-mode cap for any long wrapping text field.
                    eff_max_cm = min(eff_max_cm, max(spec.word_min_cm * 2.0, 2.8))
                elif spec.category in {
                    'full_name', 'parent_name', 'guardian_name', 'spouse_name',
                    'reporting_manager', 'address', 'email', 'allergies',
                    'medical_condition',
                }:
                    # Sparse/medium tables can afford a little more width for
                    # verbose text columns before wrapping hard.
                    eff_max_cm = min(spec.word_max_cm * 1.15, spec.word_max_cm + 0.5)

                min_cm = max(0.6, spec.word_min_cm)
                target_cm = max(min_cm, min(raw_cm, eff_max_cm))

                text_min_cm[col_idx] = min_cm
                text_target_cm[col_idx] = target_cm
                text_max_cm[col_idx] = max(target_cm, eff_max_cm)
            else:
                column_widths[col_idx] = 1.5

        text_keys = list(text_weights.keys())
        if text_keys:
            allocated = {k: text_min_cm.get(k, 0.8) for k in text_keys}
            min_total = sum(allocated.values())

            if min_total > remaining:
                # Hard fallback: if even all mins do not fit, compress to a
                # readable floor per column while still respecting page width.
                floor_cm = {}
                for k in text_keys:
                    spec = text_specs.get(k, sr_spec)
                    if k == 0:
                        base_floor = 0.6  # Sr No can be the narrowest
                    elif spec.wrap:
                        base_floor = 0.75
                    else:
                        base_floor = 0.65
                    floor_cm[k] = min(allocated[k], base_floor)

                floor_total = sum(floor_cm.values())
                if floor_total > remaining and floor_total > 0:
                    scale = remaining / floor_total
                    allocated = {k: floor_cm[k] * scale for k in text_keys}
                else:
                    allocated = dict(floor_cm)
                    extra = max(0.0, remaining - floor_total)
                    grow_need = {k: max(text_min_cm[k] - allocated[k], 0.0) for k in text_keys}
                    need_total = sum(grow_need.values())
                    if extra > 0 and need_total > 0:
                        for k in text_keys:
                            allocated[k] += extra * (grow_need[k] / need_total)
                    elif extra > 0:
                        even_add = extra / len(text_keys)
                        for k in text_keys:
                            allocated[k] += even_add
            else:
                extra = max(0.0, remaining - min_total)

                def _distribute(limit_map, extra_amount):
                    if extra_amount <= 1e-9:
                        return 0.0
                    needs = {k: max(limit_map[k] - allocated[k], 0.0) for k in text_keys}
                    need_total = sum(needs.values())
                    if need_total <= 1e-9:
                        return extra_amount
                    used = 0.0
                    for k in text_keys:
                        need = needs[k]
                        if need <= 0:
                            continue
                        add = extra_amount * (need / need_total)
                        add = min(add, need)
                        allocated[k] += add
                        used += add
                    return max(0.0, extra_amount - used)

                # Pass 1: move from min -> target
                extra = _distribute(text_target_cm, extra)
                # Pass 2: if room remains, move from target -> max
                extra = _distribute(text_max_cm, extra)
                # Pass 3: final even spread (rare rounding leftovers)
                if extra > 1e-6 and text_keys:
                    even_add = extra / len(text_keys)
                    for k in text_keys:
                        allocated[k] += even_add

            rounded = {k: round(max(0.35, allocated[k]), 2) for k in text_keys}
            drift = round(remaining - sum(rounded.values()), 2)
            if abs(drift) >= 0.01 and text_keys:
                adj_key = max(rounded, key=lambda x: rounded[x])
                rounded[adj_key] = round(max(0.35, rounded[adj_key] + drift), 2)

            for k in text_keys:
                column_widths[k] = rounded[k]

        return column_widths

    def _enforce_fixed_table_layout(self, table_obj, column_widths, num_cols, Cm):
        """Force fixed table layout with explicit column grid definitions.

        Without ``w:tblLayout type="fixed"`` and a ``w:tblGrid``, Word
        auto-resizes columns so images overflow their cell boundaries.
        This must be called once immediately after the table is created.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tbl = table_obj._tbl

        # ── 1. tblPr: set tblLayout=fixed and total table width ─────
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Remove stale tblLayout / tblW before re-adding
        for tag in ('w:tblLayout', 'w:tblW'):
            for el in tblPr.findall(qn(tag)):
                tblPr.remove(el)

        # fixed layout — Word must honour explicit column widths
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # Total table width in twips (dxa)
        total_cm = sum(column_widths.get(i, 2.0) for i in range(num_cols))
        total_twips = int(Cm(total_cm).twips)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(total_twips))
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # ── 2. tblGrid: one gridCol per column (twips) ───────────────
        for existing in tbl.findall(qn('w:tblGrid')):
            tbl.remove(existing)

        tblGrid = OxmlElement('w:tblGrid')
        for col_idx in range(num_cols):
            w_cm = column_widths.get(col_idx, 2.0)
            w_twips = int(Cm(w_cm).twips)
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(w_twips))
            tblGrid.append(gridCol)

        # tblGrid must come directly after tblPr in OOXML order
        tbl_children = list(tbl)
        tblPr_idx = tbl_children.index(tblPr)
        tbl.insert(tblPr_idx + 1, tblGrid)

    @classmethod
    def _is_nowrap_field_word(cls, field_name: str) -> bool:
        """Check if a field contains non-wrappable data (phone, DOB, etc.).

        Delegates to column_spec intelligence.
        """
        return is_nowrap_column(field_name)

    def _prefetch_image_bytes_batch(self, cards_batch, image_field_names, cancel_check=None):
        """Pre-fetch original image bytes for a batch of cards using 5 concurrent threads.

        Opens up to 5 cloud-storage connections simultaneously instead of
        serialising every read. For 826 cards × 3 photo fields the effective
        IO wait drops from ~N×latency to ~ceil(N/5)×latency.

        Args:
            cards_batch:       Slice of cards_list to pre-fetch for.
            image_field_names: List of image field names to fetch per card.
            cancel_check:      Optional callable; aborts remaining work if True.

        Returns:
            dict mapping (card.id, field_name) → bytes or None.
            None means the image was not found / could not be read.
        """
        from concurrent.futures import ThreadPoolExecutor, as_completed

        result = {}
        if not cards_batch or not image_field_names:
            return result

        def _fetch_one(card, field_name):
            card_id = getattr(card, 'id', None) or id(card)
            if callable(cancel_check) and cancel_check():
                return (card_id, field_name), None
            try:
                img_path = ImageService.get_image_path_for_export(
                    card=card,
                    field_name=field_name,
                    prefer_thumbnail=False,   # always fetch full original
                    fallback_to_field_data=True,
                )
                if not img_path or not is_valid_image_path(img_path):
                    return (card_id, field_name), None
                with default_storage.open(img_path, 'rb') as fh:
                    data = fh.read()
                if data and len(data) >= 100:
                    return (card_id, field_name), data
                return (card_id, field_name), None
            except Exception as exc:
                logger.debug(
                    'DOCX prefetch failed card=%s field=%s: %s',
                    card_id, field_name, exc,
                )
                return (card_id, field_name), None

        # Submit all (card, field) pairs; executor runs at most 5 at a time
        with ThreadPoolExecutor(max_workers=5) as pool:
            futures = {
                pool.submit(_fetch_one, card, field_name): (getattr(card, 'id', None) or id(card), field_name)
                for card in cards_batch
                for field_name in image_field_names
            }
            for future in as_completed(futures):
                try:
                    key, data = future.result(timeout=60)
                    result[key] = data
                except Exception as exc:
                    result[futures[future]] = None
                    logger.debug('DOCX prefetch future error %s: %s', futures[future], exc)

        return result

    
    def _create_data_tables(self, doc, cards_list, ordered_fields, column_widths,
                            num_cols, Cm, Pt, RGBColor, WD_TABLE_ALIGNMENT,
                            WD_ALIGN_PARAGRAPH, parse_xml, nsdecls, OxmlElement,
                            qn, Image, ImageOps, class_field_name=None,
                            section_field_name=None, break_mode='none',
                            progress_callback=None, cancel_check=None):
        """Create ONE continuous table with all card data.

        The table is never split into separate tables, so selecting a
        row in Word lets you extend the selection to the very last row.

        Page-break rules (inserted inside the table via pageBreakBefore):
          1. Every ENTRIES_PER_PAGE rows → page break before next row
          2. When class/section value changes (according to break_mode) → force page break
             even if current page has room

        The column-heading row appears on the first page ONLY.
        """
        sr_no = 1
        rows_on_current_page = 0
        prev_group_key = None

        # Pre-compute fixed image dimensions per image field (from subtype)
        image_fixed_widths = {}
        image_fixed_heights = {}
        max_image_height = 0
        for field in ordered_fields:
            if field['is_image']:
                w = field.get('image_width_cm', self.IMAGE_DEFAULT_WIDTH_CM)
                h = field.get('image_height_cm', self.IMAGE_HEIGHT_CM)
                image_fixed_widths[field['name']] = w
                image_fixed_heights[field['name']] = h
                max_image_height = max(max_image_height, h)
        # Row height policy:
        # - With image columns, keep a consistent visual baseline (2.5 cm min)
        #   but allow rows to expand for long wrapped text (e.g., address).
        # - Without image columns, keep compact text baseline.
        if max_image_height > 0:
            row_height_cm = max(self.ROW_HEIGHT_CM, round(max_image_height, 2))
        else:
            row_height_cm = 0.8
        _row_h_rule = 'atLeast'

        # Create ONE table with a header row
        table_obj = doc.add_table(rows=1, cols=num_cols)
        table_obj.style = 'Table Grid'
        table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._set_table_borders(table_obj, parse_xml, nsdecls)
        # CRITICAL: enforce fixed layout + explicit column grid so images
        # cannot overflow their cell and column widths are respected.
        self._enforce_fixed_table_layout(table_obj, column_widths, num_cols, Cm)

        # Auto-select font size: >18 cols → 7pt, >15 cols → 8pt, else 9pt
        _font_pt = 7 if num_cols > 18 else (8 if num_cols > 15 else 9)
        _font_pt = max(7, _font_pt)

        # Style the header row and mark it to repeat on every page
        self._style_header_row(
            table_obj.rows[0].cells, ordered_fields, column_widths,
            Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
            font_pt=_font_pt
        )

        # Mark header row to repeat at top of each page when table spans pages
        header_row = table_obj.rows[0]
        trPr = header_row._tr.get_or_add_trPr()
        tblHeader = parse_xml(r'<w:tblHeader {} />'.format(nsdecls('w')))
        trPr.append(tblHeader)

        # Collect image field names once for the pre-fetch helper.
        _image_field_names = [f['name'] for f in ordered_fields if f['is_image']]

        # Process cards in chunks.  Before rendering each chunk the image bytes
        # are fetched concurrently (5 threads), collapsing N serial cloud-storage
        # round-trips into ~N/5 parallel ones.  Rendering stays strictly serial
        # because python-docx is not thread-safe for writes.
        _PREFETCH_CHUNK = 25

        for _chunk_start in range(0, len(cards_list), _PREFETCH_CHUNK):
            _chunk = cards_list[_chunk_start:_chunk_start + _PREFETCH_CHUNK]

            if callable(cancel_check) and cancel_check():
                return False

            # ------ concurrent image pre-fetch (5 threads) ------
            _image_cache = {}
            if _image_field_names:
                _image_cache = self._prefetch_image_bytes_batch(
                    _chunk, _image_field_names, cancel_check=cancel_check,
                )
            # ----------------------------------------------------

            for _idx_in_chunk, card in enumerate(_chunk):
                card_idx = _chunk_start + _idx_in_chunk

                if callable(cancel_check) and cancel_check():
                    return False
                fd = card.field_data or {}
                
                # Check resolved break_mode
                resolved_break_mode = str(break_mode or 'none').strip().lower()
                if resolved_break_mode not in ('class_only', 'class_section', 'none'):
                    resolved_break_mode = 'none'

                cur_class_val = (
                    str(fd.get(class_field_name, '') or '').strip().upper()
                    if class_field_name and resolved_break_mode in ('class_only', 'class_section') else None
                )
                cur_section_val = (
                    str(fd.get(section_field_name, '') or '').strip().upper()
                    if section_field_name and resolved_break_mode == 'class_section' else None
                )
                cur_group_key = (cur_class_val, cur_section_val)

                # Decide whether to insert a page break before this row.
                need_page_break = False
                if rows_on_current_page >= self.ENTRIES_PER_PAGE:
                    need_page_break = True
                elif resolved_break_mode != 'none' and prev_group_key is not None and cur_group_key != prev_group_key:
                    need_page_break = True

                # Add the data row (images served from pre-fetched cache)
                self._add_data_row(
                    table_obj, card, ordered_fields, column_widths, sr_no,
                    Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                    Image, ImageOps, image_fixed_widths=image_fixed_widths,
                    image_fixed_heights=image_fixed_heights, row_height_cm=row_height_cm,
                    font_pt=_font_pt, h_rule=_row_h_rule, cancel_check=cancel_check,
                    image_cache=_image_cache,
                )

                # If a page break is needed, set it on the FIRST paragraph
                # of the first cell of the row we just added.
                if need_page_break:
                    new_row = table_obj.rows[-1]
                    first_para = new_row.cells[0].paragraphs[0]
                    pPr = first_para._p.get_or_add_pPr()
                    pPr.append(parse_xml(
                        r'<w:pageBreakBefore {} />'.format(nsdecls('w'))
                    ))
                    rows_on_current_page = 0

                sr_no += 1
                rows_on_current_page += 1
                prev_group_key = cur_group_key

                if callable(progress_callback):
                    # Keep callback cadence light to avoid excessive DB writes.
                    processed = card_idx + 1
                    if (processed % 10 == 0) or (processed == len(cards_list)):
                        try:
                            progress_callback(processed, len(cards_list))
                        except Exception:
                            pass

        return True
    
    def _style_header_row(self, cells, ordered_fields, column_widths,
                          Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                          font_pt=9):
        """Style the header row of a table."""
        col_idx = 0
        
        # Sr No header
        cells[col_idx].text = 'Sr No.'
        self._style_header_cell(cells[col_idx], column_widths[col_idx],
                                Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                                font_pt=font_pt)
        col_idx += 1
        
        # Field headers
        for field in ordered_fields:
            _label = humanize_label(field['name'].upper())
            cells[col_idx].text = _label
            self._style_header_cell(cells[col_idx], column_widths[col_idx],
                                    Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                                    font_pt=font_pt)
            col_idx += 1
    
    def _style_header_cell(self, cell, width, Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH,
                           parse_xml, nsdecls, font_pt=9):
        """Apply styling to a header cell with wrapping and padding."""
        para = cell.paragraphs[0]
        header_text = para.runs[0].text if para.runs else cell.text
        if para.runs:
            run = para.runs[0]
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(font_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 1px-equivalent inner padding for compact header layout
        self._set_cell_margins(cell, parse_xml, nsdecls, 15, 15, 15, 15)
        self._set_cell_vertical_align(cell, parse_xml, nsdecls)
        self._set_para_spacing(para, parse_xml, nsdecls)
        cell.width = Cm(width)
        from docx.oxml.ns import qn as _qn
        tcPr = cell._tc.get_or_add_tcPr()
        # Always allow wrapping in header cells so text cannot overflow outside
        # fixed column boundaries in dense tables.
        for nw in tcPr.findall(_qn('w:noWrap')):
            tcPr.remove(nw)
    
    def _add_data_row(self, table, card, ordered_fields, column_widths, sr_no,
                      Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                      Image, ImageOps, image_fixed_widths=None,
                      image_fixed_heights=None, row_height_cm=None, font_pt=9,
                      h_rule='atLeast', cancel_check=None, image_cache=None):
        """Add a data row to the table.

        Args:
            image_cache: dict from _prefetch_image_bytes_batch mapping
                (card.id, field_name) → bytes-or-None.  When an entry exists
                the storage read inside _add_image_to_cell is skipped entirely.
        """
        new_row = table.add_row()
        cells = new_row.cells
        
        # Set row height — "atLeast" so text can wrap and expand if needed
        tr = new_row._tr
        trPr = tr.get_or_add_trPr()
        effective_row_h = row_height_cm if row_height_cm else self.ROW_HEIGHT_CM
        row_height_twips = int(Cm(effective_row_h).twips)
        trHeight = parse_xml(
            r'<w:trHeight {} w:val="{}" w:hRule="{}"/>'.format(nsdecls('w'), row_height_twips, h_rule)
        )
        trPr.append(trHeight)
        
        col_idx = 0
        field_data = card.field_data or {}
        if image_fixed_widths is None:
            image_fixed_widths = {}
        if image_fixed_heights is None:
            image_fixed_heights = {}
        
        # Sr No
        cells[col_idx].text = str(sr_no)
        self._style_data_cell(cells[col_idx], column_widths[col_idx], False,
                              Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                              font_pt=font_pt)
        cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        col_idx += 1
        
        # Field values
        for field in ordered_fields:
            cell = cells[col_idx]
            cell.width = Cm(column_widths[col_idx])
            
            if field['is_image']:
                img_fixed_w = image_fixed_widths.get(field['name'], self.IMAGE_DEFAULT_WIDTH_CM)
                img_fixed_h = image_fixed_heights.get(field['name'], self.IMAGE_HEIGHT_CM)

                # Check whether bytes were pre-fetched for this (card, field) pair.
                # If so, skip the storage read entirely in _add_image_to_cell.
                card_id = getattr(card, 'id', None) or id(card)
                _cache_key = (card_id, field['name'])
                _in_cache = image_cache is not None and _cache_key in image_cache
                _prefetched = image_cache[_cache_key] if _in_cache else None

                # Only look up the path when the cache missed (fallback / logging).
                if _in_cache:
                    image_path = ''   # bytes already available; path not needed
                else:
                    image_path = ImageService.get_image_path_for_export(
                        card=card,
                        field_name=field['name'],
                        prefer_thumbnail=False,   # always full originals for print
                        fallback_to_field_data=True,
                    ) or ''

                self._add_image_to_cell(
                    cell, image_path,
                    Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                    Image, ImageOps, fixed_width_cm=img_fixed_w,
                    fixed_height_cm=img_fixed_h,
                    image_subtype=field.get('image_subtype'),
                    field_name=field.get('name'),
                    cancel_check=cancel_check,
                    prefetched_bytes=_prefetched,
                    skip_storage_read=_in_cache,
                )
            else:
                value = format_field_value(field_data.get(field['name'], ''), uppercase=True)
                value = self._prepare_text_for_word(value)
                cell.text = value
                self._style_data_cell(cell, column_widths[col_idx], False,
                                      Cm, Pt, RGBColor, WD_ALIGN_PARAGRAPH, parse_xml, nsdecls,
                                      font_pt=font_pt)
            
            col_idx += 1
    
    @staticmethod
    def _prepare_text_for_word(text: str) -> str:
        """Insert zero-width spaces in very long unbreakable words.

        Rules:
          1. Words <= 12 chars: leave untouched.
          2. Words > 12 chars with no natural break opportunity: insert
              U+200B (ZERO-WIDTH SPACE) every 10 chars so Word can wrap
           the text without inserting ANY visible character or dash.
          3. Natural separators (, / ; : - @ . _) already act as break points.
        4. NEVER insert U+00AD (soft hyphen) — it can render as a dash
           when the text is copied from the document.
        """
        import re as _re
        if not text or len(text) <= 14:
            return text

        ZWSP = '\u200B'  # zero-width space — breaks without any visible char

        def _insert_zwsp(word):
            """Insert ZWSP every 10 chars in a long unbroken word."""
            if len(word) <= 12:
                return word
            parts = [word[i:i+10] for i in range(0, len(word), 10)]
            return ZWSP.join(parts)

        tokens = text.split(' ')
        processed = []
        for token in tokens:
            if not token or len(token) <= 12:
                processed.append(token)
                continue
            # Split on natural separators first; process each sub-part
            sub_parts = _re.split(r'([,/;:\-@._])', token)
            result = []
            for sp in sub_parts:
                if sp in (',', '/', ';', ':', '-', '@', '.', '_'):
                    result.append(sp)
                elif len(sp) > 12:
                    result.append(_insert_zwsp(sp))
                else:
                    result.append(sp)
            processed.append(''.join(result))
        return ' '.join(processed)
